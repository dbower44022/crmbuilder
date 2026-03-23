"""DOCX renderer for the documentation generator."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tools.docgen.models import DocDocument, DocParagraph, DocSection, DocTable

# Colors
HEADING_COLOR = RGBColor(0x1F, 0x49, 0x7D)
TABLE_HEADER_BG = "1F497D"
TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_ALT_ROW_BG = "F2F2F2"
STATUS_BOX_BG = "DAE8FC"
STATUS_BOX_BORDER = "6C8EBF"


def render(doc: DocDocument, output_path: Path) -> None:
    """Render a DocDocument to a .docx file.

    :param doc: Document to render.
    :param output_path: Output file path.
    """
    document = Document()

    # Page setup - US Letter
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set default font
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 5):
        heading_style = document.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri Light"
        heading_style.font.color.rgb = HEADING_COLOR
        if level == 1:
            heading_style.font.size = Pt(20)
        elif level == 2:
            heading_style.font.size = Pt(16)
        elif level == 3:
            heading_style.font.size = Pt(13)
        else:
            heading_style.font.size = Pt(11)

    # Title page
    _add_title_page(document, doc)

    # Table of contents placeholder
    _add_toc(document)

    # Sections
    for sec in doc.sections:
        _render_section(document, sec)

    # Add page numbers in footer
    _add_page_numbers(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def _add_title_page(document: Document, doc: DocDocument) -> None:
    """Add the title page."""
    # Spacer
    for _ in range(4):
        document.add_paragraph("")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cleveland Business Mentors")
    run.font.size = Pt(28)
    run.font.color.rgb = HEADING_COLOR
    run.font.name = "Calibri Light"

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(doc.subtitle)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_paragraph("")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Generated from YAML program files\n").font.size = Pt(11)
    meta.add_run(f"Version: {doc.version}\n").font.size = Pt(11)
    meta.add_run(f"Generated: {doc.timestamp}").font.size = Pt(11)

    document.add_paragraph("")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This document defines the EspoCRM configuration required to support "
        "the requirements specified in the CBM PRD documents. It is generated "
        "automatically from the YAML program files and must not be edited manually."
    )
    run.font.size = Pt(10)
    run.font.italic = True

    document.add_page_break()


def _add_toc(document: Document) -> None:
    """Add a Table of Contents field."""
    document.add_heading("Table of Contents", level=1)
    para = document.add_paragraph()
    run = para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)

    para.add_run(
        "\n(Right-click and select 'Update Field' to populate "
        "the table of contents.)"
    ).font.italic = True

    document.add_page_break()


def _render_section(document: Document, section: DocSection) -> None:
    """Render a section recursively."""
    level = min(section.level, 4)
    document.add_heading(section.title, level=level)

    for item in section.content:
        if isinstance(item, DocSection):
            _render_section(document, item)
        elif isinstance(item, DocTable):
            _render_table(document, item)
        elif isinstance(item, DocParagraph):
            _render_paragraph(document, item)


def _render_table(document: Document, table: DocTable) -> None:
    """Render a table."""
    if table.caption:
        cap = document.add_paragraph()
        run = cap.add_run(table.caption)
        run.font.italic = True
        run.font.size = Pt(10)

    num_cols = len(table.headers)
    doc_table = document.add_table(
        rows=1 + len(table.rows), cols=num_cols
    )
    doc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    doc_table.style = "Table Grid"

    # Header row
    header_row = doc_table.rows[0]
    for i, header in enumerate(table.headers):
        cell = header_row.cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = TABLE_HEADER_TEXT
                run.font.size = Pt(10)
        _set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(table.rows):
        doc_row = doc_table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = doc_row.cells[col_idx]
                # Strip backtick formatting for docx
                clean = cell_text.replace("`", "")
                cell.text = clean
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Alternate row shading
        if row_idx % 2 == 1:
            for col_idx in range(num_cols):
                if col_idx < len(doc_row.cells):
                    _set_cell_shading(
                        doc_row.cells[col_idx], TABLE_ALT_ROW_BG
                    )

    document.add_paragraph("")


def _render_paragraph(document: Document, para: DocParagraph) -> None:
    """Render a paragraph."""
    if para.style == "status":
        _add_status_box(document, para.text)
    elif para.style == "note":
        p = document.add_paragraph()
        run = p.add_run(f"Note: {para.text}")
        run.font.italic = True
        run.font.size = Pt(10)
    elif para.style == "code":
        p = document.add_paragraph()
        run = p.add_run(para.text)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    else:
        # Handle bold markers **text**
        text = para.text
        if text.startswith("**") and "**" in text[2:]:
            end = text.index("**", 2)
            bold_text = text[2:end]
            rest = text[end + 2:]
            p = document.add_paragraph()
            run = p.add_run(bold_text)
            run.font.bold = True
            if rest:
                p.add_run(rest)
        else:
            document.add_paragraph(text)


def _add_status_box(document: Document, text: str) -> None:
    """Add a status info box (light blue shaded single-cell table)."""
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(11)
    _set_cell_shading(cell, STATUS_BOX_BG)
    document.add_paragraph("")


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set the background shading of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_page_numbers(document: Document) -> None:
    """Add centered page numbers to the footer."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
