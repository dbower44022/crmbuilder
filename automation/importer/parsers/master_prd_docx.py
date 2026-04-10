"""Parse CBM-Master-PRD.docx into a Path B envelope JSON string.

Adapter that turns a Word document into a Path B envelope suitable for
processing by the existing ImportProcessor seven-stage pipeline.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

from automation.importer.parsers import MasterPrdParseError, ParseReport

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Separator: em-dash (U+2014), en-dash (U+2013), or hyphen-minus (U+002D).
# Hyphen must be last in the character class to avoid range interpretation.
_DASH = r"[—–-]"

# Section 2: persona headings (Heading 3)
_PERSONA_RE = re.compile(rf"^(MST-PER-\d+)\s*{_DASH}\s*(.+)$")

# Section 3: domain headings (Heading 2)
_DOMAIN_H2_RE = re.compile(
    rf"^\d+(?:\.\d+)?\s+MST-DOM-\d+\s*{_DASH}\s*(.+?)\s*\(([A-Z]+)\)\s*$"
)

# Section 4: service headings — v2.6 format with MST-SVC code (Heading 2)
_SERVICE_V26_RE = re.compile(
    rf"^\d+(?:\.\d+)?\s+MST-SVC-\d+\s*{_DASH}\s*(.+?)\s*\(([A-Z]+)\)\s*$"
)

# Section 4: service headings — v2.3 fallback, no MST-SVC code (Heading 2)
_SERVICE_V23_RE = re.compile(r"^\d+(?:\.\d+)?\s+(.+)$")

# Organization overview H2 subsection prefix
_ORG_H2_RE = re.compile(r"^1\.\d+\s+(.+)$")

# Code + name line in prose sections (e.g. "MN-INTAKE — Client Intake")
_CODE_NAME_RE = re.compile(rf"^([A-Z]+(?:-[A-Z]+)+)\s*{_DASH}\s*(.+)$")

# Table 2 domain column: extract parenthesized code
_PAREN_CODE_RE = re.compile(r"\(([A-Z]+)\)")

# Expected header cells for Table 2 (lowercased, trimmed)
_TABLE2_HEADERS = ["code", "process / sub-domain", "domain", "tier"]

_VALID_TIERS = frozenset({"core", "important", "enhancement"})


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _style_name(p: Any) -> str:
    """Return the paragraph's style name, or '' if unavailable."""
    if p.style is None:
        return ""
    name = p.style.name
    return name if name else ""


def _is_heading(p: Any, level: int | None = None) -> bool:
    """Check whether paragraph *p* is a Word heading.

    :param level: If given, match exactly that heading level.
    """
    name = _style_name(p)
    if level is not None:
        return name == f"Heading {level}"
    return name.startswith("Heading ")


def _find_heading1(paragraphs: list, pattern: str) -> int | None:
    """Return the index of the first Heading 1 matching *pattern*."""
    regex = re.compile(pattern)
    for i, p in enumerate(paragraphs):
        if _is_heading(p, 1) and regex.match(p.text.strip()):
            return i
    return None


def _domain_section_range(
    paragraphs: list, domain_code: str,
) -> tuple[int | None, int | None]:
    """Return (start, end) paragraph indices for a domain's Heading 2 section."""
    for i, p in enumerate(paragraphs):
        if _is_heading(p, 2):
            m = _DOMAIN_H2_RE.match(p.text.strip())
            if m and m.group(2) == domain_code:
                end = len(paragraphs)
                for j in range(i + 1, len(paragraphs)):
                    if _is_heading(paragraphs[j], 2) or _is_heading(paragraphs[j], 1):
                        end = j
                        break
                return i, end
    return None, None


def _next_plain_text(paragraphs: list, start: int, end: int) -> str:
    """Return the text of the next non-empty, non-heading, non-code paragraph.

    Stops at headings and at the next code-name line to avoid leaking
    a description from a neighboring entry.
    """
    for j in range(start, end):
        p = paragraphs[j]
        if _is_heading(p):
            return ""
        text = p.text.strip()
        if not text:
            continue
        if _CODE_NAME_RE.match(text):
            return ""  # hit next code line — no description for this entry
        return text
    return ""


# ---------------------------------------------------------------------------
# Section parsers
# ---------------------------------------------------------------------------


def _parse_org_overview(paragraphs: list, report: ParseReport) -> str:
    """Section 4.1 of spec — Organization Overview (Section 1 of source)."""
    h1 = _find_heading1(paragraphs, r"^1\.\s+Organization Overview$")
    if h1 is None:
        report.warn(
            "missing_section", "Organization Overview",
            "No Heading 1 for Section 1",
        )
        return ""

    subsections: list[str] = []
    i = h1 + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading(p, 1):
            break
        if _is_heading(p, 2):
            m = _ORG_H2_RE.match(p.text.strip())
            if m:
                title = m.group(1).strip()
                lines: list[str] = []
                i += 1
                while i < len(paragraphs):
                    pp = paragraphs[i]
                    if _is_heading(pp, 2) or _is_heading(pp, 1):
                        break
                    text = pp.text.strip()
                    if text:
                        lines.append(text)
                    i += 1
                subsections.append(f"## {title}\n\n" + "\n\n".join(lines))
                continue
        i += 1

    if not subsections:
        report.warn(
            "missing_content", "Organization Overview",
            "Zero subsections found",
        )
        return ""

    return "\n\n".join(subsections)


def _parse_personas(paragraphs: list, report: ParseReport) -> list[dict]:
    """Section 4.2 of spec — Personas (Section 2 of source)."""
    h1 = _find_heading1(paragraphs, r"^2\.\s+Personas$")
    if h1 is None:
        raise MasterPrdParseError(
            "Document has no Heading 1 matching '^2. Personas$'"
        )

    personas: list[dict] = []
    i = h1 + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading(p, 1):
            break
        if _is_heading(p, 3):
            m = _PERSONA_RE.match(p.text.strip())
            if m:
                identifier = m.group(1)
                name = m.group(2).strip()
                desc_lines: list[str] = []
                i += 1
                while i < len(paragraphs):
                    pp = paragraphs[i]
                    if _is_heading(pp, 3) or _is_heading(pp, 1):
                        break
                    text = pp.text.strip()
                    if text:
                        desc_lines.append(text)
                    i += 1
                description = "\n\n".join(desc_lines)
                if not description:
                    report.warn(
                        "missing_description",
                        f"Persona {identifier}",
                        "Empty description",
                    )
                personas.append({
                    "identifier": identifier,
                    "name": name,
                    "description": description,
                })
                continue
        i += 1

    if not personas:
        raise MasterPrdParseError(
            "Zero personas parsed despite Section 2 existing"
        )
    return personas


def _parse_domains(paragraphs: list, report: ParseReport) -> list[dict]:
    """Section 4.3 of spec — Top-level Domains (Section 3 of source)."""
    h1 = _find_heading1(paragraphs, r"^3\.\s+Key Business Domains$")
    if h1 is None:
        raise MasterPrdParseError(
            "Document has no Heading 1 matching '^3. Key Business Domains$'"
        )

    domains: list[dict] = []
    sort_order = 0
    i = h1 + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading(p, 1):
            break
        if _is_heading(p, 2):
            m = _DOMAIN_H2_RE.match(p.text.strip())
            if m:
                sort_order += 1
                name = m.group(1).strip()
                code = m.group(2).strip()

                # Find "Domain Purpose" Heading 3
                description = ""
                j = i + 1
                while j < len(paragraphs):
                    pj = paragraphs[j]
                    if _is_heading(pj, 2) or _is_heading(pj, 1):
                        break
                    if _is_heading(pj, 3) and pj.text.strip() == "Domain Purpose":
                        j += 1
                        desc_lines: list[str] = []
                        while j < len(paragraphs):
                            pk = paragraphs[j]
                            if _is_heading(pk):
                                break
                            text = pk.text.strip()
                            if text:
                                desc_lines.append(text)
                            j += 1
                        description = "\n\n".join(desc_lines)
                        break
                    j += 1

                if not description:
                    report.warn(
                        "missing_description", f"Domain {code}",
                        "Missing Domain Purpose subsection",
                    )

                domains.append({
                    "name": name,
                    "code": code,
                    "description": description,
                    "sort_order": sort_order,
                    "is_service": False,
                    "sub_domains": [],
                })
        i += 1

    if not domains:
        raise MasterPrdParseError(
            "Zero top-level domains parsed from Section 3"
        )
    return domains


def _find_table2(tables: list) -> Any:
    """Locate and validate the Process Tier Summary table (Table 2)."""
    candidate_headers: list[str] | None = None
    for table in tables:
        if not table.rows:
            continue
        cells = [cell.text.strip() for cell in table.rows[0].cells]
        cells_lower = [c.lower() for c in cells]
        if cells_lower == _TABLE2_HEADERS:
            return table
        # Track partial match for a better error message
        if cells_lower and cells_lower[0] == "code":
            candidate_headers = cells

    if candidate_headers is not None:
        raise MasterPrdParseError(
            f"Process Tier Summary table has wrong column structure: "
            f"found {candidate_headers}, expected "
            f"['Code', 'Process / Sub-Domain', 'Domain', 'Tier']"
        )
    raise MasterPrdParseError("Process Tier Summary table not found")


def _parse_table2(
    table: Any, report: ParseReport,
) -> tuple[dict[str, list[dict]], list[dict]]:
    """Parse Table 2 into sub-domains (keyed by parent code) and processes."""
    sub_domains_by_parent: dict[str, list[dict]] = {}
    processes: list[dict] = []
    proc_sort = 0
    sd_sort_by_parent: dict[str, int] = {}

    for row_idx in range(1, len(table.rows)):
        cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
        if len(cells) < 4 or not cells[0]:
            continue
        code, name_cell, domain_cell, tier_cell = (
            cells[0], cells[1], cells[2], cells[3],
        )

        # Extract parent domain code from domain column
        m = _PAREN_CODE_RE.search(domain_cell)
        parent_code = m.group(1) if m else domain_cell.strip()

        if name_cell.rstrip().lower().endswith("(sub-domain)"):
            # Sub-domain row
            name = re.sub(
                r"\s*\((?i:sub-domain)\)\s*$", "", name_cell,
            ).strip()
            sd_sort_by_parent.setdefault(parent_code, 0)
            sd_sort_by_parent[parent_code] += 1
            sub_domains_by_parent.setdefault(parent_code, []).append({
                "name": name,
                "code": code,
                "description": "",
                "sort_order": sd_sort_by_parent[parent_code],
                "is_service": False,
            })
        else:
            # Process row
            proc_sort += 1
            tier_raw = tier_cell.strip()
            tier = tier_raw.lower()
            if tier not in _VALID_TIERS:
                report.warn(
                    "bad_tier", f"Process {code}",
                    f"Unexpected tier '{tier_raw}'",
                )
                tier = None
            processes.append({
                "name": name_cell.strip(),
                "code": code,
                "description": "",
                "sort_order": proc_sort,
                "tier": tier,
                "domain_code": parent_code,
            })

    if not processes:
        raise MasterPrdParseError("Zero processes parsed from Table 2")
    return sub_domains_by_parent, processes


def _parse_services(paragraphs: list, report: ParseReport) -> list[dict]:
    """Section 4.5 of spec — Cross-Domain Services (Section 4 of source)."""
    h1 = _find_heading1(paragraphs, r"^4\.\s+Cross-Domain Services$")
    if h1 is None:
        report.warn(
            "missing_section", "Cross-Domain Services",
            "Section 4 not found",
        )
        return []

    services: list[dict] = []
    sort_order = 100
    i = h1 + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading(p, 1):
            break
        if _is_heading(p, 2):
            text = p.text.strip()
            # Try v2.6 format first (MST-SVC-NNN — Name (CODE))
            m = _SERVICE_V26_RE.match(text)
            if m:
                name = m.group(1).strip()
                code = m.group(2).strip()
            else:
                # v2.3 fallback (N.N ServiceName)
                m = _SERVICE_V23_RE.match(text)
                if not m:
                    i += 1
                    continue
                raw_name = m.group(1).strip()
                name = raw_name
                # Derive code: strip " Service" suffix, uppercase
                if " Service" in raw_name:
                    code = raw_name.split(" Service")[0].strip().upper()
                else:
                    code = raw_name.upper().replace(" ", "")

            sort_order += 1

            # Find Purpose paragraph and collect description
            description = _extract_service_purpose(paragraphs, i + 1)
            if not description:
                report.warn(
                    "missing_description", f"Service {code}",
                    "Empty Purpose section",
                )

            services.append({
                "name": name,
                "code": code,
                "description": description,
                "sort_order": sort_order,
                "is_service": True,
                "sub_domains": [],
            })
        i += 1

    return services


def _extract_service_purpose(paragraphs: list, start: int) -> str:
    """Extract the Purpose text from a service section starting at *start*."""
    j = start
    while j < len(paragraphs):
        pj = paragraphs[j]
        if _is_heading(pj, 2) or _is_heading(pj, 1):
            break
        if not _is_heading(pj) and pj.text.strip() == "Purpose":
            j += 1
            desc_lines: list[str] = []
            while j < len(paragraphs):
                pk = paragraphs[j]
                if _is_heading(pk, 2) or _is_heading(pk, 1):
                    break
                text = pk.text.strip()
                if text == "Capabilities":
                    break
                if text.startswith("Consuming Domains:"):
                    break
                if text.startswith("Entities Owned:"):
                    break
                if text and not _is_heading(pk):
                    desc_lines.append(text)
                j += 1
            return "\n\n".join(desc_lines)
        j += 1
    return ""


# ---------------------------------------------------------------------------
# Prose enrichment
# ---------------------------------------------------------------------------


def _enrich_subdomain_descriptions(
    sub_domains_by_parent: dict[str, list[dict]],
    paragraphs: list,
    domains: list[dict],
    report: ParseReport,
) -> None:
    """Enrich sub-domain descriptions from prose in Section 3.

    Pass 1: search under the "Sub-Domains" Heading 3 (per spec).
    Pass 2: for any sub-domain still missing, search the entire domain section
    (handles v2.3 documents where sub-domain descriptions are scattered).
    """
    all_sd_codes: set[str] = set()
    for sds in sub_domains_by_parent.values():
        for sd in sds:
            all_sd_codes.add(sd["code"])
    if not all_sd_codes:
        return

    prose_lookup: dict[str, str] = {}

    for domain in domains:
        parent_code = domain["code"]
        if parent_code not in sub_domains_by_parent:
            continue
        start, end = _domain_section_range(paragraphs, parent_code)
        if start is None or end is None:
            continue

        parent_sd_codes = {sd["code"] for sd in sub_domains_by_parent[parent_code]}

        # Pass 1: under "Sub-Domains" Heading 3
        i = start
        while i < end:
            p = paragraphs[i]
            if _is_heading(p, 3) and p.text.strip() == "Sub-Domains":
                i += 1
                while i < end:
                    pp = paragraphs[i]
                    if _is_heading(pp):
                        break
                    text = pp.text.strip()
                    if text:
                        m = _CODE_NAME_RE.match(text)
                        if m:
                            code = m.group(1)
                            desc = _next_plain_text(paragraphs, i + 1, end)
                            prose_lookup[code] = desc
                    i += 1
                break
            i += 1

        # Pass 2: broader search for sub-domains still missing descriptions
        missing_codes = parent_sd_codes - set(prose_lookup.keys())
        if missing_codes:
            for i in range(start, end):
                p = paragraphs[i]
                if _is_heading(p):
                    continue
                text = p.text.strip()
                if text:
                    m = _CODE_NAME_RE.match(text)
                    if m and m.group(1) in missing_codes:
                        code = m.group(1)
                        desc = _next_plain_text(paragraphs, i + 1, end)
                        prose_lookup[code] = desc
                        missing_codes.discard(code)
                        if not missing_codes:
                            break

    # Merge descriptions into sub-domain dicts
    for sds in sub_domains_by_parent.values():
        for sd in sds:
            if sd["code"] in prose_lookup:
                sd["description"] = prose_lookup[sd["code"]]
            else:
                report.warn(
                    "missing_description",
                    f"Sub-domain {sd['code']}",
                    "No prose description found",
                )

    # Orphan prose entries (found in doc but not in Table 2)
    for code in prose_lookup:
        if code not in all_sd_codes:
            report.warn(
                "orphan_prose",
                f"Sub-Domains prose for {code}",
                "Not in Table 2 sub-domains",
            )


def _enrich_process_descriptions(
    processes: list[dict],
    paragraphs: list,
    domains: list[dict],
    report: ParseReport,
) -> None:
    """Enrich process descriptions from prose under Business Processes headings."""
    prose_lookup: dict[str, str] = {}
    proc_codes = {p["code"] for p in processes}

    for domain in domains:
        start, end = _domain_section_range(paragraphs, domain["code"])
        if start is None or end is None:
            continue

        # Find "Business Processes" Heading 3
        i = start
        while i < end:
            p = paragraphs[i]
            if _is_heading(p, 3) and p.text.strip() == "Business Processes":
                i += 1
                while i < end:
                    pp = paragraphs[i]
                    if _is_heading(pp):
                        break
                    text = pp.text.strip()
                    if text:
                        m = _CODE_NAME_RE.match(text)
                        if m:
                            code = m.group(1)
                            desc = _next_plain_text(paragraphs, i + 1, end)
                            prose_lookup[code] = desc
                    i += 1
                break
            i += 1

    # Merge
    for proc in processes:
        if proc["code"] in prose_lookup:
            proc["description"] = prose_lookup[proc["code"]]
        else:
            report.warn(
                "missing_description", f"Process {proc['code']}",
                "No prose description found",
            )

    # Orphan prose entries (in doc but not in Table 2)
    for code in prose_lookup:
        if code not in proc_codes:
            report.warn(
                "orphan_prose",
                f"Business Processes prose for {code}",
                "Not in Table 2 processes",
            )


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse(
    path: str | Path,
    work_item: dict,
    session_type: str = "initial",
) -> tuple[str, ParseReport]:
    """Parse CBM-Master-PRD.docx into a Path B envelope JSON string.

    :param path: Path to the .docx file.
    :param work_item: Work item dict with keys 'id' and 'item_type'.
                      item_type must be 'master_prd' or ValueError is raised.
    :param session_type: Session type for the envelope. Default 'initial'.
    :returns: Tuple of (envelope_json_string, ParseReport).
    :raises FileNotFoundError: If path does not exist.
    :raises MasterPrdParseError: If the document is structurally unparseable.
    :raises ValueError: If work_item['item_type'] != 'master_prd'.
    """
    if work_item.get("item_type") != "master_prd":
        raise ValueError(
            f"work_item['item_type'] must be 'master_prd', "
            f"got {work_item.get('item_type')!r}"
        )

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    report = ParseReport()
    doc = Document(str(path))
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)

    # Parse each section
    org_overview = _parse_org_overview(paragraphs, report)
    personas = _parse_personas(paragraphs, report)
    domains = _parse_domains(paragraphs, report)
    table2 = _find_table2(tables)
    sub_domains_by_parent, processes = _parse_table2(table2, report)
    services = _parse_services(paragraphs, report)

    # Enrich descriptions from prose
    _enrich_subdomain_descriptions(
        sub_domains_by_parent, paragraphs, domains, report,
    )
    _enrich_process_descriptions(processes, paragraphs, domains, report)

    # Attach sub-domains to their parent domain entries
    for domain in domains:
        domain["sub_domains"] = sub_domains_by_parent.get(domain["code"], [])

    # Combine top-level domains and services into one list
    all_domains = domains + services

    # Build payload
    payload: dict[str, Any] = {
        "organization_overview": org_overview,
        "personas": personas,
        "domains": all_domains,
        "processes": processes,
    }

    # Record parsed counts
    report.parsed_counts = {
        "persona": len(personas),
        "domain": len(domains),
        "sub_domain": sum(len(d["sub_domains"]) for d in domains),
        "service": len(services),
        "process": len(processes),
    }

    # Build envelope
    envelope: dict[str, Any] = {
        "output_version": "1.0",
        "work_item_type": "master_prd",
        "work_item_id": work_item["id"],
        "session_type": session_type,
        "payload": payload,
        "decisions": [],
        "open_issues": [],
    }

    envelope_json = json.dumps(envelope, ensure_ascii=False, indent=2)
    return envelope_json, report
