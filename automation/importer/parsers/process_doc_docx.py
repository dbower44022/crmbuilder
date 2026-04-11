"""Parse a CBM Process Document .docx into a Path B envelope JSON string.

Adapter that turns a Word process document into a Path B envelope suitable
for processing by the existing ImportProcessor seven-stage pipeline.

Replaces the Path A parser at automation.cbm_import.parsers.process_document.
Fixes Bug 4 (domain_code extraction) and Bug 5 (sub-domain assignment).
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

from automation.importer.parsers import ParseReport, ProcessDocParseError

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Header table: "Name (CODE)" where CODE can be multi-segment like CR-PARTNER
_NAME_CODE_RE = re.compile(r"^(.+?)\s*\(([A-Z][A-Z0-9-]*)\)\s*$")

# Section heading patterns (Heading 1)
_SECTION_RES = {
    1: re.compile(r"^1\.\s+Process Purpose$"),
    2: re.compile(r"^2\.\s+Process Triggers?$"),
    3: re.compile(r"^3\.\s+Personas Involved$"),
    4: re.compile(r"^4\.\s+Process Workflow$"),
    5: re.compile(r"^5\.\s+Process Completion$"),
    6: re.compile(r"^6\.\s+System Requirements$"),
    7: re.compile(r"^7\.\s+Process Data$"),
    8: re.compile(r"^8\.\s+Data Collected$"),
    9: re.compile(r"^9\.\s+Open Issues$"),
}

# Persona paragraph pattern (Format A): "Name (MST-PER-NNN)"
_PERSONA_PARA_RE = re.compile(r"^(.+?)\s*\((MST-PER-\d+)\)\s*$")

# Leading number token in Heading 2 (e.g. "4.1 " or "7 ")
_LEADING_NUM_RE = re.compile(r"^\d+(?:\.\d+)?\s+")

# Entity heading in Sections 7/8: "Entity: Name (qualifier)"
_ENTITY_HEADING_RE = re.compile(r"^Entity:\s*(.+)$")

# Field table header (six columns)
_FIELD_TABLE_HEADERS = ["field name", "type", "required", "values", "default", "id"]

# Requirement identifier pattern
_REQ_ID_RE = re.compile(r"-REQ-\d+$")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _style_name(p: Any) -> str:
    """Return the paragraph's style name, or '' if unavailable."""
    if p.style is None:
        return ""
    name = p.style.name
    return name if name else ""


def _is_heading(p: Any, level: int | None = None) -> bool:
    """Check whether paragraph *p* is a Word heading."""
    name = _style_name(p)
    if level is not None:
        return name == f"Heading {level}"
    return name.startswith("Heading ")


def _find_heading1(paragraphs: list, pattern: re.Pattern) -> int | None:
    """Return the index of the first Heading 1 matching *pattern*."""
    for i, p in enumerate(paragraphs):
        if _is_heading(p, 1) and pattern.match(p.text.strip()):
            return i
    return None


def _section_range(paragraphs: list, start: int) -> int:
    """Return the end index (exclusive) for a Heading 1 section."""
    for j in range(start + 1, len(paragraphs)):
        if _is_heading(paragraphs[j], 1):
            return j
    return len(paragraphs)


def _collect_prose(paragraphs: list, start: int, end: int) -> str:
    """Collect all non-heading paragraph text between start and end."""
    parts: list[str] = []
    for i in range(start, end):
        p = paragraphs[i]
        if _is_heading(p):
            continue
        text = p.text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _tables_in_range(doc: Document, paragraphs: list, start: int, end: int) -> list:
    """Return all tables whose position falls between paragraph indices start and end.

    Uses the document's XML element ordering to determine table position
    relative to paragraphs.
    """
    body = doc.element.body
    elements = list(body)

    # Build paragraph index map: XML element -> paragraph index
    para_elem_indices: dict[int, int] = {}
    pi = 0
    for ei, elem in enumerate(elements):
        if elem.tag.endswith("}p"):
            if pi < len(paragraphs):
                para_elem_indices[ei] = pi
                pi += 1

    # Find element indices for start and end paragraphs
    start_elem_idx = None
    end_elem_idx = None
    for ei, pi in sorted(para_elem_indices.items()):
        if pi == start and start_elem_idx is None:
            start_elem_idx = ei
        if pi >= end and end_elem_idx is None:
            end_elem_idx = ei
            break
    if start_elem_idx is None:
        return []
    if end_elem_idx is None:
        end_elem_idx = len(elements)

    # Collect tables between those element indices
    result = []
    table_idx = 0
    for ei, elem in enumerate(elements):
        if elem.tag.endswith("}tbl"):
            if table_idx < len(doc.tables):
                if start_elem_idx <= ei < end_elem_idx:
                    result.append(doc.tables[table_idx])
                table_idx += 1
    return result


def _table_header(table: Any) -> list[str]:
    """Return lowercased, stripped header row cells."""
    if not table.rows:
        return []
    return [cell.text.strip().lower() for cell in table.rows[0].cells]


def _table_rows(table: Any) -> list[list[str]]:
    """Return data rows (excluding header) as lists of stripped cell strings."""
    rows = []
    for i in range(1, len(table.rows)):
        rows.append([cell.text.strip() for cell in table.rows[i].cells])
    return rows


def _parse_required(text: str, report: ParseReport, location: str) -> bool:
    """Parse a required field value to bool."""
    cleaned = text.strip().lower()
    if cleaned.startswith("yes"):
        return True
    if cleaned.startswith("no"):
        return False
    report.warn(
        "unclear_required", location,
        f"Unclear required value: {text!r}, treating as True",
    )
    return True


def _derive_role(description: str) -> str:
    """Derive persona role from description text."""
    lower = description.lower()
    if "initiat" in lower:
        return "initiator"
    if "approv" in lower:
        return "approver"
    if "receiv" in lower or "notif" in lower:
        return "recipient"
    return "performer"


# ---------------------------------------------------------------------------
# Section parsers
# ---------------------------------------------------------------------------


def _parse_header_table(
    doc: Document, report: ParseReport,
) -> dict[str, str]:
    """Parse the first table (header table) for metadata fields."""
    if not doc.tables:
        raise ProcessDocParseError("No header table found in document")

    table = doc.tables[0]
    meta: dict[str, str] = {}
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            meta[cells[0]] = cells[1]

    # Process Code is required
    process_code = meta.get("Process Code", "")
    if not process_code:
        raise ProcessDocParseError(
            "Header table is missing a 'Process Code' row"
        )

    result: dict[str, str] = {"process_code": process_code}

    # Domain — required with Name (CODE) pattern
    domain_raw = meta.get("Domain", "")
    if domain_raw:
        m = _NAME_CODE_RE.match(domain_raw)
        if not m:
            raise ProcessDocParseError(
                f"Domain row does not match 'Name (CODE)' pattern: {domain_raw!r}"
            )
        result["domain_name"] = m.group(1).strip()
        result["domain_code"] = m.group(2)
    else:
        report.warn("missing_field", "header_table", "No Domain row found")

    # Sub-Domain — optional
    sub_domain_raw = meta.get("Sub-Domain", "")
    if sub_domain_raw:
        m = _NAME_CODE_RE.match(sub_domain_raw)
        if not m:
            raise ProcessDocParseError(
                f"Sub-Domain row does not match 'Name (CODE)' pattern: "
                f"{sub_domain_raw!r}"
            )
        result["sub_domain_name"] = m.group(1).strip()
        result["sub_domain_code"] = m.group(2)

    # Validate process code prefix
    domain_code = result.get("domain_code", "")
    sub_domain_code = result.get("sub_domain_code", "")
    if sub_domain_code:
        if not process_code.startswith(f"{sub_domain_code}-"):
            raise ProcessDocParseError(
                f"Process code {process_code!r} does not start with "
                f"sub-domain code prefix '{sub_domain_code}-'"
            )
    elif domain_code:
        if not process_code.startswith(f"{domain_code}-"):
            raise ProcessDocParseError(
                f"Process code {process_code!r} does not start with "
                f"domain code prefix '{domain_code}-'"
            )

    # Process Name
    result["process_name"] = meta.get("Process Name", "")

    # Optional metadata fields
    for field in ("Version", "Status", "Last Updated"):
        value = meta.get(field, "")
        if value:
            result[field.lower().replace(" ", "_")] = value
        else:
            report.warn(
                "missing_field", "header_table",
                f"Missing {field!r} row in header table",
            )

    return result


def _parse_prose_section(
    paragraphs: list, section_num: int, section_name: str,
) -> str:
    """Parse a prose section (Sections 1, 2, 5) — hard-fail if missing or empty."""
    pattern = _SECTION_RES[section_num]
    h1 = _find_heading1(paragraphs, pattern)
    if h1 is None:
        raise ProcessDocParseError(
            f"Document has no Heading 1 matching Section {section_num} "
            f"({section_name})"
        )
    end = _section_range(paragraphs, h1)
    text = _collect_prose(paragraphs, h1 + 1, end)
    if not text:
        raise ProcessDocParseError(
            f"Section {section_num} ({section_name}) exists but is empty"
        )
    return text


def _parse_personas(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 3 — Personas Involved."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[3])
    if h1 is None:
        raise ProcessDocParseError(
            "Document has no Heading 1 matching Section 3 (Personas Involved)"
        )
    end = _section_range(paragraphs, h1)

    # Check for persona table (Format B)
    tables = _tables_in_range(doc, paragraphs, h1, end)
    for table in tables:
        header = _table_header(table)
        has_id = any("id" in h for h in header)
        has_persona = any("persona" in h for h in header)
        has_role = any("role" in h for h in header)
        if has_id and has_persona and has_role:
            personas: list[dict] = []
            for row_cells in _table_rows(table):
                if len(row_cells) < 3 or not row_cells[0].strip():
                    continue
                identifier = row_cells[0].strip()
                name = row_cells[1].strip()
                description = row_cells[2].strip()
                role = _derive_role(description)
                personas.append({
                    "identifier": identifier,
                    "name": name,
                    "role": role,
                    "description": description,
                })
            if personas:
                return personas

    # Format A — paragraph scan
    personas = []
    i = h1 + 1
    while i < end:
        p = paragraphs[i]
        if _is_heading(p):
            i += 1
            continue
        text = p.text.strip()
        if not text:
            i += 1
            continue
        m = _PERSONA_PARA_RE.match(text)
        if m:
            name = m.group(1).strip()
            identifier = m.group(2)
            # Next non-empty, non-heading paragraph is description
            description = ""
            for j in range(i + 1, end):
                pj = paragraphs[j]
                if _is_heading(pj):
                    break
                desc_text = pj.text.strip()
                if desc_text:
                    # Check it's not another persona line
                    if _PERSONA_PARA_RE.match(desc_text):
                        break
                    description = desc_text
                    break
            role = _derive_role(description)
            personas.append({
                "identifier": identifier,
                "name": name,
                "role": role,
                "description": description,
            })
        i += 1

    if not personas:
        raise ProcessDocParseError(
            "Section 3 (Personas Involved) exists but zero personas were parsed"
        )
    return personas


def _parse_workflow(
    paragraphs: list, report: ParseReport,
) -> list[dict]:
    """Section 4 — Process Workflow."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[4])
    if h1 is None:
        raise ProcessDocParseError(
            "Document has no Heading 1 matching Section 4 (Process Workflow)"
        )
    end = _section_range(paragraphs, h1)

    # Determine format: any Heading 2 in range → Format B (activity areas)
    has_h2 = any(
        _is_heading(paragraphs[i], 2) for i in range(h1 + 1, end)
    )

    steps: list[dict] = []

    if has_h2:
        # Format B — activity areas
        # Warn about framing prose between H1 and first H2
        first_h2_idx = None
        for i in range(h1 + 1, end):
            if _is_heading(paragraphs[i], 2):
                first_h2_idx = i
                break
        if first_h2_idx is not None:
            framing = _collect_prose(paragraphs, h1 + 1, first_h2_idx)
            if framing:
                report.warn(
                    "framing_prose", "Section 4",
                    "Prose between Section 4 heading and first activity area — dropped",
                )

        # Each Heading 2 is one step
        sort_order = 0
        i = h1 + 1
        while i < end:
            p = paragraphs[i]
            if _is_heading(p, 2):
                sort_order += 1
                raw_name = p.text.strip()
                name = _LEADING_NUM_RE.sub("", raw_name)[:200]

                # Collect description until next H2 or H1
                step_end = end
                for j in range(i + 1, end):
                    if _is_heading(paragraphs[j], 2) or _is_heading(paragraphs[j], 1):
                        step_end = j
                        break

                # Include all non-heading paragraphs (Normal + List Paragraph)
                desc_parts: list[str] = []
                for j in range(i + 1, step_end):
                    pj = paragraphs[j]
                    if _is_heading(pj):
                        continue
                    text = pj.text.strip()
                    if text:
                        desc_parts.append(text)

                steps.append({
                    "name": name,
                    "description": "\n\n".join(desc_parts),
                    "step_type": "action",
                    "sort_order": sort_order,
                })
                i = step_end
                continue
            i += 1
    else:
        # Format A — flat list
        sort_order = 0
        last_step_idx: int | None = None
        for i in range(h1 + 1, end):
            p = paragraphs[i]
            if _is_heading(p):
                continue
            text = p.text.strip()
            if not text:
                continue

            style = _style_name(p)
            is_list_item = "list" in style.lower()

            if is_list_item:
                sort_order += 1
                steps.append({
                    "name": text[:200],
                    "description": text,
                    "step_type": "action",
                    "sort_order": sort_order,
                })
                last_step_idx = len(steps) - 1
            elif last_step_idx is not None:
                # Normal paragraph following a list item — append to description
                steps[last_step_idx]["description"] += " " + text
            else:
                # Normal paragraph with no preceding step — drop with warning
                report.warn(
                    "orphan_prose", "Section 4",
                    f"Normal paragraph before any step — dropped: {text[:80]}",
                )

    if not steps:
        raise ProcessDocParseError(
            "Section 4 (Process Workflow) exists but zero steps were parsed"
        )
    return steps


def _parse_requirements(
    paragraphs: list, doc: Document, process_code: str, report: ParseReport,
) -> list[dict]:
    """Section 6 — System Requirements."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[6])
    if h1 is None:
        raise ProcessDocParseError(
            "Document has no Heading 1 matching Section 6 (System Requirements)"
        )
    end = _section_range(paragraphs, h1)

    # Find matching table: header row ["ID", "Requirement"] (case-insensitive)
    tables = _tables_in_range(doc, paragraphs, h1, end)
    req_table = None
    for table in tables:
        header = _table_header(table)
        if len(header) >= 2 and header[0] == "id" and "requirement" in header[1]:
            req_table = table
            break

    if req_table is None:
        raise ProcessDocParseError(
            "Section 6 (System Requirements) has no matching "
            "['ID', 'Requirement'] table"
        )

    requirements: list[dict] = []
    expected_prefix_re = re.compile(rf"^{re.escape(process_code)}-REQ-\d+$")
    for row_cells in _table_rows(req_table):
        if len(row_cells) < 2:
            continue
        identifier = row_cells[0].strip()
        description = row_cells[1].strip()
        if not identifier or not description:
            continue
        if "-REQ-" not in identifier:
            continue
        if not expected_prefix_re.match(identifier):
            report.warn(
                "identifier_mismatch", f"Requirement {identifier}",
                f"Identifier does not match expected pattern "
                f"{process_code}-REQ-NNN",
            )
        requirements.append({
            "identifier": identifier,
            "description": description,
            "priority": "must",
        })

    if not requirements:
        raise ProcessDocParseError(
            "Section 6 (System Requirements) table has zero valid requirements"
        )
    return requirements


def _parse_field_table(
    table: Any, report: ParseReport, location: str,
) -> list[dict]:
    """Parse a six-column, two-row-per-field table into field dicts."""
    header = _table_header(table)
    if len(header) < 6:
        return []
    # Check if this is a field table
    if header != _FIELD_TABLE_HEADERS:
        return []

    data_rows = _table_rows(table)
    if not data_rows:
        return []

    fields: list[dict] = []
    i = 0
    while i < len(data_rows):
        meta_row = data_rows[i]
        if len(meta_row) < 6 or not meta_row[0].strip():
            i += 1
            continue

        field_name = meta_row[0].strip()
        field_type = meta_row[1].strip()
        required_text = meta_row[2].strip()
        values = meta_row[3].strip()
        default_value = meta_row[4].strip()
        identifier = meta_row[5].strip()

        is_required = _parse_required(
            required_text, report, f"{location}.{field_name}",
        )

        # Description row
        description = ""
        if i + 1 < len(data_rows):
            desc_row = data_rows[i + 1]
            description = desc_row[0].strip() if desc_row else ""
            # Check if description cells are consistent (soft-warn if not)
            if len(desc_row) >= 2:
                unique_cells = {c.strip() for c in desc_row if c.strip()}
                if len(unique_cells) > 1:
                    report.warn(
                        "inconsistent_description", f"{location}.{field_name}",
                        "Description row cells don't all match",
                    )
            i += 2
        else:
            report.warn(
                "odd_row_count", location,
                f"Field table has odd row count — missing description for {field_name}",
            )
            i += 1

        fields.append({
            "field_name": field_name,
            "label": field_name,
            "field_type": field_type,
            "is_required": is_required,
            "values": values,
            "default_value": default_value,
            "identifier": identifier,
            "description": description,
        })

    return fields


def _parse_entity_subsections(
    paragraphs: list, doc: Document, h1: int, end: int,
    report: ParseReport, section_num: int,
) -> list[tuple[str, str, list[dict]]]:
    """Parse Entity: subsections within Sections 7 or 8.

    Returns list of (entity_name, description, fields).
    """
    entities: list[tuple[str, str, list[dict]]] = []

    # Find each Heading 2 starting with "Entity:"
    i = h1 + 1
    while i < end:
        p = paragraphs[i]
        if _is_heading(p, 2):
            m = _ENTITY_HEADING_RE.match(p.text.strip())
            if m:
                raw = m.group(1).strip()
                # Extract entity name: strip parenthetical qualifier
                paren_match = re.match(r"^(\w[\w\s]*?)(?:\s*\(.+\))?\s*$", raw)
                entity_name = paren_match.group(1).strip() if paren_match else raw

                # Qualifier text for context
                qualifier = ""
                qual_match = re.search(r"\((.+)\)", raw)
                if qual_match:
                    qualifier = qual_match.group(1)

                # Find subsection end
                sub_end = end
                for j in range(i + 1, end):
                    if _is_heading(paragraphs[j], 2) or _is_heading(paragraphs[j], 1):
                        sub_end = j
                        break

                # Collect description prose
                desc_parts: list[str] = []
                if qualifier:
                    desc_parts.append(f"({qualifier})")
                for j in range(i + 1, sub_end):
                    pj = paragraphs[j]
                    if _is_heading(pj):
                        continue
                    text = pj.text.strip()
                    if text:
                        desc_parts.append(text)
                description = "\n\n".join(desc_parts)

                # Parse field tables in this subsection
                tables = _tables_in_range(doc, paragraphs, i, sub_end)
                fields: list[dict] = []
                loc = f"Section {section_num}.{entity_name}"
                for table in tables:
                    fields.extend(_parse_field_table(table, report, loc))

                entities.append((entity_name, description, fields))
                i = sub_end
                continue
        i += 1

    return entities


def _parse_process_data(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 7 — Process Data."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[7])
    if h1 is None:
        raise ProcessDocParseError(
            "Document has no Heading 1 matching Section 7 (Process Data)"
        )
    end = _section_range(paragraphs, h1)

    entity_subs = _parse_entity_subsections(
        paragraphs, doc, h1, end, report, section_num=7,
    )

    if not entity_subs:
        report.warn(
            "empty_section", "Section 7",
            "Section 7 has no entity subsections",
        )
        return []

    process_data: list[dict] = []
    for entity_name, description, fields in entity_subs:
        field_refs = [
            {
                "name": f["field_name"],
                "usage": "displayed",
                "description": f["description"],
            }
            for f in fields
        ]
        process_data.append({
            "entity_name": entity_name,
            "role": "referenced",
            "description": description,
            "field_references": field_refs,
        })

    return process_data


def _parse_data_collected(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 8 — Data Collected."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[8])
    if h1 is None:
        raise ProcessDocParseError(
            "Document has no Heading 1 matching Section 8 (Data Collected)"
        )
    end = _section_range(paragraphs, h1)

    entity_subs = _parse_entity_subsections(
        paragraphs, doc, h1, end, report, section_num=8,
    )

    if not entity_subs:
        report.warn(
            "empty_section", "Section 8",
            "Section 8 has no entity subsections",
        )

    data_collected: list[dict] = []
    for entity_name, _description, fields in entity_subs:
        data_collected.append({
            "entity_name": entity_name,
            "new_fields": fields,
        })

    return data_collected


def _parse_open_issues(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 9 — Open Issues."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[9])
    if h1 is None:
        raise ProcessDocParseError(
            "Document has no Heading 1 matching Section 9 (Open Issues)"
        )
    end = _section_range(paragraphs, h1)

    # Find two-column tables with header ["ID", "Issue"] or ["ID", "Description"]
    tables = _tables_in_range(doc, paragraphs, h1, end)
    issue_tables: list[Any] = []
    for table in tables:
        header = _table_header(table)
        if (len(header) >= 2 and header[0] == "id"
                and header[1] in ("issue", "description", "requirement")):
            issue_tables.append(table)

    if len(issue_tables) > 1:
        report.warn(
            "multiple_tables", "Section 9",
            "Section 9 has multiple tables; parsing only the first "
            "(process-owned issues). Subsequent tables likely contain "
            "inherited/upstream issues.",
        )

    issues: list[dict] = []
    if issue_tables:
        for row_cells in _table_rows(issue_tables[0]):
            if len(row_cells) < 2:
                continue
            identifier = row_cells[0].strip()
            description = row_cells[1].strip()
            if not identifier:
                continue
            if description.upper().startswith(("CLOSED", "RESOLVED")):
                report.warn(
                    "closed_issue", f"Issue {identifier}",
                    f"Issue description starts with CLOSED/RESOLVED: {description[:60]}",
                )
            issues.append({
                "identifier": identifier,
                "description": description,
                "status": "open",
            })

    return issues


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse(
    path: str | Path,
    work_item: dict,
    session_type: str = "initial",
) -> tuple[str, ParseReport]:
    """Parse a process document .docx into a Path B envelope JSON string.

    :param path: Path to the .docx file.
    :param work_item: Must have item_type == 'process_definition' and
                      a 'process_id' key (the target Process row id).
    :param session_type: Session type for the envelope. Default 'initial'.
    :returns: Tuple of (envelope_json_string, ParseReport).
    :raises ValueError: If work_item['item_type'] != 'process_definition'.
    :raises FileNotFoundError: If path does not exist.
    :raises ProcessDocParseError: If the document is structurally unparseable.
    """
    if work_item.get("item_type") != "process_definition":
        raise ValueError(
            f"work_item['item_type'] must be 'process_definition', "
            f"got {work_item.get('item_type')!r}"
        )

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    report = ParseReport()
    doc = Document(str(path))
    paragraphs = list(doc.paragraphs)

    # Header table
    header_meta = _parse_header_table(doc, report)

    # Prose sections
    process_purpose = _parse_prose_section(paragraphs, 1, "Process Purpose")
    triggers = _parse_prose_section(paragraphs, 2, "Process Triggers")
    completion = _parse_prose_section(paragraphs, 5, "Process Completion")

    # Structured sections
    personas = _parse_personas(paragraphs, doc, report)
    workflow = _parse_workflow(paragraphs, report)
    requirements = _parse_requirements(
        paragraphs, doc, header_meta["process_code"], report,
    )
    process_data = _parse_process_data(paragraphs, doc, report)
    data_collected = _parse_data_collected(paragraphs, doc, report)
    open_issues = _parse_open_issues(paragraphs, doc, report)

    # Build source_metadata
    source_metadata: dict[str, Any] = {
        "process_code": header_meta["process_code"],
        "process_name": header_meta.get("process_name", ""),
        "domain_code": header_meta.get("domain_code", ""),
        "version": header_meta.get("version", ""),
        "status": header_meta.get("status", ""),
        "last_updated": header_meta.get("last_updated", ""),
    }
    if "sub_domain_code" in header_meta:
        source_metadata["sub_domain_code"] = header_meta["sub_domain_code"]

    # Record parsed counts
    report.parsed_counts = {
        "persona": len(personas),
        "workflow_step": len(workflow),
        "requirement": len(requirements),
        "process_data_entity": len(process_data),
        "data_collected_entity": len(data_collected),
        "open_issue": len(open_issues),
    }

    # Build payload — triggers and completion wrapped as dicts to satisfy
    # Layer 3 validation (parser.py PAYLOAD_KEYS expects dict for these)
    payload: dict[str, Any] = {
        "source_metadata": source_metadata,
        "process_purpose": process_purpose,
        "triggers": {"description": triggers},
        "personas": personas,
        "workflow": workflow,
        "completion": {"description": completion},
        "system_requirements": requirements,
        "process_data": process_data,
        "data_collected": data_collected,
    }

    # Build envelope
    envelope: dict[str, Any] = {
        "output_version": "1.0",
        "work_item_type": "process_definition",
        "work_item_id": work_item["id"],
        "session_type": session_type,
        "payload": payload,
        "decisions": [],
        "open_issues": open_issues,
    }

    envelope_json = json.dumps(envelope, ensure_ascii=False, indent=2)
    return envelope_json, report
