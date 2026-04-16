"""Parse an Entity PRD .docx into a Path B envelope JSON string.

Adapter that turns a Word entity PRD document into a Path B envelope
suitable for processing by the existing ImportProcessor pipeline.

Replaces the Path A parser at automation.cbm_import.parsers.entity_prd.
Fixes Bug 6 (primary_domain_id always NULL) by extracting domain
references from the Entity Overview table.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

from automation.importer.parsers import EntityPrdParseError, ParseReport

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Header table Entity value: "Contact (Native — Person Type)"
_ENTITY_HEADER_RE = re.compile(
    r"^(.+?)\s*\(\s*(Native|Custom)\s*[—–\-]\s*(.+?)\s*\)\s*$"
)

# "Name (CODE)" pattern for domains
_NAME_CODE_RE = re.compile(r"^(.+?)\s*\(([A-Z][A-Z0-9-]*)\)\s*$")

# Section heading patterns (Heading 1)
_SECTION_RES: dict[int, re.Pattern] = {
    1: re.compile(r"^1\.\s+Entity Overview$"),
    2: re.compile(r"^2\.\s+Native Fields$"),
    3: re.compile(r"^3\.\s+Custom Fields$"),
    4: re.compile(r"^4\.\s+Relationships$"),
    5: re.compile(r"^5\.\s+Dynamic Logic Rules$"),
    6: re.compile(r"^6\.\s+Layout Guidance$"),
    7: re.compile(r"^7\.\s+Implementation Notes$"),
    8: re.compile(r"^8\.\s+Open Issues$"),
    9: re.compile(r"^9\.\s+Decisions Made$"),
}

# Leading number prefix: "3.1 " or "5 "
_LEADING_NUM_RE = re.compile(r"^\d+(?:\.\d+)?\s+")

# Expected six-column header for custom field tables
_FIELD_TABLE_HEADERS = ["field name", "type", "required", "values", "default", "id"]

# Match subsection titles like "Incomplete Domain Fields"
_INCOMPLETE_DOMAIN_RE = re.compile(r"^.*incomplete\s+domain.*$", re.IGNORECASE)

# Valid EspoCRM link types
_VALID_LINK_TYPES = frozenset({"oneToMany", "manyToOne", "manyToMany", "oneToOne"})

# Decision identifier pattern
_DEC_ID_RE = re.compile(r"^[A-Z]+-DEC-\d+$")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _style_name(p: Any) -> str:
    """Return the paragraph's style name, or '' if unavailable."""
    if p.style is None:
        return ""
    name = p.style.name
    return name if name else ""


def _is_heading(p: Any, level: int | None = None) -> bool:
    """Check whether paragraph *p* is a Word heading."""
    name = _style_name(p)
    if level is not None:
        return name == f"Heading {level}"
    return name.startswith("Heading ")


def _find_heading1(paragraphs: list, pattern: re.Pattern) -> int | None:
    """Return the index of the first Heading 1 matching *pattern*."""
    for i, p in enumerate(paragraphs):
        if _is_heading(p, 1) and pattern.match(p.text.strip()):
            return i
    return None


def _section_range(paragraphs: list, start: int) -> int:
    """Return the end index (exclusive) for a Heading 1 section."""
    for j in range(start + 1, len(paragraphs)):
        if _is_heading(paragraphs[j], 1):
            return j
    return len(paragraphs)


def _collect_prose(paragraphs: list, start: int, end: int) -> str:
    """Collect all non-heading paragraph text between start and end."""
    parts: list[str] = []
    for i in range(start, end):
        p = paragraphs[i]
        if _is_heading(p):
            continue
        text = p.text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _tables_in_range(
    doc: Document, paragraphs: list, start: int, end: int,
) -> list:
    """Return all tables whose position falls between paragraph indices.

    Uses the document's XML element ordering to determine table position
    relative to paragraphs.
    """
    body = doc.element.body
    elements = list(body)

    # Build paragraph index map: XML element -> paragraph index
    para_elem_indices: dict[int, int] = {}
    pi = 0
    for ei, elem in enumerate(elements):
        if elem.tag.endswith("}p"):
            if pi < len(paragraphs):
                para_elem_indices[ei] = pi
                pi += 1

    # Find element indices for start and end paragraphs
    start_elem_idx = None
    end_elem_idx = None
    for ei, pi_val in sorted(para_elem_indices.items()):
        if pi_val == start and start_elem_idx is None:
            start_elem_idx = ei
        if pi_val >= end and end_elem_idx is None:
            end_elem_idx = ei
            break
    if start_elem_idx is None:
        return []
    if end_elem_idx is None:
        end_elem_idx = len(elements)

    # Collect tables between those element indices
    result = []
    table_idx = 0
    for ei, elem in enumerate(elements):
        if elem.tag.endswith("}tbl"):
            if table_idx < len(doc.tables):
                if start_elem_idx <= ei < end_elem_idx:
                    result.append(doc.tables[table_idx])
                table_idx += 1
    return result


def _table_header(table: Any) -> list[str]:
    """Return lowercased, stripped header row cells."""
    if not table.rows:
        return []
    return [cell.text.strip().lower() for cell in table.rows[0].cells]


def _table_rows(table: Any) -> list[list[str]]:
    """Return data rows (excluding header) as lists of stripped cell strings."""
    rows = []
    for i in range(1, len(table.rows)):
        rows.append([cell.text.strip() for cell in table.rows[i].cells])
    return rows


def _to_label(field_name: str) -> str:
    """Convert camelCase field name to a human-readable label."""
    label = re.sub(r"([a-z])([A-Z])", r"\1 \2", field_name)
    return label[0].upper() + label[1:] if label else label


def _clean_default(value: str) -> str | None:
    """Normalize placeholder default values to None."""
    if value.strip() in ("—", "–", "-", "N/A", "n/a", ""):
        return None
    return value.strip()


def _parse_required(text: str, report: ParseReport, location: str) -> bool:
    """Parse a required field value to bool."""
    cleaned = text.strip().lower()
    if cleaned.startswith("yes"):
        return True
    if cleaned.startswith("no"):
        return False
    report.warn(
        "unclear_required", location,
        f"Unclear required value: {text!r}, treating as True",
    )
    return True


# ---------------------------------------------------------------------------
# Section parsers
# ---------------------------------------------------------------------------


def _parse_header_table(
    doc: Document, report: ParseReport,
) -> dict[str, Any]:
    """Parse the header table (Table 0) for document metadata."""
    if not doc.tables:
        raise EntityPrdParseError("No header table found in document")

    table = doc.tables[0]
    meta: dict[str, str] = {}
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            meta[cells[0]] = cells[1]

    result: dict[str, Any] = {}

    # Entity row — required
    entity_raw = meta.get("Entity", meta.get("Entity Name", ""))
    if not entity_raw:
        raise EntityPrdParseError(
            "Header table is missing an 'Entity' row"
        )

    m = _ENTITY_HEADER_RE.match(entity_raw)
    if m:
        result["name"] = m.group(1).strip()
        result["is_native"] = m.group(2) == "Native"
        entity_type = m.group(3).strip()
        # Strip trailing " Type" if present
        if entity_type.endswith(" Type"):
            entity_type = entity_type[:-5].strip()
        result["entity_type"] = entity_type
    else:
        report.warn(
            "entity_value_no_parens", "Header Table.Entity",
            f"Entity value lacks parenthetical: {entity_raw!r} — "
            "is_native and entity_type will be null",
        )
        result["name"] = entity_raw.strip()
        result["is_native"] = None
        result["entity_type"] = None

    # Optional metadata fields
    result["version"] = meta.get("Version", "")
    if not result["version"]:
        report.warn("missing_field", "Header Table", "Missing Version row")

    result["status"] = meta.get("Status", "")
    if not result["status"]:
        report.warn("missing_field", "Header Table", "Missing Status row")

    result["last_updated"] = meta.get("Last Updated", "")
    if not result["last_updated"]:
        report.warn(
            "missing_field", "Header Table", "Missing Last Updated row",
        )

    result["source_documents"] = meta.get("Source Documents", "")
    if not result["source_documents"]:
        report.warn(
            "missing_field", "Header Table",
            "Missing Source Documents row",
        )

    return result


def _parse_entity_overview_table(
    doc: Document,
    report: ParseReport,
    header_meta: dict[str, Any],
) -> dict[str, Any]:
    """Parse the Entity Overview table for entity metadata."""
    # Find the Entity Overview table — first 2-col KV table after header
    overview_table = None
    for table in doc.tables[1:]:  # skip header table
        if not table.rows:
            continue
        first_cell = table.rows[0].cells[0].text.strip().lower()
        if first_cell in ("crm entity name", "entity name"):
            overview_table = table
            break

    if overview_table is None:
        raise EntityPrdParseError(
            "No Entity Overview table found (expected first two-column "
            "key/value table after header with 'CRM Entity Name' or "
            "'Entity Name' in first cell)"
        )

    # Build row lookup (case-insensitive keys)
    ov: dict[str, str] = {}
    for row in overview_table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            ov[cells[0].lower()] = cells[1]

    # CRM Entity Name — required, cross-check with header
    crm_entity_name = ov.get("crm entity name", ov.get("entity name", ""))
    if not crm_entity_name:
        raise EntityPrdParseError(
            "Entity Overview table is missing 'CRM Entity Name' value"
        )

    # Lenient cross-check against header name
    header_name = header_meta.get("name", "")
    if header_name and crm_entity_name:
        if header_name.lower().strip() != crm_entity_name.lower().strip():
            report.warn(
                "name_mismatch", "Entity Overview",
                f"Header name {header_name!r} differs from overview "
                f"name {crm_entity_name!r}",
            )

    result: dict[str, Any] = {"name": crm_entity_name}

    # Native / Custom cross-check
    native_custom = ov.get("native / custom", "")
    if native_custom:
        ov_is_native = "native" in native_custom.lower()
        header_is_native = header_meta.get("is_native")
        if header_is_native is not None and ov_is_native != header_is_native:
            report.warn(
                "native_mismatch", "Entity Overview",
                f"Header says {'Native' if header_is_native else 'Custom'} "
                f"but overview says {native_custom!r}",
            )

    # Entity Type cross-check
    entity_type = ov.get("entity type", "")
    if entity_type:
        header_type = header_meta.get("entity_type")
        if header_type and header_type.lower() != entity_type.lower().strip():
            report.warn(
                "entity_type_mismatch", "Entity Overview",
                f"Header entity type {header_type!r} differs from "
                f"overview {entity_type!r}",
            )
        result["entity_type"] = entity_type.strip()

    # Labels
    result["singular_label"] = ov.get("display label (singular)", "")
    result["plural_label"] = ov.get("display label (plural)", "")

    # Activity Stream
    activity = ov.get("activity stream", "")
    if activity:
        result["activity_stream"] = activity.lower().strip() == "yes"
    else:
        report.warn(
            "missing_field", "Entity Overview",
            "Missing Activity Stream row — defaulting to False",
        )
        result["activity_stream"] = False

    # --- Bug 6: Domain extraction ---

    # Primary Domain — optional
    primary_domain_raw = ov.get("primary domain", "")
    primary_domain_code: str | None = None
    if primary_domain_raw:
        m = _NAME_CODE_RE.match(primary_domain_raw)
        if m:
            primary_domain_code = m.group(2)
        else:
            # Try bare code
            primary_domain_code = primary_domain_raw.strip()

    # Contributing Domains — required
    contributing_raw = ov.get("contributing domains", "")
    if not contributing_raw:
        raise EntityPrdParseError(
            "Entity Overview table is missing 'Contributing Domains' "
            "row or the value is empty"
        )

    contributing_codes: list[str] = []
    entries = [e.strip() for e in contributing_raw.split(",")]
    for entry in entries:
        if not entry:
            continue
        m = _NAME_CODE_RE.match(entry)
        if m:
            contributing_codes.append(m.group(2))
        else:
            code = entry.strip()
            if code:
                contributing_codes.append(code)
                report.warn(
                    "contributing_domain_format", "Contributing Domains",
                    f"Entry {entry!r} does not match 'Name (CODE)' format",
                )

    if not contributing_codes:
        raise EntityPrdParseError(
            "Contributing Domains present but no valid codes could be "
            "extracted"
        )

    result["contributing_domain_codes"] = contributing_codes

    # Bug 6 fallback logic
    if primary_domain_code:
        result["primary_domain_code"] = primary_domain_code
        if primary_domain_code not in contributing_codes:
            report.warn(
                "primary_domain_not_in_contributing",
                "Primary Domain",
                f"Primary Domain code {primary_domain_code!r} is not in "
                f"Contributing Domains list {contributing_codes}",
            )
    else:
        # Fallback to first contributing domain
        result["primary_domain_code"] = contributing_codes[0]
        report.warn(
            "primary_domain_fallback", "Primary Domain",
            f"Primary Domain row missing — falling back to first "
            f"contributing domain: {contributing_codes[0]}",
        )

    # Discriminator fields (optional, pass-through)
    result["discriminator_field"] = ov.get("discriminator field", "")
    result["discriminator_values"] = ov.get("discriminator values", "")

    return result


def _parse_entity_overview_prose(
    paragraphs: list, report: ParseReport,
) -> str:
    """Section 1 — Entity Overview prose description."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[1])
    if h1 is None:
        raise EntityPrdParseError(
            "Document has no Heading 1 matching Section 1 (Entity Overview)"
        )
    end = _section_range(paragraphs, h1)
    text = _collect_prose(paragraphs, h1 + 1, end)
    if not text:
        report.warn(
            "empty_section", "Section 1",
            "Entity Overview section is empty",
        )
    return text


def _parse_native_fields(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 2 — Native Fields (single-row-per-field, 4+ columns)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[2])
    if h1 is None:
        raise EntityPrdParseError(
            "Document has no Heading 1 matching Section 2 (Native Fields)"
        )
    end = _section_range(paragraphs, h1)

    # Find matching table — first cell of header matches
    tables = _tables_in_range(doc, paragraphs, h1, end)
    native_table = None
    for table in tables:
        if not table.rows:
            continue
        first_cell = table.rows[0].cells[0].text.strip().lower()
        if first_cell in ("native field", "field name"):
            native_table = table
            break

    if native_table is None:
        raise EntityPrdParseError(
            "Section 2 (Native Fields) has no matching native fields table"
        )

    fields: list[dict] = []
    for row_cells in _table_rows(native_table):
        if not row_cells or not row_cells[0].strip():
            continue

        name = row_cells[0].strip()
        field_type = row_cells[1].strip().lower() if len(row_cells) > 1 else ""
        if not field_type:
            report.warn(
                "empty_field_type", f"Section 2.{name}",
                "Empty field_type for native field",
            )

        description = row_cells[2] if len(row_cells) > 2 else ""
        referenced_by = row_cells[3] if len(row_cells) > 3 else ""

        fields.append({
            "name": name,
            "field_type": field_type,
            "label": _to_label(name),
            "is_required": False,
            "default_value": None,
            "description": description,
            "identifier": None,
            "referenced_by": referenced_by,
        })

    if not fields:
        raise EntityPrdParseError(
            "Section 2 (Native Fields) table has zero data rows"
        )

    return fields


def _parse_custom_fields(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 3 — Custom Fields, organized by H2 subsections."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[3])
    if h1 is None:
        raise EntityPrdParseError(
            "Document has no Heading 1 matching Section 3 (Custom Fields)"
        )
    end = _section_range(paragraphs, h1)

    # Find all H2 subsections within Section 3
    h2_indices: list[int] = []
    for i in range(h1 + 1, end):
        if _is_heading(paragraphs[i], 2):
            h2_indices.append(i)

    subsections: list[tuple[int, int, str]] = []
    for idx, h2_idx in enumerate(h2_indices):
        sub_end = h2_indices[idx + 1] if idx + 1 < len(h2_indices) else end
        raw_title = paragraphs[h2_idx].text.strip()
        title = _LEADING_NUM_RE.sub("", raw_title)
        subsections.append((h2_idx, sub_end, title))

    fields: list[dict] = []
    for sub_start, sub_end, subsection_title in subsections:
        is_incomplete = bool(_INCOMPLETE_DOMAIN_RE.match(subsection_title))
        sub_tables = _tables_in_range(doc, paragraphs, sub_start, sub_end)
        sub_fields: list[dict] = []

        for table in sub_tables:
            header = _table_header(table)
            if len(header) < 6:
                continue
            if header[:6] != _FIELD_TABLE_HEADERS:
                continue

            data_rows = _table_rows(table)
            if not data_rows:
                continue

            i = 0
            while i < len(data_rows):
                meta_row = data_rows[i]
                if len(meta_row) < 6 or not meta_row[0].strip():
                    i += 1
                    continue

                name = meta_row[0].strip()
                field_type = meta_row[1].strip().lower()
                required_text = meta_row[2].strip()
                values = meta_row[3].strip()
                default_raw = meta_row[4].strip()
                identifier = meta_row[5].strip()

                is_required = _parse_required(
                    required_text, report, f"Section 3.{name}",
                )

                # Description row (two-row-per-field pattern)
                description = ""
                if i + 1 < len(data_rows):
                    desc_row = data_rows[i + 1]
                    description = desc_row[0].strip() if desc_row else ""
                    i += 2
                else:
                    report.warn(
                        "odd_row_count",
                        f"Section 3.{subsection_title}",
                        f"Field table has odd row count — missing "
                        f"description for {name}",
                    )
                    i += 1

                field: dict[str, Any] = {
                    "name": name,
                    "field_type": field_type,
                    "label": _to_label(name),
                    "is_required": is_required,
                    "values": values,
                    "default_value": _clean_default(default_raw),
                    "identifier": identifier,
                    "description": description,
                    "subsection": subsection_title,
                }

                # Parse enum options
                if field_type in ("enum", "multienum") and values:
                    cleaned = values.strip()
                    if cleaned not in ("—", "–", "-", "N/A", ""):
                        if "TBD" in cleaned.upper():
                            report.warn(
                                "tbd_values", f"Section 3.{name}",
                                f"Values contain TBD: {values!r}",
                            )
                        field["value_options"] = [
                            v.strip()
                            for v in cleaned.split(",")
                            if v.strip()
                        ]

                # Incomplete domain subsection per-field warning
                if is_incomplete:
                    report.warn(
                        "incomplete_domain_field", f"Field {name}",
                        "In Incomplete Domain subsection — may need "
                        "revision when source domain processes are "
                        "defined",
                    )

                sub_fields.append(field)

        if not sub_fields:
            report.warn(
                "subsection_no_field_tables",
                f"Section 3.{subsection_title}",
                "Subsection has no field tables",
            )

        fields.extend(sub_fields)

    if not fields:
        raise EntityPrdParseError(
            "Section 3 (Custom Fields) has zero fields across all "
            "subsections"
        )

    return fields


def _parse_relationships(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 4 — Relationships (5-column table)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[4])
    if h1 is None:
        raise EntityPrdParseError(
            "Document has no Heading 1 matching Section 4 (Relationships)"
        )
    end = _section_range(paragraphs, h1)

    tables = _tables_in_range(doc, paragraphs, h1, end)
    rel_table = None
    for table in tables:
        header = _table_header(table)
        if len(header) >= 5 and header[0] == "relationship":
            rel_table = table
            break

    if rel_table is None:
        raise EntityPrdParseError(
            "Section 4 (Relationships) has no matching 5-column table"
        )

    relationships: list[dict] = []
    for row_cells in _table_rows(rel_table):
        if len(row_cells) < 5 or not row_cells[0].strip():
            continue

        link_type = row_cells[2].strip()
        if link_type and link_type not in _VALID_LINK_TYPES:
            report.warn(
                "invalid_link_type",
                f"Relationship {row_cells[0].strip()}",
                f"link_type {link_type!r} not in valid set",
            )

        relationships.append({
            "name": row_cells[0].strip(),
            "related_entity": row_cells[1].strip(),
            "link_type": link_type,
            "prd_reference": row_cells[3].strip(),
            "domains": row_cells[4].strip(),
        })

    if not relationships:
        raise EntityPrdParseError(
            "Section 4 (Relationships) table has zero data rows"
        )

    return relationships


def _parse_dynamic_logic(
    paragraphs: list, report: ParseReport,
) -> list[dict]:
    """Section 5 — Dynamic Logic Rules (optional, H2 subsections)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[5])
    if h1 is None:
        report.info(
            "optional_section_absent", "Section 5",
            "Section 5 (Dynamic Logic Rules) not present — this is "
            "normal for entities without dynamic logic",
        )
        return []

    end = _section_range(paragraphs, h1)

    entries: list[dict] = []
    i = h1 + 1
    while i < end:
        p = paragraphs[i]
        if _is_heading(p, 2):
            raw_title = p.text.strip()
            num_match = re.match(r"^(\d+(?:\.\d+)?)\s+(.+)$", raw_title)
            if num_match:
                section = num_match.group(1)
                title = num_match.group(2).strip()
            else:
                section = ""
                title = raw_title

            # Find subsection end
            sub_end = end
            for j in range(i + 1, end):
                if (_is_heading(paragraphs[j], 2)
                        or _is_heading(paragraphs[j], 1)):
                    sub_end = j
                    break

            description = _collect_prose(paragraphs, i + 1, sub_end)
            if not description:
                report.warn(
                    "empty_subsection", f"Section {section}",
                    "Subsection has no prose",
                )

            entries.append({
                "section": section,
                "title": title,
                "description": description,
            })
            i = sub_end
            continue
        i += 1

    if not entries:
        report.warn(
            "empty_section", "Section 5",
            "Section 5 present but has zero subsections",
        )

    return entries


def _parse_layout_guidance(
    paragraphs: list,
) -> str:
    """Section 6 — Layout Guidance (optional, flat prose)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[6])
    if h1 is None:
        return ""
    end = _section_range(paragraphs, h1)
    return _collect_prose(paragraphs, h1 + 1, end)


def _parse_implementation_notes(paragraphs: list) -> str:
    """Section 7 — Implementation Notes (optional, flat prose)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[7])
    if h1 is None:
        return ""
    end = _section_range(paragraphs, h1)
    return _collect_prose(paragraphs, h1 + 1, end)


def _parse_open_issues(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 8 — Open Issues (required, two-column table)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[8])
    if h1 is None:
        raise EntityPrdParseError(
            "Document has no Heading 1 matching Section 8 (Open Issues)"
        )
    end = _section_range(paragraphs, h1)

    tables = _tables_in_range(doc, paragraphs, h1, end)
    issue_table = None
    for table in tables:
        header = _table_header(table)
        if (len(header) >= 2
                and header[0] == "id"
                and header[1] in ("issue", "description")):
            issue_table = table
            break

    issues: list[dict] = []
    if issue_table is not None:
        for row_cells in _table_rows(issue_table):
            if len(row_cells) < 2:
                continue
            identifier = row_cells[0].strip()
            description = row_cells[1].strip()
            if not identifier:
                continue

            if description.upper().startswith(("CLOSED", "RESOLVED")):
                report.warn(
                    "closed_issue", f"Issue {identifier}",
                    f"Issue marked as CLOSED/RESOLVED: "
                    f"{description[:60]}",
                )

            issues.append({
                "identifier": identifier,
                "description": description,
                "status": "open",
            })

    return issues


def _parse_decisions(
    paragraphs: list, doc: Document, report: ParseReport,
) -> list[dict]:
    """Section 9 — Decisions Made (required, two-column table)."""
    h1 = _find_heading1(paragraphs, _SECTION_RES[9])
    if h1 is None:
        raise EntityPrdParseError(
            "Document has no Heading 1 matching Section 9 (Decisions Made)"
        )
    end = _section_range(paragraphs, h1)

    tables = _tables_in_range(doc, paragraphs, h1, end)
    dec_table = None
    for table in tables:
        header = _table_header(table)
        if (len(header) >= 2
                and header[0] == "id"
                and header[1] == "decision"):
            dec_table = table
            break

    decisions: list[dict] = []
    if dec_table is not None:
        for row_cells in _table_rows(dec_table):
            if len(row_cells) < 2:
                continue
            identifier = row_cells[0].strip()
            description = row_cells[1].strip()
            if not identifier:
                continue

            if not _DEC_ID_RE.match(identifier):
                report.warn(
                    "decision_id_format",
                    f"Decision {identifier}",
                    "Identifier does not match pattern "
                    "^[A-Z]+-DEC-\\d+$",
                )

            decisions.append({
                "identifier": identifier,
                "description": description,
                "status": "accepted",
            })

    return decisions


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse(
    path: str | Path,
    work_item: dict,
    session_type: str = "initial",
) -> tuple[str, ParseReport]:
    """Parse an Entity PRD .docx into a Path B envelope JSON string.

    :param path: Path to the .docx file.
    :param work_item: Must have item_type == 'entity_prd' and 'entity_id' key.
    :param session_type: Session type for the envelope. Default 'initial'.
    :returns: Tuple of (envelope_json_string, ParseReport).
    :raises ValueError: If work_item['item_type'] != 'entity_prd'.
    :raises FileNotFoundError: If path does not exist.
    :raises EntityPrdParseError: If the document is structurally unparseable.
    """
    if work_item.get("item_type") != "entity_prd":
        raise ValueError(
            f"work_item['item_type'] must be 'entity_prd', "
            f"got {work_item.get('item_type')!r}"
        )

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    report = ParseReport()
    doc = Document(str(path))
    paragraphs = list(doc.paragraphs)

    # Header table
    header_meta = _parse_header_table(doc, report)

    # Entity Overview table
    overview = _parse_entity_overview_table(doc, report, header_meta)

    # Section 1: Entity Overview prose
    description = _parse_entity_overview_prose(paragraphs, report)

    # Section 2: Native Fields
    native_fields = _parse_native_fields(paragraphs, doc, report)

    # Section 3: Custom Fields
    custom_fields = _parse_custom_fields(paragraphs, doc, report)

    # Section 4: Relationships
    relationships = _parse_relationships(paragraphs, doc, report)

    # Section 5: Dynamic Logic Rules (optional)
    dynamic_logic = _parse_dynamic_logic(paragraphs, report)

    # Section 6: Layout Guidance (optional)
    layout_guidance = _parse_layout_guidance(paragraphs)

    # Section 7: Implementation Notes (optional)
    implementation_notes = _parse_implementation_notes(paragraphs)

    # Section 8: Open Issues
    open_issues = _parse_open_issues(paragraphs, doc, report)

    # Section 9: Decisions Made
    decisions = _parse_decisions(paragraphs, doc, report)

    # --- Build envelope ---

    entity_metadata: dict[str, Any] = {
        "name": overview.get("name", header_meta.get("name", "")),
        "entity_type": overview.get(
            "entity_type", header_meta.get("entity_type"),
        ),
        "is_native": header_meta.get("is_native"),
        "singular_label": overview.get("singular_label", ""),
        "plural_label": overview.get("plural_label", ""),
        "activity_stream": overview.get("activity_stream", False),
        "primary_domain_code": overview.get("primary_domain_code", ""),
        "contributing_domain_codes": overview.get(
            "contributing_domain_codes", [],
        ),
        "discriminator_field": overview.get("discriminator_field", ""),
        "discriminator_values": overview.get("discriminator_values", ""),
        "description": description,
    }

    source_metadata: dict[str, Any] = {
        "entity_name": entity_metadata["name"],
        "version": header_meta.get("version", ""),
        "status": header_meta.get("status", ""),
        "last_updated": header_meta.get("last_updated", ""),
        "source_documents": header_meta.get("source_documents", ""),
    }

    # Record parsed counts
    report.parsed_counts = {
        "native_field": len(native_fields),
        "custom_field": len(custom_fields),
        "relationship": len(relationships),
        "dynamic_logic_entry": len(dynamic_logic),
        "open_issue": len(open_issues),
        "decision": len(decisions),
    }

    payload: dict[str, Any] = {
        "source_metadata": source_metadata,
        "entity_metadata": entity_metadata,
        "native_fields": native_fields,
        "custom_fields": custom_fields,
        "relationships": relationships,
        "dynamic_logic": dynamic_logic,
        "layout_guidance": layout_guidance,
        "implementation_notes": implementation_notes,
    }

    envelope: dict[str, Any] = {
        "output_version": "1.0",
        "work_item_type": "entity_prd",
        "work_item_id": work_item["id"],
        "session_type": session_type,
        "payload": payload,
        "decisions": decisions,
        "open_issues": open_issues,
    }

    envelope_json = json.dumps(envelope, ensure_ascii=False, indent=2)
    return envelope_json, report
