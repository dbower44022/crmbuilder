"""Parse a Domain PRD .docx into a Path B envelope JSON string.

Adapter that turns a Word Domain PRD (reconciliation) document into a
Path B envelope suitable for processing by the domain_reconciliation mapper.

Replaces the Path A parser at automation.cbm_import.parsers.domain_prd.
Fixes Bug 8 (only text blobs extracted) by parsing the full document
structure including personas, per-process requirements, 7-column
consolidated field tables, 4-column decisions, and 4-column open issues.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

from automation.importer.parsers import DomainPrdParseError, ParseReport

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# "Name (CODE)" pattern for domains and processes
_NAME_CODE_RE = re.compile(r"^(.+?)\s*\(([A-Z][A-Z0-9-]*)\)\s*$")

# Leading number prefix: "3.1 " or "5 "
_LEADING_NUM_RE = re.compile(r"^\d+(?:\.\d+)?\s+")

# Section heading patterns (Heading 1)
_SECTION_RES: dict[int, re.Pattern] = {
    1: re.compile(r"^1\.\s+Domain Overview$"),
    2: re.compile(r"^2\.\s+Personas$"),
    3: re.compile(r"^3\.\s+Business Processes$"),
    4: re.compile(r"^4\.\s+Data Reference$"),
    5: re.compile(r"^5\.\s+Decisions Made$"),
    6: re.compile(r"^6\.\s+Open Issues$"),
}

# Persona paragraph pattern: "Name (IDENTIFIER)"
_PERSONA_PARA_RE = re.compile(r"^(.+?)\s*\(([A-Z]+-PER-\d+)\)\s*$")

# Entity heading pattern: "Entity: Name" or "Entity: Name (qualifier)"
_ENTITY_HEADING_RE = re.compile(r"^Entity:\s*(.+)$")

# Decision identifier pattern: MN-DEC-001, MN-RECON-DEC-001, MN-INTAKE-DEC-001
_DEC_ID_RE = re.compile(r"^[A-Z]+(?:-[A-Z]+)*-DEC-\d+$")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _style_name(p: Any) -> str:
    """Return the paragraph's style name, or '' if unavailable."""
    if p.style is None:
        return ""
    name = p.style.name
    return name if name else ""


def _is_heading(p: Any, level: int | None = None) -> bool:
    """Check whether paragraph *p* is a Word heading."""
    name = _style_name(p)
    if level is not None:
        return name == f"Heading {level}"
    return name.startswith("Heading ")


def _find_heading1(paragraphs: list, pattern: re.Pattern) -> int | None:
    """Return the index of the first Heading 1 matching *pattern*."""
    for i, p in enumerate(paragraphs):
        if _is_heading(p, 1) and pattern.match(p.text.strip()):
            return i
    return None


def _section_range(paragraphs: list, start: int) -> int:
    """Return the end index (exclusive) for a Heading 1 section."""
    for j in range(start + 1, len(paragraphs)):
        if _is_heading(paragraphs[j], 1):
            return j
    return len(paragraphs)


def _collect_prose(paragraphs: list, start: int, end: int) -> str:
    """Collect all non-heading paragraph text between start and end."""
    parts: list[str] = []
    for i in range(start, end):
        p = paragraphs[i]
        if _is_heading(p):
            continue
        text = p.text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _tables_in_range(
    doc: Document, paragraphs: list, start: int, end: int,
) -> list:
    """Return all tables whose position falls between paragraph indices."""
    body = doc.element.body
    elements = list(body)

    # Build paragraph index map: XML element -> paragraph index
    para_elem_indices: dict[int, int] = {}
    pi = 0
    for ei, elem in enumerate(elements):
        if elem.tag.endswith("}p"):
            if pi < len(paragraphs):
                para_elem_indices[ei] = pi
                pi += 1

    # Find element indices for start and end paragraphs
    start_elem_idx = None
    end_elem_idx = None
    for ei, pi_val in sorted(para_elem_indices.items()):
        if pi_val == start and start_elem_idx is None:
            start_elem_idx = ei
        if pi_val >= end and end_elem_idx is None:
            end_elem_idx = ei
            break
    if start_elem_idx is None:
        return []
    if end_elem_idx is None:
        end_elem_idx = len(elements)

    # Collect tables between those element indices
    result = []
    table_idx = 0
    for ei, elem in enumerate(elements):
        if elem.tag.endswith("}tbl"):
            if table_idx < len(doc.tables):
                if start_elem_idx <= ei < end_elem_idx:
                    result.append(doc.tables[table_idx])
                table_idx += 1
    return result


def _table_header(table: Any) -> list[str]:
    """Return lowercased, stripped header row cells."""
    if not table.rows:
        return []
    return [cell.text.strip().lower() for cell in table.rows[0].cells]


def _table_rows(table: Any) -> list[list[str]]:
    """Return data rows (excluding header) as lists of stripped cell strings."""
    rows = []
    for i in range(1, len(table.rows)):
        rows.append([cell.text.strip() for cell in table.rows[i].cells])
    return rows


# ---------------------------------------------------------------------------
# Section parsers
# ---------------------------------------------------------------------------


def _parse_header_table(
    doc: Document, report: ParseReport,
) -> dict[str, str]:
    """Extract source metadata from the header key/value table."""
    # Find header table: first 2-col table with "domain" in col 0
    header_table = None
    for table in doc.tables:
        if len(table.columns) != 2:
            continue
        for row in table.rows:
            if "domain" in row.cells[0].text.strip().lower():
                header_table = table
                break
        if header_table:
            break

    if header_table is None:
        raise DomainPrdParseError("No header metadata table found")

    metadata: dict[str, str] = {}
    key_map = {
        "domain": "domain_raw",
        "domain code": "domain_code",
        "version": "version",
        "status": "status",
        "last updated": "last_updated",
        "source": "source",
        "processes": "processes",
    }

    for row in header_table.rows:
        key = row.cells[0].text.strip().lower()
        value = row.cells[1].text.strip()
        mapped = key_map.get(key)
        if mapped:
            metadata[mapped] = value

    # Extract domain code from "Domain" row
    domain_raw = metadata.pop("domain_raw", "")
    domain_name = ""
    domain_code = ""

    if domain_raw:
        m = _NAME_CODE_RE.match(domain_raw)
        if m:
            domain_name = m.group(1).strip()
            domain_code = m.group(2)
        else:
            domain_name = domain_raw
            report.warn(
                "domain_name_code_parse",
                "Header table",
                f"Domain value '{domain_raw}' doesn't match 'Name (CODE)' format",
            )

    # Cross-check with Domain Code row
    explicit_code = metadata.pop("domain_code", "")
    if domain_code and explicit_code and domain_code != explicit_code:
        report.warn(
            "domain_code_mismatch",
            "Header table",
            f"Domain row code '{domain_code}' differs from Domain Code "
            f"row '{explicit_code}'",
        )

    # Fallback
    if not domain_code:
        domain_code = explicit_code

    if not domain_code:
        raise DomainPrdParseError(
            "No domain code extractable from header table "
            "(neither Domain row nor Domain Code row)"
        )

    metadata["domain_name"] = domain_name
    metadata["domain_code"] = domain_code

    # Warn on missing optional fields
    for field in ("version", "status"):
        if field not in metadata:
            report.warn(
                "missing_header_field",
                "Header table",
                f"Missing expected row: {field}",
            )

    return metadata


def _parse_domain_overview(
    paragraphs: list, report: ParseReport,
) -> str:
    """Extract Section 1 Domain Overview prose."""
    idx = _find_heading1(paragraphs, _SECTION_RES[1])
    if idx is None:
        raise DomainPrdParseError("Missing Section 1: Domain Overview")

    end = _section_range(paragraphs, idx)
    prose = _collect_prose(paragraphs, idx + 1, end)

    if not prose:
        report.warn(
            "empty_section",
            "Section 1",
            "Domain Overview section is empty",
        )

    return prose


def _parse_personas(
    paragraphs: list, report: ParseReport,
) -> list[dict[str, Any]]:
    """Extract Section 2 Personas.

    Handles paragraph-based format where each persona is a Normal paragraph
    with "Name (IDENTIFIER)" followed by a description paragraph.
    """
    idx = _find_heading1(paragraphs, _SECTION_RES[2])
    if idx is None:
        report.info(
            "missing_section",
            "Section 2",
            "Personas section not found — empty list",
        )
        return []

    end = _section_range(paragraphs, idx)
    personas: list[dict[str, Any]] = []

    # Try paragraph-based format: "Name (IDENTIFIER)" followed by description
    i = idx + 1
    while i < end:
        p = paragraphs[i]
        if _is_heading(p):
            # If H2-based, extract from heading
            if _is_heading(p, 2):
                name = _LEADING_NUM_RE.sub("", p.text.strip())
                # Collect prose until next H2 or H1
                j = i + 1
                while j < end and not _is_heading(paragraphs[j], 2) and not _is_heading(paragraphs[j], 1):
                    j += 1
                desc = _collect_prose(paragraphs, i + 1, j)
                personas.append({
                    "identifier": None,
                    "code": name,
                    "name": name,
                    "consolidated_role": desc,
                    "description": desc,
                })
                i = j
                continue
            i += 1
            continue

        text = p.text.strip()
        if not text:
            i += 1
            continue

        # Check if this is a persona name line
        m = _PERSONA_PARA_RE.match(text)
        if m:
            name = m.group(1).strip()
            identifier = m.group(2)
            # Next non-empty paragraph is description
            desc = ""
            j = i + 1
            while j < end:
                next_p = paragraphs[j]
                next_text = next_p.text.strip()
                if not next_text or _is_heading(next_p):
                    j += 1
                    continue
                # Check if it's another persona name
                if _PERSONA_PARA_RE.match(next_text):
                    break
                desc = next_text
                j += 1
                break
            personas.append({
                "identifier": identifier,
                "code": identifier,
                "name": name,
                "consolidated_role": desc,
                "description": desc,
            })
            i = j
            continue

        i += 1

    if not personas:
        report.warn(
            "empty_personas",
            "Section 2",
            "Personas section present but no personas found",
        )

    return personas


def _parse_process_summaries(
    doc: Document, paragraphs: list, report: ParseReport,
) -> list[dict[str, Any]]:
    """Extract Section 3 Business Processes — per-process summaries."""
    idx = _find_heading1(paragraphs, _SECTION_RES[3])
    if idx is None:
        report.warn(
            "missing_section",
            "Section 3",
            "Business Processes section not found",
        )
        return []

    end = _section_range(paragraphs, idx)
    summaries: list[dict[str, Any]] = []

    # Find all H2 subsections
    h2_indices: list[int] = []
    for i in range(idx + 1, end):
        if _is_heading(paragraphs[i], 2):
            h2_indices.append(i)

    for hi, h2_idx in enumerate(h2_indices):
        h2_text = paragraphs[h2_idx].text.strip()
        stripped = _LEADING_NUM_RE.sub("", h2_text)

        # Extract process code from "Name (CODE)" pattern
        m = _NAME_CODE_RE.match(stripped)
        if m:
            process_name = m.group(1).strip()
            process_code = m.group(2)
        else:
            process_name = stripped
            process_code = stripped
            report.warn(
                "process_code_missing",
                f"Section 3 {h2_text}",
                f"Could not extract process code from heading: {stripped!r}",
            )

        # Subsection range
        sub_end = h2_indices[hi + 1] if hi + 1 < len(h2_indices) else end

        # Collect prose (including H3 content)
        desc = _collect_prose(paragraphs, h2_idx + 1, sub_end)

        # Find requirement tables in this subsection
        tables = _tables_in_range(doc, paragraphs, h2_idx, sub_end)
        requirements: list[dict[str, str]] = []
        for table in tables:
            header = _table_header(table)
            if len(header) >= 2 and "id" in header[0] and "requirement" in header[1]:
                for row_data in _table_rows(table):
                    if len(row_data) >= 2 and row_data[0].strip():
                        identifier = row_data[0].strip()
                        requirements.append({
                            "identifier": identifier,
                            "description": row_data[1].strip(),
                        })

        summaries.append({
            "process_code": process_code,
            "process_name": process_name,
            "description": desc,
            "requirements": requirements,
        })

    return summaries


def _parse_data_reference(
    doc: Document, paragraphs: list, report: ParseReport,
) -> list[dict[str, Any]]:
    """Extract Section 4 Data Reference — per-entity field tables."""
    idx = _find_heading1(paragraphs, _SECTION_RES[4])
    if idx is None:
        report.info(
            "missing_section",
            "Section 4",
            "Data Reference section not found",
        )
        return []

    end = _section_range(paragraphs, idx)
    entities: list[dict[str, Any]] = []

    # Find all H2 subsections matching "Entity: ..."
    h2_indices: list[tuple[int, str]] = []
    for i in range(idx + 1, end):
        if _is_heading(paragraphs[i], 2):
            text = paragraphs[i].text.strip()
            m = _ENTITY_HEADING_RE.match(text)
            if m:
                entity_name = m.group(1).strip()
                # Strip parenthetical qualifier if present
                paren_idx = entity_name.find("(")
                if paren_idx > 0:
                    entity_name = entity_name[:paren_idx].strip()
                h2_indices.append((i, entity_name))

    for hi, (h2_idx, entity_name) in enumerate(h2_indices):
        sub_end_idx = h2_indices[hi + 1][0] if hi + 1 < len(h2_indices) else end

        # Find 7-column field tables in this subsection
        tables = _tables_in_range(doc, paragraphs, h2_idx, sub_end_idx)
        fields: list[dict[str, str]] = []

        for table in tables:
            header = _table_header(table)
            if len(header) < 7:
                continue
            if "field" not in header[0] and "name" not in header[0]:
                continue

            # Two-row-per-field pattern
            data_rows = _table_rows(table)
            i = 0
            while i < len(data_rows):
                meta_row = data_rows[i]
                if len(meta_row) < 7:
                    i += 1
                    continue

                # Description row follows
                desc = ""
                if i + 1 < len(data_rows):
                    desc_row = data_rows[i + 1]
                    desc = desc_row[0] if desc_row else ""
                    i += 2
                else:
                    report.warn(
                        "missing_description_row",
                        f"Entity {entity_name}",
                        f"Field '{meta_row[0]}' missing description row",
                    )
                    i += 1

                fields.append({
                    "name": meta_row[0],
                    "field_type": meta_row[1],
                    "is_required": meta_row[2],
                    "values": meta_row[3],
                    "default_value": meta_row[4],
                    "identifier": meta_row[5],
                    "defined_in": meta_row[6],
                    "description": desc,
                })

        if not fields:
            report.warn(
                "no_field_tables",
                f"Entity {entity_name}",
                "No 7-column field tables found in entity subsection",
            )

        entities.append({
            "entity_name": entity_name,
            "deduplicated_fields": fields,
        })

    return entities


def _parse_decisions(
    doc: Document, paragraphs: list, report: ParseReport,
) -> list[dict[str, str]]:
    """Extract Section 5 Decisions Made — 4-column table."""
    idx = _find_heading1(paragraphs, _SECTION_RES[5])
    if idx is None:
        report.warn(
            "missing_section",
            "Section 5",
            "Decisions Made section not found",
        )
        return []

    end = _section_range(paragraphs, idx)
    tables = _tables_in_range(doc, paragraphs, idx, end)

    decisions: list[dict[str, str]] = []
    for table in tables:
        header = _table_header(table)
        if len(header) < 4:
            continue
        if "id" not in header[0] or "decision" not in header[1]:
            continue

        for row_data in _table_rows(table):
            if len(row_data) < 4:
                continue
            identifier = row_data[0].strip()
            if not identifier or identifier.lower() == "id":
                continue

            if not _DEC_ID_RE.match(identifier):
                report.warn(
                    "decision_id_format",
                    f"Decision {identifier}",
                    f"Identifier doesn't match expected pattern: {identifier!r}",
                )

            decisions.append({
                "identifier": identifier,
                "description": row_data[1].strip(),
                "rationale": row_data[2].strip(),
                "made_during": row_data[3].strip(),
                "status": "locked",
            })

    return decisions


def _parse_open_issues(
    doc: Document, paragraphs: list, report: ParseReport,
) -> list[dict[str, str]]:
    """Extract Section 6 Open Issues — 4-column table."""
    idx = _find_heading1(paragraphs, _SECTION_RES[6])
    if idx is None:
        report.warn(
            "missing_section",
            "Section 6",
            "Open Issues section not found",
        )
        return []

    end = _section_range(paragraphs, idx)
    tables = _tables_in_range(doc, paragraphs, idx, end)

    issues: list[dict[str, str]] = []
    for table in tables:
        header = _table_header(table)
        if len(header) < 4:
            continue
        if "id" not in header[0] or "issue" not in header[1]:
            continue

        for row_data in _table_rows(table):
            if len(row_data) < 4:
                continue
            identifier = row_data[0].strip()
            if not identifier or identifier.lower() == "id":
                continue

            issues.append({
                "identifier": identifier,
                "description": row_data[1].strip(),
                "question": row_data[2].strip(),
                "needs_input_from": row_data[3].strip(),
                "status": "open",
            })

    return issues


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse(
    path: str | Path,
    work_item: dict,
    session_type: str = "initial",
) -> tuple[str, ParseReport]:
    """Parse a Domain PRD .docx into a Path B envelope JSON string.

    :param path: Path to Domain PRD .docx file.
    :param work_item: Must have item_type == 'domain_reconciliation'.
    :param session_type: Session type for envelope metadata.
    :returns: Tuple of (JSON envelope string, ParseReport).
    :raises ValueError: If work_item['item_type'] != 'domain_reconciliation'.
    :raises FileNotFoundError: If path does not exist.
    :raises DomainPrdParseError: If document is structurally unparseable.
    """
    if work_item.get("item_type") != "domain_reconciliation":
        msg = (
            f"Expected work_item item_type='domain_reconciliation', "
            f"got {work_item.get('item_type')!r}"
        )
        raise ValueError(msg)

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Domain PRD file not found: {path}")

    report = ParseReport()
    doc = Document(str(path))
    paragraphs = list(doc.paragraphs)

    # Parse all sections
    source_metadata = _parse_header_table(doc, report)
    domain_overview = _parse_domain_overview(paragraphs, report)
    personas = _parse_personas(paragraphs, report)
    process_summaries = _parse_process_summaries(doc, paragraphs, report)
    data_reference = _parse_data_reference(doc, paragraphs, report)
    decisions = _parse_decisions(doc, paragraphs, report)
    open_issues = _parse_open_issues(doc, paragraphs, report)

    # Update parsed counts
    report.parsed_counts["personas"] = len(personas)
    report.parsed_counts["process_summaries"] = len(process_summaries)
    report.parsed_counts["data_reference_entities"] = len(data_reference)
    report.parsed_counts["decisions"] = len(decisions)
    report.parsed_counts["open_issues"] = len(open_issues)

    # Build envelope
    envelope = {
        "output_version": "1.0",
        "work_item_type": "domain_reconciliation",
        "work_item_id": work_item["id"],
        "session_type": session_type,
        "payload": {
            "source_metadata": source_metadata,
            "domain_overview_narrative": domain_overview,
            "personas": personas,
            "process_summaries": process_summaries,
            "consolidated_data_reference": data_reference,
        },
        "decisions": decisions,
        "open_issues": open_issues,
    }

    return json.dumps(envelope, indent=2, ensure_ascii=False), report
