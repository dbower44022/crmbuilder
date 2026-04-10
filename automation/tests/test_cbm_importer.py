"""Tests for CBM importer parsers against fixture subset documents."""

from pathlib import Path

from automation.cbm_import.parser_logic import (
    extract_enum_values,
    extract_section_text,
    find_section_by_heading,
    parse_field_table,
    parse_header_table,
    parse_persona_list,
    parse_requirement_list,
)
from automation.cbm_import.reporter import ImportReport

FIXTURES = Path(__file__).parent / "fixtures" / "cbm_subset"


# ---------------------------------------------------------------------------
# parser_logic tests (pure Python, no docx)
# ---------------------------------------------------------------------------


class TestFindSectionByHeading:

    def test_finds_heading(self):
        paras = ["Intro", "2. Personas", "Some content"]
        assert find_section_by_heading(paras, r"personas") == 1

    def test_returns_none_when_missing(self):
        paras = ["Intro", "Content"]
        assert find_section_by_heading(paras, r"personas") is None

    def test_case_insensitive(self):
        paras = ["1. ORGANIZATION OVERVIEW"]
        assert find_section_by_heading(paras, r"organization\s+overview") == 0


class TestExtractSectionText:

    def test_extracts_until_next_heading(self):
        paras = ["1. Heading", "Content line 1", "Content line 2", "2. Next heading"]
        text = extract_section_text(paras, 0)
        assert "Content line 1" in text
        assert "Content line 2" in text
        assert "Next heading" not in text

    def test_extracts_to_end(self):
        paras = ["1. Heading", "Only content"]
        text = extract_section_text(paras, 0)
        assert text == "Only content"


class TestParseHeaderTable:

    def test_parses_key_value_pairs(self):
        rows = [
            ["Document Type", "Master PRD"],
            ["Version", "2.3"],
        ]
        result = parse_header_table(rows)
        assert result["Document Type"] == "Master PRD"
        assert result["Version"] == "2.3"


class TestParseFieldTable:

    def test_parses_fields(self):
        rows = [
            ["Field Name", "Type", "Required", "Values", "Default", "ID"],
            ["contactType", "multiEnum", "Yes", "Client | Mentor", "—", "DAT-001"],
            ["preferredName", "varchar", "No", "", "—", "DAT-002"],
        ]
        fields = parse_field_table(rows)
        assert len(fields) == 2
        assert fields[0]["field_name"] == "contactType"
        assert fields[0]["field_type"] == "multiEnum"
        assert fields[1]["field_name"] == "preferredName"

    def test_skips_empty_rows(self):
        rows = [
            ["Field Name", "Type"],
            ["", ""],
            ["email", "email"],
        ]
        fields = parse_field_table(rows)
        assert len(fields) == 1

    def test_merged_description_rows_not_parsed_as_fields(self):
        """Merged description rows (python-docx repeats text in every cell) must be
        attached to the previous field, not created as separate field records."""
        desc1 = "The discriminator field for the Contact entity type classification."
        desc2 = "The name the contact prefers to be called in informal settings."
        desc3 = "Primary email address used for all CRM communications."
        rows = [
            ["Field Name", "Type", "Required", "Values", "Default", "ID"],
            ["contactType", "multiEnum", "Yes", "Client | Mentor", "—", "DAT-001"],
            # Merged description row: same text in every cell
            [desc1, desc1, desc1, desc1, desc1, desc1],
            ["preferredName", "varchar", "No", "", "—", "DAT-002"],
            [desc2, desc2, desc2, desc2, desc2, desc2],
            ["primaryEmail", "email", "Yes", "", "—", "DAT-003"],
            [desc3, desc3, desc3, desc3, desc3, desc3],
        ]
        fields = parse_field_table(rows)
        assert len(fields) == 3, f"Expected 3 fields, got {len(fields)}: {[f.get('field_name') for f in fields]}"
        assert fields[0]["field_name"] == "contactType"
        assert fields[0]["description"] == desc1
        assert fields[1]["field_name"] == "preferredName"
        assert fields[1]["description"] == desc2
        assert fields[2]["field_name"] == "primaryEmail"
        assert fields[2]["description"] == desc3

    def test_long_field_name_treated_as_description(self):
        """Field names over 200 chars are descriptions, not real fields."""
        long_desc = "A" * 250
        rows = [
            ["Field Name", "Type"],
            ["contactType", "enum"],
            [long_desc, ""],
        ]
        fields = parse_field_table(rows)
        assert len(fields) == 1
        assert fields[0]["field_name"] == "contactType"
        assert fields[0]["description"] == long_desc


class TestExtractEnumValues:

    def test_pipe_separated(self):
        values = extract_enum_values("Client | Mentor | Partner")
        assert values == ["Client", "Mentor", "Partner"]

    def test_comma_separated(self):
        values = extract_enum_values("Active, Inactive, On Hold")
        assert values == ["Active", "Inactive", "On Hold"]

    def test_dash_returns_empty(self):
        assert extract_enum_values("—") == []

    def test_empty_returns_empty(self):
        assert extract_enum_values("") == []


class TestParsePersonaList:

    def test_extracts_personas(self):
        paras = [
            "2. Personas",
            "MST-PER-001: System Administrator — Manages CRM",
            "MST-PER-002: Executive Member — Strategic oversight",
            "3. Next section",
        ]
        personas = parse_persona_list(paras, 0)
        assert len(personas) == 2
        assert personas[0]["code"] == "MST-PER-001"
        assert personas[0]["name"] == "System Administrator"


class TestParseRequirementList:

    def test_extracts_requirements(self):
        paras = [
            "6. System Requirements",
            "MN-INTAKE-REQ-001: Accept form submissions",
            "MN-INTAKE-REQ-002: Auto-create records",
            "7. Process Data",
        ]
        reqs = parse_requirement_list(paras, 0)
        assert len(reqs) == 2
        assert reqs[0]["identifier"] == "MN-INTAKE-REQ-001"


# ---------------------------------------------------------------------------
# Reporter tests
# ---------------------------------------------------------------------------


class TestImportReport:

    def test_record_counts(self):
        report = ImportReport()
        report.record_parsed("Entity", 3)
        report.record_imported("Entity", 2)
        assert report.total_parsed == 3
        assert report.total_imported == 2

    def test_skipped_records(self):
        report = ImportReport()
        report.record_skipped("test.docx", "Field", "email", "Duplicate")
        assert len(report.skipped) == 1
        assert report.skipped[0].reason == "Duplicate"

    def test_merge(self):
        r1 = ImportReport()
        r1.record_parsed("Entity", 2)
        r2 = ImportReport()
        r2.record_parsed("Entity", 3)
        r2.add_warning("test warning")
        r1.merge(r2)
        assert r1.parsed["Entity"] == 5
        assert len(r1.warnings) == 1

    def test_summary(self):
        report = ImportReport()
        report.record_parsed("Entity", 3)
        report.record_imported("Entity", 2)
        text = report.summary()
        assert "Parsed:   3 records" in text
        assert "Imported: 2 records" in text


# ---------------------------------------------------------------------------
# Parser tests against fixture documents
# ---------------------------------------------------------------------------


class TestMasterPrdParser:

    def test_parse_fixture(self):
        from automation.cbm_import.parsers.master_prd import parse
        data, report = parse(FIXTURES / "CBM-Master-PRD.docx")

        assert len(data["personas"]) >= 3
        codes = {p["code"] for p in data["personas"]}
        assert "MST-PER-001" in codes
        assert "MST-PER-003" in codes

        assert len(data["domains"]) >= 2
        domain_codes = {d["code"] for d in data["domains"]}
        assert "MN" in domain_codes
        assert "MR" in domain_codes

        assert len(data["processes"]) >= 2
        proc_codes = {p["code"] for p in data["processes"]}
        assert "MN-INTAKE" in proc_codes

        assert not report.errors


class TestEntityInventoryParser:

    def test_parse_fixture(self):
        from automation.cbm_import.parsers.entity_inventory import parse
        data, report = parse(FIXTURES / "CBM-Entity-Inventory.docx")

        assert len(data["business_objects"]) >= 2
        assert len(data["entities"]) >= 2

        entity_names = {e["name"] for e in data["entities"]}
        assert "Contact" in entity_names
        assert "Account" in entity_names

        assert not report.errors


class TestEntityPrdParser:

    def test_parse_contact_fixture(self):
        from automation.cbm_import.parsers.entity_prd import parse
        data, report = parse(FIXTURES / "entities" / "Contact-Entity-PRD.docx")

        assert data["entity"]["name"] == "Contact"
        assert data["entity"]["is_native"] is True

        assert len(data["fields"]) >= 3
        field_names = {f["name"] for f in data["fields"]}
        assert "contactType" in field_names
        assert "preferredName" in field_names

        # Check enum options extracted
        assert len(data["field_options"]) >= 3

        assert not report.errors


class TestProcessDocumentParser:

    def test_parse_intake_fixture(self):
        from automation.cbm_import.parsers.process_document import parse
        data, report = parse(FIXTURES / "MN" / "MN-INTAKE.docx")

        assert data["process"]["code"] == "MN-INTAKE"
        assert data["process"]["domain_code"] == "MN"
        assert "Client Intake" in data["process"]["name"]

        assert len(data["steps"]) >= 3
        assert len(data["requirements"]) >= 2

        req_ids = {r["identifier"] for r in data["requirements"]}
        assert "MN-INTAKE-REQ-001" in req_ids

        assert not report.errors


class TestProcessDocumentHeading2Steps:
    """Bug #5: workflow step extractor must not break on Heading 2 subsections."""

    def test_heading2_subsections_yield_all_steps(self):
        from automation.cbm_import.parsers.process_document import parse
        data, report = parse(FIXTURES / "TD" / "TD-SUBSECT.docx")

        assert len(data["steps"]) == 5, (
            f"Expected 5 steps across two Heading 2 subsections, got {len(data['steps'])}"
        )
        assert not report.errors


class TestProcessDocumentHeaderlessRequirements:
    """Bug #6: requirements table without a header row must be detected."""

    def test_headerless_requirements_detected(self):
        from automation.cbm_import.parsers.process_document import parse
        data, report = parse(FIXTURES / "TD" / "TD-NOHDR.docx")

        assert len(data["requirements"]) == 2, (
            f"Expected 2 requirements from headerless table, got {len(data['requirements'])}"
        )
        req_ids = {r["identifier"] for r in data["requirements"]}
        assert "TD-NOHDR-REQ-001" in req_ids
        assert "TD-NOHDR-REQ-002" in req_ids
        assert not report.errors

    def test_dat_table_not_detected_as_requirements(self):
        """Field tables (-DAT- identifiers) must not be mistaken for requirements."""
        from automation.cbm_import.parsers.process_document import parse
        data, _ = parse(FIXTURES / "TD" / "TD-NOHDR.docx")

        for req in data["requirements"]:
            assert "-DAT-" not in req["identifier"], (
                f"Field table row wrongly detected as requirement: {req['identifier']}"
            )


class TestImporterSubdomainRecursion:
    """Bug #8: process discovery must recurse into sub-domain directories."""

    def test_rglob_discovers_nested_processes(self, tmp_path):
        """Process docs in sub-directories must be found by the importer."""
        from docx import Document

        # Build a minimal fixture PRDs tree with nested sub-domain
        prds = tmp_path / "PRDs"
        prds.mkdir()

        # Master PRD (required by importer)
        doc = Document()
        doc.add_table(rows=1, cols=2)
        doc.save(str(prds / "CBM-Master-PRD.docx"))

        # Entity Inventory (required by importer)
        doc = Document()
        doc.add_table(rows=1, cols=2)
        doc.save(str(prds / "CBM-Entity-Inventory.docx"))

        # Domain directory with a top-level process and a nested one
        td_dir = prds / "TD"
        td_dir.mkdir()

        # Domain overview (should be filtered out)
        doc = Document()
        doc.save(str(td_dir / "CBM-Domain-Overview-TestDomain.docx"))

        # Top-level process
        doc = Document()
        header = doc.add_table(rows=2, cols=2)
        header.cell(0, 0).text = "Process Code"
        header.cell(0, 1).text = "TD-DIRECT"
        header.cell(1, 0).text = "Domain"
        header.cell(1, 1).text = "TD"
        doc.add_heading("4. Process Workflow", level=1)
        doc.add_paragraph("1. Direct step", style="List Paragraph")
        doc.save(str(td_dir / "TD-DIRECT.docx"))

        # Sub-domain directory
        sub_dir = td_dir / "SUB"
        sub_dir.mkdir()

        # SubDomain overview (should be filtered out)
        doc = Document()
        doc.save(str(sub_dir / "CBM-SubDomain-Overview-Sub.docx"))

        # Nested process
        doc = Document()
        header = doc.add_table(rows=2, cols=2)
        header.cell(0, 0).text = "Process Code"
        header.cell(0, 1).text = "TD-SUB-NESTED"
        header.cell(1, 0).text = "Domain"
        header.cell(1, 1).text = "TD"
        doc.add_heading("4. Process Workflow", level=1)
        doc.add_paragraph("1. Nested step", style="List Paragraph")
        doc.save(str(sub_dir / "TD-SUB-NESTED.docx"))

        # Run dry-run via the importer
        from automation.cbm_import.importer import CBMImporter
        importer = CBMImporter(
            client_db_path=tmp_path / "client.db",
            master_db_path=tmp_path / "master.db",
            cbm_repo_path=tmp_path,
        )
        report = importer.import_all(dry_run=True)

        # Both processes should be discovered
        assert report.parsed.get("Process", 0) >= 2, (
            f"Expected at least 2 processes (top-level + nested), "
            f"got {report.parsed.get('Process', 0)}"
        )


class TestImportEntityPrdRelationships:
    """Bug #9: relationships parsed by Entity PRD must be written to DB."""

    def _make_entity_prd_with_relationship(self, path, entity_name, rel_name, rel_target, link_type):
        """Create a minimal Entity PRD .docx with one relationship."""
        from docx import Document
        doc = Document()
        header = doc.add_table(rows=3, cols=2)
        header.cell(0, 0).text = "Entity"
        header.cell(0, 1).text = entity_name
        header.cell(1, 0).text = "Entity Type"
        header.cell(1, 1).text = "Base"
        header.cell(2, 0).text = "Native/Custom"
        header.cell(2, 1).text = "Custom"

        doc.add_heading("3. Relationships", level=1)
        rel_table = doc.add_table(rows=2, cols=3)
        rel_table.cell(0, 0).text = "Relationship Name"
        rel_table.cell(0, 1).text = "Type"
        rel_table.cell(0, 2).text = "Related Entity"
        rel_table.cell(1, 0).text = rel_name
        rel_table.cell(1, 1).text = link_type
        rel_table.cell(1, 2).text = rel_target
        doc.save(str(path))

    def test_relationship_imported_to_db(self, tmp_path):
        from automation.cbm_import.importer import CBMImporter
        from automation.db.migrations import (
            run_client_migrations,
            run_master_migrations,
        )

        client_db = tmp_path / "client.db"
        master_db = tmp_path / "master.db"
        master_conn = run_master_migrations(str(master_db))
        master_conn.close()
        conn = run_client_migrations(str(client_db))

        # Seed two entities
        conn.execute(
            "INSERT INTO Entity (name, code, entity_type, is_native) VALUES (?, ?, ?, ?)",
            ("Contact", "CONTACT", "Person", True),
        )
        conn.execute(
            "INSERT INTO Entity (name, code, entity_type, is_native) VALUES (?, ?, ?, ?)",
            ("Engagement", "ENGAGEMENT", "Base", False),
        )
        conn.commit()

        # Create a minimal work item so _create_session works
        conn.execute(
            "INSERT INTO WorkItem (item_type, status) VALUES ('entity_prd', 'in_progress')"
        )
        conn.commit()
        conn.close()

        prd_path = tmp_path / "PRDs" / "entities"
        prd_path.mkdir(parents=True)
        self._make_entity_prd_with_relationship(
            prd_path / "Contact-Entity-PRD.docx",
            "Contact", "engagements", "Engagement", "oneToMany",
        )

        importer = CBMImporter(str(client_db), str(master_db), str(tmp_path))
        importer._conn = run_client_migrations(str(client_db))
        report = importer.import_entity_prd("Contact", prd_path / "Contact-Entity-PRD.docx")
        importer._conn.close()

        import sqlite3
        conn = sqlite3.connect(str(client_db))
        rows = conn.execute(
            "SELECT name, link_type, entity_id, entity_foreign_id FROM Relationship"
        ).fetchall()
        conn.close()

        assert len(rows) == 1
        assert rows[0][0] == "engagements"
        assert rows[0][1] == "oneToMany"
        assert report.imported.get("Relationship", 0) == 1

    def test_relationship_skipped_when_target_missing(self, tmp_path):
        from automation.cbm_import.importer import CBMImporter
        from automation.db.migrations import (
            run_client_migrations,
            run_master_migrations,
        )

        client_db = tmp_path / "client.db"
        master_db = tmp_path / "master.db"
        master_conn = run_master_migrations(str(master_db))
        master_conn.close()
        conn = run_client_migrations(str(client_db))

        # Only seed Contact — Engagement does NOT exist
        conn.execute(
            "INSERT INTO Entity (name, code, entity_type, is_native) VALUES (?, ?, ?, ?)",
            ("Contact", "CONTACT", "Person", True),
        )
        conn.execute(
            "INSERT INTO WorkItem (item_type, status) VALUES ('entity_prd', 'in_progress')"
        )
        conn.commit()
        conn.close()

        prd_path = tmp_path / "PRDs" / "entities"
        prd_path.mkdir(parents=True)
        self._make_entity_prd_with_relationship(
            prd_path / "Contact-Entity-PRD.docx",
            "Contact", "engagements", "Engagement", "oneToMany",
        )

        importer = CBMImporter(str(client_db), str(master_db), str(tmp_path))
        importer._conn = run_client_migrations(str(client_db))
        report = importer.import_entity_prd("Contact", prd_path / "Contact-Entity-PRD.docx")
        importer._conn.close()

        import sqlite3
        conn = sqlite3.connect(str(client_db))
        rel_count = conn.execute("SELECT COUNT(*) FROM Relationship").fetchone()[0]
        conn.close()

        assert rel_count == 0
        assert report.imported.get("Relationship", 0) == 0
        assert any("Target entity not found" in s.reason for s in report.skipped)


class TestImportServiceProcess:
    """Bug #10: service documents must auto-create a Services domain."""

    def _make_service_process(self, path, code):
        """Create a minimal service process .docx with empty Domain."""
        from docx import Document
        doc = Document()
        header = doc.add_table(rows=3, cols=2)
        header.cell(0, 0).text = "Process Code"
        header.cell(0, 1).text = code
        header.cell(1, 0).text = "Domain"
        header.cell(1, 1).text = ""  # Empty — services have no domain
        header.cell(2, 0).text = "Process Name"
        header.cell(2, 1).text = "Test Service Process"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("A service process for testing.")
        doc.add_heading("4. Process Workflow", level=1)
        doc.add_paragraph("1. Single service step")
        doc.save(str(path))

    def test_service_process_creates_svc_domain(self, tmp_path):
        from automation.cbm_import.importer import CBMImporter
        from automation.db.migrations import (
            run_client_migrations,
            run_master_migrations,
        )

        client_db = tmp_path / "client.db"
        master_db = tmp_path / "master.db"
        master_conn = run_master_migrations(str(master_db))
        master_conn.close()
        conn = run_client_migrations(str(client_db))
        conn.execute(
            "INSERT INTO WorkItem (item_type, status) VALUES ('process_definition', 'in_progress')"
        )
        conn.commit()
        conn.close()

        proc_path = tmp_path / "SVC-TEST.docx"
        self._make_service_process(proc_path, "SVC-TEST")

        importer = CBMImporter(str(client_db), str(master_db), str(tmp_path))
        importer._conn = run_client_migrations(str(client_db))
        report = importer.import_process("SVC-TEST", proc_path, is_service=True)
        importer._conn.close()

        import sqlite3
        conn = sqlite3.connect(str(client_db))

        # Services domain should exist
        svc_domain = conn.execute(
            "SELECT id, name, code, is_service FROM Domain WHERE code = 'SVC'"
        ).fetchone()
        assert svc_domain is not None
        assert svc_domain[3] == 1  # is_service=True

        # Process should reference the SVC domain
        proc = conn.execute(
            "SELECT code, domain_id FROM Process WHERE code = 'SVC-TEST'"
        ).fetchone()
        assert proc is not None
        assert proc[1] == svc_domain[0]

        conn.close()
        assert report.imported.get("Process", 0) == 1

    def test_regular_process_does_not_use_svc_domain(self, tmp_path):
        from automation.cbm_import.importer import CBMImporter
        from automation.db.migrations import (
            run_client_migrations,
            run_master_migrations,
        )

        client_db = tmp_path / "client.db"
        master_db = tmp_path / "master.db"
        master_conn = run_master_migrations(str(master_db))
        master_conn.close()
        conn = run_client_migrations(str(client_db))
        conn.execute(
            "INSERT INTO WorkItem (item_type, status) VALUES ('process_definition', 'in_progress')"
        )
        conn.commit()
        conn.close()

        # Service process with empty domain but is_service=False
        proc_path = tmp_path / "XX-TEST.docx"
        self._make_service_process(proc_path, "XX-TEST")

        importer = CBMImporter(str(client_db), str(master_db), str(tmp_path))
        importer._conn = run_client_migrations(str(client_db))
        report = importer.import_process("XX-TEST", proc_path, is_service=False)
        importer._conn.close()

        # Should be skipped — no domain, not a service
        assert report.imported.get("Process", 0) == 0
        assert any("Could not resolve domain" in s.reason for s in report.skipped)


class TestDomainPrdParser:

    def test_parse_mentoring_fixture(self):
        from automation.cbm_import.parsers.domain_prd import parse
        data, report = parse(FIXTURES / "MN" / "CBM-Domain-PRD-Mentoring.docx")

        assert data["domain_code"] == "MN"
        assert len(data["domain_overview_text"]) > 50

        assert len(data["decisions"]) >= 2
        dec_ids = {d["identifier"] for d in data["decisions"]}
        assert "MN-DEC-001" in dec_ids

        assert not report.errors
