"""Tests for automation.importer.parsers.domain_prd_docx."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from automation.importer.parsers import DomainPrdParseError, ParseReport
from automation.importer.parsers.domain_prd_docx import parse

FIXTURE = Path(__file__).parent / "fixtures" / "cbm-domain-prd-mentoring-v1.0.docx"

skip_fixture = pytest.mark.skipif(
    not FIXTURE.exists(),
    reason="cbm-domain-prd-mentoring-v1.0.docx fixture not committed",
)


def _work_item(wi_id: int = 99) -> dict:
    return {"id": wi_id, "item_type": "domain_reconciliation"}


# -------------------------------------------------------------------------
# Helpers for building synthetic documents
# -------------------------------------------------------------------------


def _save(doc: Document, tmp_path: Path, name: str = "test.docx") -> Path:
    p = tmp_path / name
    doc.save(str(p))
    return p


def _add_row(table, key: str, value: str) -> None:
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value


def _build_minimal_doc(
    *,
    include_header: bool = True,
    domain_value: str = "Test Domain (TD)",
    domain_code_value: str = "TD",
    include_overview: bool = True,
    overview_text: str = "This is the domain overview.",
    include_personas: bool = False,
    include_processes: bool = False,
) -> Document:
    """Build a minimal valid Domain PRD document."""
    doc = Document()

    # Header table
    if include_header:
        ht = doc.add_table(rows=1, cols=2)
        ht.rows[0].cells[0].text = "Domain"
        ht.rows[0].cells[1].text = domain_value
        _add_row(ht, "Domain Code", domain_code_value)
        _add_row(ht, "Version", "1.0")
        _add_row(ht, "Status", "Draft")

    # Section 1 — Domain Overview
    if include_overview:
        doc.add_heading("1. Domain Overview", level=1)
        doc.add_paragraph(overview_text)

    # Section 2 — Personas (optional)
    if include_personas:
        doc.add_heading("2. Personas", level=1)
        doc.add_paragraph("Test Persona (TST-PER-001)")
        doc.add_paragraph("This persona does test things.")

    # Section 3 — Business Processes (optional)
    if include_processes:
        doc.add_heading("3. Business Processes", level=1)
        doc.add_heading("3.1 Test Process (TD-TEST)", level=2)
        doc.add_paragraph("Test process description.")

    return doc


# =========================================================================
# Fixture tests against CBM-Domain-PRD-Mentoring.docx v1.0
# =========================================================================


class TestMentoringFixture:
    """Tests against the real CBM Domain PRD Mentoring fixture."""

    @pytest.fixture()
    def parsed(self) -> tuple[dict, ParseReport]:
        json_str, report = parse(FIXTURE, _work_item())
        return json.loads(json_str), report

    @skip_fixture
    def test_parse_succeeds_and_roundtrips(self, parsed):
        envelope, _report = parsed
        assert json.dumps(envelope)

    @skip_fixture
    def test_source_metadata(self, parsed):
        envelope, _report = parsed
        meta = envelope["payload"]["source_metadata"]
        assert meta["domain_code"] == "MN"
        assert meta["domain_name"] == "Mentoring"
        assert meta["version"] == "1.0"

    @skip_fixture
    def test_envelope_structure(self, parsed):
        envelope, _report = parsed
        assert envelope["work_item_type"] == "domain_reconciliation"
        assert envelope["output_version"] == "1.0"
        assert envelope["work_item_id"] == 99
        assert envelope["session_type"] == "initial"

    @skip_fixture
    def test_domain_overview_narrative(self, parsed):
        envelope, _report = parsed
        narrative = envelope["payload"]["domain_overview_narrative"]
        assert isinstance(narrative, str)
        assert len(narrative) > 0

    @skip_fixture
    def test_personas(self, parsed):
        envelope, _report = parsed
        personas = envelope["payload"]["personas"]
        assert len(personas) >= 1
        for p in personas:
            assert "name" in p
            assert "consolidated_role" in p
            assert p["consolidated_role"]  # non-empty

    @skip_fixture
    def test_process_summaries(self, parsed):
        envelope, _report = parsed
        summaries = envelope["payload"]["process_summaries"]
        assert len(summaries) >= 1
        for s in summaries:
            assert "process_code" in s
            assert "process_name" in s

    @skip_fixture
    def test_process_requirements(self, parsed):
        """At least one process summary has requirements."""
        envelope, _report = parsed
        summaries = envelope["payload"]["process_summaries"]
        has_reqs = any(len(s["requirements"]) > 0 for s in summaries)
        assert has_reqs

    @skip_fixture
    def test_data_reference(self, parsed):
        envelope, _report = parsed
        data_ref = envelope["payload"]["consolidated_data_reference"]
        assert len(data_ref) >= 1
        for entry in data_ref:
            assert "entity_name" in entry
            assert "deduplicated_fields" in entry

    @skip_fixture
    def test_data_reference_defined_in(self, parsed):
        """Field dicts have the 7th column 'defined_in' key."""
        envelope, _report = parsed
        data_ref = envelope["payload"]["consolidated_data_reference"]
        # Find first entity with fields
        for entry in data_ref:
            if entry["deduplicated_fields"]:
                field = entry["deduplicated_fields"][0]
                assert "defined_in" in field
                break

    @skip_fixture
    def test_decisions(self, parsed):
        envelope, _report = parsed
        decisions = envelope["decisions"]
        assert len(decisions) >= 1
        for d in decisions:
            assert "identifier" in d
            assert "description" in d
            assert "rationale" in d
            assert "made_during" in d
            assert d["status"] == "locked"

    @skip_fixture
    def test_open_issues(self, parsed):
        envelope, _report = parsed
        issues = envelope["open_issues"]
        assert len(issues) >= 1
        for oi in issues:
            assert "identifier" in oi
            assert "description" in oi
            assert "question" in oi
            assert "needs_input_from" in oi
            assert oi["status"] == "open"


# =========================================================================
# Hard-failure synthetic tests
# =========================================================================


class TestHardFailures:
    """Tests for structural parse errors."""

    def test_wrong_work_item_type(self, tmp_path):
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        with pytest.raises(ValueError, match="domain_reconciliation"):
            parse(p, {"id": 1, "item_type": "wrong_type"})

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            parse("/nonexistent/path.docx", _work_item())

    def test_no_header_table(self, tmp_path):
        doc = _build_minimal_doc(include_header=False)
        p = _save(doc, tmp_path)
        with pytest.raises(DomainPrdParseError, match="header"):
            parse(p, _work_item())

    def test_missing_domain_overview(self, tmp_path):
        doc = _build_minimal_doc(include_overview=False)
        p = _save(doc, tmp_path)
        with pytest.raises(DomainPrdParseError, match="Domain Overview"):
            parse(p, _work_item())


# =========================================================================
# Soft-warning synthetic tests
# =========================================================================


class TestSoftWarnings:
    """Tests for soft warnings."""

    def test_domain_name_code_parse_fallback(self, tmp_path):
        """Domain row doesn't match Name (CODE) — falls back to Domain Code row."""
        doc = _build_minimal_doc(
            domain_value="Bad Format",
            domain_code_value="BF",
        )
        p = _save(doc, tmp_path)
        json_str, report = parse(p, _work_item())
        envelope = json.loads(json_str)

        assert envelope["payload"]["source_metadata"]["domain_code"] == "BF"
        assert any(w.category == "domain_name_code_parse" for w in report.warnings)

    def test_missing_personas_section(self, tmp_path):
        """Missing Section 2 → empty personas, no error."""
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        json_str, _report = parse(p, _work_item())
        envelope = json.loads(json_str)

        assert envelope["payload"]["personas"] == []


# =========================================================================
# Structural tests
# =========================================================================


class TestStructure:
    """Tests for envelope structure."""

    def test_envelope_keys(self, tmp_path):
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        json_str, _report = parse(p, _work_item())
        envelope = json.loads(json_str)

        required_keys = {
            "output_version", "work_item_type", "work_item_id",
            "session_type", "payload", "decisions", "open_issues",
        }
        assert required_keys.issubset(envelope.keys())

    def test_payload_keys(self, tmp_path):
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        json_str, _report = parse(p, _work_item())
        envelope = json.loads(json_str)

        payload = envelope["payload"]
        assert "source_metadata" in payload
        assert "domain_overview_narrative" in payload
        assert "personas" in payload
        assert "process_summaries" in payload
        assert "consolidated_data_reference" in payload
