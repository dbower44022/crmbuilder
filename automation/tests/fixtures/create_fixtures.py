"""Script to create minimal CBM-style fixture documents for testing.

Run once to generate the fixture .docx files:
    python automation/tests/fixtures/create_fixtures.py
"""

from pathlib import Path

from docx import Document

FIXTURES = Path(__file__).parent / "cbm_subset"


def create_master_prd() -> None:
    """Create a minimal Master PRD fixture."""
    doc = Document()

    # Header table
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Document Type"
    table.cell(0, 1).text = "Master PRD"
    table.cell(1, 0).text = "Version"
    table.cell(1, 1).text = "1.0"
    table.cell(2, 0).text = "Status"
    table.cell(2, 1).text = "Active"
    table.cell(3, 0).text = "Organization"
    table.cell(3, 1).text = "Cleveland Business Mentors"

    doc.add_heading("1. Organization Overview", level=1)
    doc.add_paragraph(
        "Cleveland Business Mentors (CBM) is a nonprofit providing free "
        "business mentoring to entrepreneurs and small businesses in "
        "Northeast Ohio. CBM connects volunteer mentors with clients "
        "through a structured engagement process."
    )

    doc.add_heading("2. Personas", level=1)
    doc.add_paragraph("MST-PER-001: System Administrator — Manages CRM configuration and user access")
    doc.add_paragraph("MST-PER-002: Executive Member — Strategic oversight and reporting")
    doc.add_paragraph("MST-PER-003: Client Administrator — Manages client intake and engagement lifecycle")
    doc.add_paragraph("MST-PER-011: Mentor — Volunteer mentor providing business guidance")
    doc.add_paragraph("MST-PER-013: Client — Individual or business receiving mentoring")

    doc.add_heading("3. Business Domains", level=1)
    doc.add_paragraph("MN — Mentoring: Core engagement lifecycle from intake through closure")
    doc.add_paragraph("MR — Mentor Recruitment: Recruiting, onboarding, and managing mentor volunteers")

    doc.add_heading("4. Process Inventory", level=1)

    # Process table
    proc_table = doc.add_table(rows=4, cols=3)
    proc_table.cell(0, 0).text = "Code"
    proc_table.cell(0, 1).text = "Name"
    proc_table.cell(0, 2).text = "Domain"
    proc_table.cell(1, 0).text = "MN-INTAKE"
    proc_table.cell(1, 1).text = "Client Intake"
    proc_table.cell(1, 2).text = "MN"
    proc_table.cell(2, 0).text = "MN-MATCH"
    proc_table.cell(2, 1).text = "Mentor Matching"
    proc_table.cell(2, 2).text = "MN"
    proc_table.cell(3, 0).text = "MR-RECRUIT"
    proc_table.cell(3, 1).text = "Mentor Recruitment"
    proc_table.cell(3, 2).text = "MR"

    doc.save(str(FIXTURES / "CBM-Master-PRD.docx"))


def create_entity_inventory() -> None:
    """Create a minimal Entity Inventory fixture."""
    doc = Document()

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Document Type"
    table.cell(0, 1).text = "Entity Inventory"

    doc.add_heading("Entity Inventory", level=1)

    inv_table = doc.add_table(rows=4, cols=4)
    inv_table.cell(0, 0).text = "Business Concept"
    inv_table.cell(0, 1).text = "CRM Entity"
    inv_table.cell(0, 2).text = "Type"
    inv_table.cell(0, 3).text = "Native/Custom"
    inv_table.cell(1, 0).text = "Client Contact"
    inv_table.cell(1, 1).text = "Contact"
    inv_table.cell(1, 2).text = "Person"
    inv_table.cell(1, 3).text = "Native (Person)"
    inv_table.cell(2, 0).text = "Client Organization"
    inv_table.cell(2, 1).text = "Account"
    inv_table.cell(2, 2).text = "Company"
    inv_table.cell(2, 3).text = "Native (Company)"
    inv_table.cell(3, 0).text = "Mentoring Engagement"
    inv_table.cell(3, 1).text = "Engagement"
    inv_table.cell(3, 2).text = "Base"
    inv_table.cell(3, 3).text = "Custom"

    # Decoy detail table (Bug #3 regression test — should be ignored)
    detail = doc.add_table(rows=3, cols=2)
    detail.cell(0, 0).text = "Entity Type"
    detail.cell(0, 1).text = "Person"
    detail.cell(1, 0).text = "Display Name"
    detail.cell(1, 1).text = "Contact"
    detail.cell(2, 0).text = "Activity Stream"
    detail.cell(2, 1).text = "Yes"

    doc.save(str(FIXTURES / "CBM-Entity-Inventory.docx"))


def create_contact_entity_prd() -> None:
    """Create a minimal Contact Entity PRD fixture."""
    doc = Document()

    # Header table
    header = doc.add_table(rows=5, cols=2)
    header.cell(0, 0).text = "Entity"
    header.cell(0, 1).text = "Contact"
    header.cell(1, 0).text = "Entity Type"
    header.cell(1, 1).text = "Person"
    header.cell(2, 0).text = "Native/Custom"
    header.cell(2, 1).text = "Native"
    header.cell(3, 0).text = "Singular Label"
    header.cell(3, 1).text = "Contact"
    header.cell(4, 0).text = "Plural Label"
    header.cell(4, 1).text = "Contacts"

    doc.add_heading("1. Entity Overview", level=1)
    doc.add_paragraph(
        "Contact represents any individual in the CBM system. Uses a "
        "multiEnum contactType discriminator field."
    )

    doc.add_heading("2. Custom Fields", level=1)

    # Field table
    field_table = doc.add_table(rows=5, cols=6)
    field_table.cell(0, 0).text = "Field Name"
    field_table.cell(0, 1).text = "Type"
    field_table.cell(0, 2).text = "Required"
    field_table.cell(0, 3).text = "Values"
    field_table.cell(0, 4).text = "Default"
    field_table.cell(0, 5).text = "ID"

    field_table.cell(1, 0).text = "contactType"
    field_table.cell(1, 1).text = "multiEnum"
    field_table.cell(1, 2).text = "Yes"
    field_table.cell(1, 3).text = "Client | Mentor | Partner | Administrator"
    field_table.cell(1, 4).text = "—"
    field_table.cell(1, 5).text = "MN-INTAKE-DAT-001"

    field_table.cell(2, 0).text = "preferredName"
    field_table.cell(2, 1).text = "varchar"
    field_table.cell(2, 2).text = "No"
    field_table.cell(2, 3).text = ""
    field_table.cell(2, 4).text = "—"
    field_table.cell(2, 5).text = "MN-INTAKE-DAT-002"

    field_table.cell(3, 0).text = "linkedInProfile"
    field_table.cell(3, 1).text = "url"
    field_table.cell(3, 2).text = "No"
    field_table.cell(3, 3).text = ""
    field_table.cell(3, 4).text = "—"
    field_table.cell(3, 5).text = "MN-INTAKE-DAT-003"

    field_table.cell(4, 0).text = "mentorStatus"
    field_table.cell(4, 1).text = "enum"
    field_table.cell(4, 2).text = "No"
    field_table.cell(4, 3).text = "Active | Inactive | On Hold"
    field_table.cell(4, 4).text = "—"
    field_table.cell(4, 5).text = "MR-RECRUIT-DAT-001"

    doc.add_heading("3. Relationships", level=1)
    doc.add_paragraph("Contact has relationships to Account and Engagement entities.")

    doc.save(str(FIXTURES / "entities" / "Contact-Entity-PRD.docx"))


def create_intake_process() -> None:
    """Create a minimal MN-INTAKE process document fixture."""
    doc = Document()

    # Header table
    header = doc.add_table(rows=4, cols=2)
    header.cell(0, 0).text = "Process Code"
    header.cell(0, 1).text = "MN-INTAKE"
    header.cell(1, 0).text = "Domain"
    header.cell(1, 1).text = "MN"
    header.cell(2, 0).text = "Process Name"
    header.cell(2, 1).text = "Client Intake"
    header.cell(3, 0).text = "Version"
    header.cell(3, 1).text = "1.0"

    doc.add_heading("1. Process Purpose", level=1)
    doc.add_paragraph(
        "The Client Intake process captures prospective client information, "
        "assesses eligibility for mentoring services, and routes approved "
        "clients to the mentor matching process."
    )

    doc.add_heading("2. Process Triggers", level=1)
    doc.add_paragraph(
        "Initiated when a prospective client submits a mentoring request form. "
        "No preconditions required — this is the first process in the engagement lifecycle."
    )

    doc.add_heading("3. Personas Involved", level=1)
    doc.add_paragraph("MST-PER-013: Client — submits the mentoring request")
    doc.add_paragraph("MST-PER-003: Client Administrator — reviews and approves intake")

    doc.add_heading("4. Process Workflow", level=1)
    # Use List Paragraph style to match real CBM documents (Bug #2)
    doc.add_paragraph("Prospective client submits Phase 1 Mentoring Request form", style="List Paragraph")
    doc.add_paragraph("System creates Account, Contact, and Engagement records", style="List Paragraph")
    doc.add_paragraph("Client Administrator receives notification of new submission", style="List Paragraph")
    doc.add_paragraph("Administrator reviews submission for eligibility", style="List Paragraph")
    doc.add_paragraph("If eligible, advances engagement to MN-MATCH process", style="List Paragraph")

    doc.add_heading("5. Process Completion", level=1)
    doc.add_paragraph(
        "The process completes when the Client Administrator either approves "
        "or declines the client's application."
    )

    doc.add_heading("6. System Requirements", level=1)
    # Use table format to match real CBM documents (Bug #1)
    req_table = doc.add_table(rows=4, cols=2)
    req_table.cell(0, 0).text = "ID"
    req_table.cell(0, 1).text = "Requirement"
    req_table.cell(1, 0).text = "MN-INTAKE-REQ-001"
    req_table.cell(1, 1).text = "System shall accept mentoring request form submissions"
    req_table.cell(2, 0).text = "MN-INTAKE-REQ-002"
    req_table.cell(2, 1).text = "System shall auto-create linked Account, Contact, and Engagement records"
    req_table.cell(3, 0).text = "MN-INTAKE-REQ-003"
    req_table.cell(3, 1).text = "System shall notify Client Administrator of new submissions"

    doc.add_heading("7. Process Data", level=1)
    doc.add_paragraph("Entities referenced: Contact, Account, Engagement")

    doc.save(str(FIXTURES / "MN" / "MN-INTAKE.docx"))


def create_domain_prd() -> None:
    """Create a minimal Domain PRD (Mentoring) fixture."""
    doc = Document()

    header = doc.add_table(rows=3, cols=2)
    header.cell(0, 0).text = "Document Type"
    header.cell(0, 1).text = "Domain PRD"
    header.cell(1, 0).text = "Domain Code"
    header.cell(1, 1).text = "MN"
    header.cell(2, 0).text = "Domain"
    header.cell(2, 1).text = "Mentoring"

    doc.add_heading("1. Domain Overview", level=1)
    doc.add_paragraph(
        "The Mentoring domain covers the full engagement lifecycle from "
        "initial client intake through engagement closure. It includes "
        "five core processes: Intake, Matching, Engage, Inactive Management, "
        "and Closure."
    )

    doc.add_heading("2. Design Decisions", level=1)
    doc.add_paragraph(
        "MN-DEC-001: Engagement records use auto-generated names combining "
        "client name, mentor name, and start year."
    )
    doc.add_paragraph(
        "MN-DEC-002: Session records are never deleted; cancelled sessions "
        "retain history for institutional memory."
    )

    doc.save(str(FIXTURES / "MN" / "CBM-Domain-PRD-Mentoring.docx"))


def create_heading2_workflow_process() -> None:
    """Create a process doc with Heading 2 subsections in the workflow section.

    Regression test for Bug #5: the step extractor used to break on any
    Heading (including Heading 2), losing all steps in documents that
    organize workflow sub-flows under Heading 2 subsections.
    """
    doc = Document()

    header = doc.add_table(rows=4, cols=2)
    header.cell(0, 0).text = "Process Code"
    header.cell(0, 1).text = "TD-SUBSECT"
    header.cell(1, 0).text = "Domain"
    header.cell(1, 1).text = "TD"
    header.cell(2, 0).text = "Process Name"
    header.cell(2, 1).text = "Test Subsection Workflow"
    header.cell(3, 0).text = "Version"
    header.cell(3, 1).text = "1.0"

    doc.add_heading("1. Process Purpose", level=1)
    doc.add_paragraph("Test process for Heading 2 subsection handling.")

    doc.add_heading("4. Process Workflow", level=1)
    doc.add_heading("4.1 First Subsection", level=2)
    doc.add_paragraph("1. Step one of first subsection")
    doc.add_paragraph("2. Step two of first subsection")
    doc.add_paragraph("3. Step three of first subsection")
    doc.add_heading("4.2 Second Subsection", level=2)
    doc.add_paragraph("1. Step four from second subsection")
    doc.add_paragraph("2. Step five from second subsection")
    doc.add_heading("5. Process Completion", level=1)
    doc.add_paragraph("Process is complete when all sub-flows finish.")

    (FIXTURES / "TD").mkdir(exist_ok=True)
    doc.save(str(FIXTURES / "TD" / "TD-SUBSECT.docx"))


def create_headerless_req_process() -> None:
    """Create a process doc with a requirements table that has no header row.

    Regression test for Bug #6: the requirements-table detector used to
    require a literal ['ID', 'Requirement'] header row, missing tables
    where the first row is already data.
    """
    doc = Document()

    header = doc.add_table(rows=4, cols=2)
    header.cell(0, 0).text = "Process Code"
    header.cell(0, 1).text = "TD-NOHDR"
    header.cell(1, 0).text = "Domain"
    header.cell(1, 1).text = "TD"
    header.cell(2, 0).text = "Process Name"
    header.cell(2, 1).text = "Test Headerless Requirements"
    header.cell(3, 0).text = "Version"
    header.cell(3, 1).text = "1.0"

    doc.add_heading("1. Process Purpose", level=1)
    doc.add_paragraph("Test process for headerless requirements table.")

    doc.add_heading("4. Process Workflow", level=1)
    doc.add_paragraph("Single step process.", style="List Paragraph")

    doc.add_heading("6. System Requirements", level=1)
    # Requirements table WITHOUT a header row — first row is data
    req_table = doc.add_table(rows=2, cols=2)
    req_table.cell(0, 0).text = "TD-NOHDR-REQ-001"
    req_table.cell(0, 1).text = "First requirement description"
    req_table.cell(1, 0).text = "TD-NOHDR-REQ-002"
    req_table.cell(1, 1).text = "Second requirement description"

    # Decoy field table — should NOT be detected as requirements
    doc.add_heading("7. Process Data", level=1)
    dat_table = doc.add_table(rows=2, cols=6)
    dat_table.cell(0, 0).text = "ID"
    dat_table.cell(0, 1).text = "Field Name"
    dat_table.cell(0, 2).text = "Type"
    dat_table.cell(0, 3).text = "Required"
    dat_table.cell(0, 4).text = "Values"
    dat_table.cell(0, 5).text = "Description"
    dat_table.cell(1, 0).text = "TD-NOHDR-DAT-001"
    dat_table.cell(1, 1).text = "testField"
    dat_table.cell(1, 2).text = "varchar"
    dat_table.cell(1, 3).text = "No"
    dat_table.cell(1, 4).text = ""
    dat_table.cell(1, 5).text = "A test field"

    (FIXTURES / "TD").mkdir(exist_ok=True)
    doc.save(str(FIXTURES / "TD" / "TD-NOHDR.docx"))


if __name__ == "__main__":
    FIXTURES.mkdir(parents=True, exist_ok=True)
    (FIXTURES / "entities").mkdir(exist_ok=True)
    (FIXTURES / "MN").mkdir(exist_ok=True)
    (FIXTURES / "TD").mkdir(exist_ok=True)

    create_master_prd()
    create_entity_inventory()
    create_contact_entity_prd()
    create_intake_process()
    create_domain_prd()
    create_heading2_workflow_process()
    create_headerless_req_process()
    print(f"Fixtures created in {FIXTURES}")
