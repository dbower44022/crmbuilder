"""Tests for automation.importer.parsers.master_prd_docx."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from automation.importer.parsers import MasterPrdParseError, ParseReport
from automation.importer.parsers.master_prd_docx import parse

FIXTURE = Path(__file__).parent / "fixtures" / "cbm-master-prd-v2.6.docx"

# Skip real-document tests if fixture not present
pytestmark_real = pytest.mark.skipif(
    not FIXTURE.exists(),
    reason="cbm-master-prd-v2.6.docx fixture not committed",
)


def _work_item(wi_id: int = 99) -> dict:
    return {"id": wi_id, "item_type": "master_prd"}


# -------------------------------------------------------------------------
# 7.1 — Real-document tests
# -------------------------------------------------------------------------


@pytestmark_real
class TestRealDocument:

    def test_parse_real_document_returns_envelope_and_report(self):
        envelope_json, report = parse(FIXTURE, _work_item())
        assert isinstance(envelope_json, str)
        assert isinstance(report, ParseReport)
        envelope = json.loads(envelope_json)
        assert "output_version" in envelope
        assert "payload" in envelope

    def test_parse_real_document_envelope_metadata(self):
        envelope_json, _ = parse(FIXTURE, _work_item(42), session_type="revision")
        envelope = json.loads(envelope_json)
        assert envelope["output_version"] == "1.0"
        assert envelope["work_item_type"] == "master_prd"
        assert envelope["work_item_id"] == 42
        assert envelope["session_type"] == "revision"

    def test_parse_real_document_persona_count(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        payload = json.loads(envelope_json)["payload"]
        personas = payload["personas"]
        assert len(personas) == 13
        codes = [p["identifier"] for p in personas]
        for i in range(1, 14):
            assert f"MST-PER-{i:03d}" in codes

    def test_parse_real_document_domain_count(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        payload = json.loads(envelope_json)["payload"]
        non_svc = [d for d in payload["domains"] if not d["is_service"]]
        assert len(non_svc) == 4
        codes = {d["code"] for d in non_svc}
        assert codes == {"MN", "MR", "CR", "FU"}
        names = {d["code"]: d["name"] for d in non_svc}
        assert names == {
            "MN": "Mentoring",
            "MR": "Mentor Recruitment",
            "CR": "Client Recruiting",
            "FU": "Fundraising",
        }

    def test_parse_real_document_subdomain_structure(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        payload = json.loads(envelope_json)["payload"]
        by_code = {d["code"]: d for d in payload["domains"] if not d["is_service"]}
        cr = by_code["CR"]
        sd_codes = [sd["code"] for sd in cr["sub_domains"]]
        assert sd_codes == ["CR-PARTNER", "CR-MARKETING", "CR-EVENTS", "CR-REACTIVATE"]
        for code in ("MN", "MR", "FU"):
            assert by_code[code]["sub_domains"] == []

    def test_parse_real_document_service_count(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        payload = json.loads(envelope_json)["payload"]
        services = [d for d in payload["domains"] if d["is_service"]]
        assert len(services) == 4
        codes = {s["code"] for s in services}
        assert codes == {"NOTES", "EMAIL", "CALENDAR", "SURVEY"}
        for s in services:
            assert s["name"].endswith(" Service") or s["name"].endswith("Service")
        sort_orders = sorted(s["sort_order"] for s in services)
        assert sort_orders == [101, 102, 103, 104]

    def test_parse_real_document_process_count(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        payload = json.loads(envelope_json)["payload"]
        processes = payload["processes"]
        assert len(processes) == 15
        sort_orders = [p["sort_order"] for p in processes]
        assert sort_orders == list(range(1, 16))

    def test_parse_real_document_process_tiers(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        payload = json.loads(envelope_json)["payload"]
        for proc in payload["processes"]:
            assert proc["tier"] in {"core", "important", "enhancement"}, (
                f"Process {proc['code']} has tier={proc['tier']!r}"
            )

    def test_parse_real_document_organization_overview_format(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        overview = json.loads(envelope_json)["payload"]["organization_overview"]
        assert "## Mission and Context" in overview
        assert "## Operating Model" in overview
        assert "## Why a CRM is Needed" in overview

    def test_parse_real_document_no_warnings_on_v26(self):
        _, report = parse(FIXTURE, _work_item())
        assert report.warnings == [], (
            f"Expected zero warnings, got {len(report.warnings)}: "
            + "; ".join(f"[{w.category}] {w.location}: {w.message}" for w in report.warnings)
        )


# -------------------------------------------------------------------------
# 7.2 — Hard-failure tests (synthetic documents)
# -------------------------------------------------------------------------


def _minimal_doc_with_sections(*heading1_texts: str) -> Document:
    """Build a minimal Document with the given Heading 1 paragraphs."""
    doc = Document()
    for text in heading1_texts:
        doc.add_heading(text, level=1)
        doc.add_paragraph("Placeholder content.")
    return doc


def _save(doc: Document, tmp_path: Path, name: str = "test.docx") -> Path:
    p = tmp_path / name
    doc.save(str(p))
    return p


class TestHardFailures:

    def test_raises_on_missing_personas_section(self, tmp_path):
        doc = _minimal_doc_with_sections(
            "1. Organization Overview", "3. Key Business Domains",
        )
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="Personas"):
            parse(path, _work_item())

    def test_raises_on_missing_domains_section(self, tmp_path):
        doc = _minimal_doc_with_sections(
            "1. Organization Overview", "2. Personas",
        )
        # Add a persona so we pass the zero-personas check
        doc.add_heading("MST-PER-001 — Test Persona", level=3)
        doc.add_paragraph("A test persona.")
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="Key Business Domains"):
            parse(path, _work_item())

    def test_raises_on_missing_table2(self, tmp_path):
        doc = Document()
        doc.add_heading("2. Personas", level=1)
        doc.add_heading("MST-PER-001 — Test", level=3)
        doc.add_paragraph("Desc.")
        doc.add_heading("3. Key Business Domains", level=1)
        doc.add_heading("3.1 MST-DOM-001 — Testing (TST)", level=2)
        h3 = doc.add_heading("Domain Purpose", level=3)  # noqa: F841
        doc.add_paragraph("Test domain purpose.")
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="table not found"):
            parse(path, _work_item())

    def test_raises_on_table2_wrong_columns(self, tmp_path):
        doc = Document()
        doc.add_heading("2. Personas", level=1)
        doc.add_heading("MST-PER-001 — Test", level=3)
        doc.add_paragraph("Desc.")
        doc.add_heading("3. Key Business Domains", level=1)
        doc.add_heading("3.1 MST-DOM-001 — Testing (TST)", level=2)
        doc.add_heading("Domain Purpose", level=3)
        doc.add_paragraph("Test domain purpose.")
        # Table with wrong columns
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Code"
        table.cell(0, 1).text = "Name"
        table.cell(0, 2).text = "Domain"
        table.cell(1, 0).text = "TST-PROC"
        table.cell(1, 1).text = "Test Process"
        table.cell(1, 2).text = "Testing (TST)"
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="wrong column structure"):
            parse(path, _work_item())

    def test_raises_on_zero_personas_parsed(self, tmp_path):
        doc = Document()
        doc.add_heading("2. Personas", level=1)
        doc.add_paragraph("No persona headings here.")
        doc.add_heading("3. Key Business Domains", level=1)
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="Zero personas"):
            parse(path, _work_item())

    def test_raises_on_zero_domains_parsed(self, tmp_path):
        doc = Document()
        doc.add_heading("2. Personas", level=1)
        doc.add_heading("MST-PER-001 — Test", level=3)
        doc.add_paragraph("Desc.")
        doc.add_heading("3. Key Business Domains", level=1)
        doc.add_paragraph("No domain headings here.")
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="Zero top-level domains"):
            parse(path, _work_item())

    def test_raises_on_zero_processes_parsed(self, tmp_path):
        doc = Document()
        doc.add_heading("2. Personas", level=1)
        doc.add_heading("MST-PER-001 — Test", level=3)
        doc.add_paragraph("Desc.")
        doc.add_heading("3. Key Business Domains", level=1)
        doc.add_heading("3.1 MST-DOM-001 — Testing (TST)", level=2)
        doc.add_heading("Domain Purpose", level=3)
        doc.add_paragraph("Test domain purpose.")
        # Table 2 with header only, no data rows
        table = doc.add_table(rows=1, cols=4)
        table.cell(0, 0).text = "Code"
        table.cell(0, 1).text = "Process / Sub-Domain"
        table.cell(0, 2).text = "Domain"
        table.cell(0, 3).text = "Tier"
        path = _save(doc, tmp_path)
        with pytest.raises(MasterPrdParseError, match="Zero processes"):
            parse(path, _work_item())

    def test_raises_filenotfound(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            parse(tmp_path / "nonexistent.docx", _work_item())

    def test_raises_valueerror_on_wrong_work_item_type(self, tmp_path):
        # File doesn't need to exist — ValueError fires before reading
        path = tmp_path / "dummy.docx"
        path.write_bytes(b"")
        with pytest.raises(ValueError, match="master_prd"):
            parse(path, {"id": 1, "item_type": "entity_prd"})


# -------------------------------------------------------------------------
# 7.3 — Soft-warning tests
# -------------------------------------------------------------------------


def _build_minimal_valid_doc() -> Document:
    """Build a minimal but structurally valid document for warning tests."""
    doc = Document()
    doc.add_heading("1. Organization Overview", level=1)
    doc.add_heading("1.1 Mission", level=2)
    doc.add_paragraph("Mission text.")

    doc.add_heading("2. Personas", level=1)
    doc.add_heading("MST-PER-001 — Test Persona", level=3)
    doc.add_paragraph("A persona description.")

    doc.add_heading("3. Key Business Domains", level=1)
    return doc


def _add_domain_with_table(doc: Document, code: str, name: str, processes: list[tuple[str, str, str]],
                            sub_domains: list[tuple[str, str]] | None = None,
                            include_purpose: bool = True) -> None:
    """Add a domain section and Table 2 rows."""
    doc.add_heading(f"3.1 MST-DOM-001 — {name} ({code})", level=2)
    if include_purpose:
        doc.add_heading("Domain Purpose", level=3)
        doc.add_paragraph(f"{name} domain purpose.")

    table = doc.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "Code"
    table.cell(0, 1).text = "Process / Sub-Domain"
    table.cell(0, 2).text = "Domain"
    table.cell(0, 3).text = "Tier"

    if sub_domains:
        for sd_code, sd_name in sub_domains:
            row = table.add_row()
            row.cells[0].text = sd_code
            row.cells[1].text = f"{sd_name} (Sub-Domain)"
            row.cells[2].text = f"{name} ({code})"
            row.cells[3].text = "Core"

    for proc_code, proc_name, tier in processes:
        row = table.add_row()
        row.cells[0].text = proc_code
        row.cells[1].text = proc_name
        row.cells[2].text = f"{name} ({code})"
        row.cells[3].text = tier


class TestSoftWarnings:

    def test_warns_on_missing_domain_purpose(self, tmp_path):
        doc = _build_minimal_valid_doc()
        _add_domain_with_table(
            doc, "MN", "Mentoring",
            [("MN-PROC", "Test Process", "Core")],
            include_purpose=False,
        )
        path = _save(doc, tmp_path)
        _, report = parse(path, _work_item())
        cats = [(w.category, w.location) for w in report.warnings]
        assert any(
            c == "missing_description" and "MN" in loc
            for c, loc in cats
        ), f"Expected missing_description for MN, got: {cats}"

    def test_warns_on_orphan_prose_subdomain(self, tmp_path):
        doc = _build_minimal_valid_doc()
        doc.add_heading("3.1 MST-DOM-001 — Recruiting (CR)", level=2)
        doc.add_heading("Domain Purpose", level=3)
        doc.add_paragraph("Recruiting purpose.")
        doc.add_heading("Sub-Domains", level=3)
        doc.add_paragraph("CR-PARTNER — Partner Management")
        doc.add_paragraph("Partner desc.")
        doc.add_paragraph("CR-GHOST — Ghost Sub-Domain")
        doc.add_paragraph("Ghost desc.")

        table = doc.add_table(rows=1, cols=4)
        table.cell(0, 0).text = "Code"
        table.cell(0, 1).text = "Process / Sub-Domain"
        table.cell(0, 2).text = "Domain"
        table.cell(0, 3).text = "Tier"
        # Only CR-PARTNER in Table 2 as sub-domain, not CR-GHOST
        row = table.add_row()
        row.cells[0].text = "CR-PARTNER"
        row.cells[1].text = "Partner Management (Sub-Domain)"
        row.cells[2].text = "Recruiting (CR)"
        row.cells[3].text = "Core"
        # Need at least one process
        row2 = table.add_row()
        row2.cells[0].text = "CR-PROC"
        row2.cells[1].text = "Test Process"
        row2.cells[2].text = "Recruiting (CR)"
        row2.cells[3].text = "Core"

        path = _save(doc, tmp_path)
        _, report = parse(path, _work_item())
        orphan_warnings = [w for w in report.warnings if w.category == "orphan_prose"]
        assert any("CR-GHOST" in w.location for w in orphan_warnings), (
            f"Expected orphan_prose for CR-GHOST, got: {orphan_warnings}"
        )

    def test_warns_on_bad_tier_value(self, tmp_path):
        doc = _build_minimal_valid_doc()
        doc.add_heading("3.1 MST-DOM-001 — Testing (TST)", level=2)
        doc.add_heading("Domain Purpose", level=3)
        doc.add_paragraph("Testing purpose.")

        table = doc.add_table(rows=2, cols=4)
        table.cell(0, 0).text = "Code"
        table.cell(0, 1).text = "Process / Sub-Domain"
        table.cell(0, 2).text = "Domain"
        table.cell(0, 3).text = "Tier"
        table.cell(1, 0).text = "TST-PROC"
        table.cell(1, 1).text = "Test Process"
        table.cell(1, 2).text = "Testing (TST)"
        table.cell(1, 3).text = "Critical"

        path = _save(doc, tmp_path)
        _, report = parse(path, _work_item())
        bad_tier_warnings = [w for w in report.warnings if w.category == "bad_tier"]
        assert len(bad_tier_warnings) == 1
        assert "Critical" in bad_tier_warnings[0].message

        envelope = json.loads(parse(path, _work_item())[0])
        proc = envelope["payload"]["processes"][0]
        assert proc["tier"] is None


# -------------------------------------------------------------------------
# 7.4 — Round-trip tests
# -------------------------------------------------------------------------


@pytestmark_real
class TestRoundTrip:

    def test_envelope_passes_path_b_validation(self):
        envelope_json, _ = parse(FIXTURE, _work_item())
        from automation.importer.parser import parse_and_validate
        # Should not raise
        parse_and_validate(
            envelope_json,
            expected_item_type="master_prd",
            expected_work_item_id=99,
            expected_session_type="initial",
        )

    def test_envelope_round_trip_through_mapper(self, tmp_path):
        envelope_json, _ = parse(FIXTURE, _work_item(1))
        envelope = json.loads(envelope_json)
        payload = envelope["payload"]

        from automation.db.migrations import run_client_migrations
        conn = run_client_migrations(str(tmp_path / "client.db"))
        try:
            from automation.importer.mappers.master_prd import map_payload
            work_item = {
                "id": 1, "item_type": "master_prd", "status": "in_progress",
                "domain_id": None, "entity_id": None, "process_id": None,
            }
            batch = map_payload(conn, work_item, payload, "initial", 1)

            # Count records by table
            by_table = batch.records_by_table()

            # 13 Persona creates
            persona_recs = by_table.get("Persona", [])
            assert len(persona_recs) == 13, f"Expected 13 Persona, got {len(persona_recs)}"

            # 8 Domain creates (MN, MR, CR, FU + NOTES, EMAIL, CALENDAR, SURVEY)
            domain_recs = [r for r in by_table.get("Domain", []) if r.values.get("code") in {
                "MN", "MR", "CR", "FU", "NOTES", "EMAIL", "CALENDAR", "SURVEY",
            }]
            assert len(domain_recs) == 8, f"Expected 8 top-level Domain, got {len(domain_recs)}"

            # 4 sub-domain Domain creates with parent_domain_id ref to batch:domain:CR
            sd_recs = [
                r for r in by_table.get("Domain", [])
                if r.intra_batch_refs.get("parent_domain_id") == "batch:domain:CR"
            ]
            assert len(sd_recs) == 4, f"Expected 4 sub-domain refs to CR, got {len(sd_recs)}"

            # 15 Process creates with domain refs
            proc_recs = by_table.get("Process", [])
            assert len(proc_recs) == 15, f"Expected 15 Process, got {len(proc_recs)}"
            for pr in proc_recs:
                has_domain = (
                    "domain_id" in pr.values
                    or "domain_id" in pr.intra_batch_refs
                )
                assert has_domain, f"Process {pr.values.get('code')} missing domain ref"
        finally:
            conn.close()


# -------------------------------------------------------------------------
# 7.5 — Parsed counts test
# -------------------------------------------------------------------------


@pytestmark_real
class TestParsedCounts:

    def test_parsed_counts_keys(self):
        _, report = parse(FIXTURE, _work_item())
        expected_keys = {"persona", "domain", "sub_domain", "service", "process"}
        assert set(report.parsed_counts.keys()) == expected_keys

    def test_parsed_counts_values(self):
        _, report = parse(FIXTURE, _work_item())
        assert report.parsed_counts == {
            "persona": 13,
            "domain": 4,
            "sub_domain": 4,
            "service": 4,
            "process": 15,
        }
