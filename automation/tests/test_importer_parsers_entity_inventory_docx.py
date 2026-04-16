"""Tests for automation.importer.parsers.entity_inventory_docx."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from automation.importer.parsers import EntityInventoryParseError, ParseReport
from automation.importer.parsers.entity_inventory_docx import parse

FIXTURE = Path(__file__).parent / "fixtures" / "cbm-entity-inventory-v1.4.docx"

skip_fixture = pytest.mark.skipif(
    not FIXTURE.exists(),
    reason="cbm-entity-inventory-v1.4.docx fixture not committed",
)


def _work_item(wi_id: int = 99) -> dict:
    return {"id": wi_id, "item_type": "business_object_discovery"}


# -------------------------------------------------------------------------
# Helpers for building synthetic documents
# -------------------------------------------------------------------------


def _save(doc: Document, tmp_path: Path, name: str = "test.docx") -> Path:
    p = tmp_path / name
    doc.save(str(p))
    return p


def _add_row(table, key: str, value: str) -> None:
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value


def _build_minimal_doc(
    *,
    include_header: bool = True,
    include_entity_map: bool = True,
    entity_map_rows: list[list[str]] | None = None,
    include_detail_cards: bool = False,
    detail_cards: list[dict] | None = None,
    include_open_issues: bool = False,
) -> Document:
    """Build a minimal valid Entity Inventory document."""
    doc = Document()

    # Header table
    if include_header:
        ht = doc.add_table(rows=1, cols=2)
        ht.rows[0].cells[0].text = "Document Type"
        ht.rows[0].cells[1].text = "Entity Inventory"
        _add_row(ht, "Version", "1.0")
        _add_row(ht, "Status", "Draft")

    # Entity Map table
    if include_entity_map:
        default_rows = [
            ["Test Entity", "TestEntity", "Custom", "Base", "—", "—", "MN"],
        ]
        rows = entity_map_rows if entity_map_rows is not None else default_rows
        em = doc.add_table(rows=1, cols=7)
        headers = [
            "PRD Entity Name", "CRM Entity", "Native / Custom",
            "Entity Type", "Discriminator", "Disc. Value", "Domain(s)",
        ]
        for i, h in enumerate(headers):
            em.rows[0].cells[i].text = h
        for row_data in rows:
            row = em.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = val

    # Detail cards
    if include_detail_cards and detail_cards:
        for card in detail_cards:
            dc = doc.add_table(rows=1, cols=2)
            dc.rows[0].cells[0].text = "Entity Type"
            dc.rows[0].cells[1].text = card.get("entity_type_val", "Base")
            _add_row(dc, "Display Label (Singular)", card.get("singular", "Test"))
            _add_row(dc, "Display Label (Plural)", card.get("plural", "Tests"))
            _add_row(dc, "Owning Domain", card.get("owning_domain", "Test (MN)"))
            _add_row(dc, "Activity Stream", card.get("activity_stream", "Yes"))

    # Open Issues table
    if include_open_issues:
        oi = doc.add_table(rows=1, cols=3)
        oi.rows[0].cells[0].text = "ID"
        oi.rows[0].cells[1].text = "Issue"
        oi.rows[0].cells[2].text = "Resolution Path"
        row = oi.add_row()
        row.cells[0].text = "OI-001"
        row.cells[1].text = "Test issue"
        row.cells[2].text = "Test resolution"

    return doc


# =========================================================================
# Fixture tests against CBM-Entity-Inventory.docx v1.4
# =========================================================================


class TestCBMFixture:
    """Tests against the real CBM Entity Inventory fixture."""

    @pytest.fixture()
    def parsed(self) -> tuple[dict, ParseReport]:
        json_str, report = parse(FIXTURE, _work_item())
        return json.loads(json_str), report

    @skip_fixture
    def test_parse_succeeds_and_roundtrips(self, parsed):
        envelope, _report = parsed
        # Round-trip through JSON
        assert json.dumps(envelope)

    @skip_fixture
    def test_source_metadata_version(self, parsed):
        envelope, _report = parsed
        meta = envelope["payload"]["source_metadata"]
        assert meta["version"] == "1.4"
        assert "document_type" in meta

    @skip_fixture
    def test_envelope_structure(self, parsed):
        envelope, _report = parsed
        assert envelope["work_item_type"] == "business_object_discovery"
        assert envelope["output_version"] == "1.0"
        assert envelope["work_item_id"] == 99
        assert envelope["session_type"] == "initial"

    @skip_fixture
    def test_business_objects_count(self, parsed):
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        # Table 1 has 28 data rows (rows 1-28, row 0 is header)
        assert len(bos) == 28

    @skip_fixture
    def test_entity_classified_bos_have_required_fields(self, parsed):
        envelope, _report = parsed
        for bo in envelope["payload"]["business_objects"]:
            if bo["status"] == "classified":
                assert "entity_name" in bo
                assert "entity_type" in bo
                assert "is_native" in bo

    @skip_fixture
    def test_native_entities_present(self, parsed):
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        native_bos = [bo for bo in bos if bo.get("is_native") is True]
        assert len(native_bos) >= 1
        native_names = {bo["entity_name"] for bo in native_bos}
        assert "Contact" in native_names
        assert "Account" in native_names

    @skip_fixture
    def test_person_entity_type(self, parsed):
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        person_bos = [bo for bo in bos if bo.get("entity_type") == "Person"]
        assert len(person_bos) >= 1

    @skip_fixture
    def test_detail_card_enrichment_labels(self, parsed):
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        # Engagement has a detail card
        engagement_bos = [bo for bo in bos if bo.get("entity_name") == "Engagement"]
        assert len(engagement_bos) >= 1
        for bo in engagement_bos:
            assert bo.get("singular_label") == "Engagement"
            assert bo.get("plural_label") == "Engagements"

    @skip_fixture
    def test_bug7_owning_domain_in_source_domains(self, parsed):
        """Bug 7 fix: detail cards set source_domains[0] to Owning Domain code."""
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        # Engagement: Owning Domain = Mentoring (MN)
        engagement_bos = [bo for bo in bos if bo.get("entity_name") == "Engagement"]
        for bo in engagement_bos:
            assert bo["source_domains"][0] == "MN"
        # Dues: Owning Domain = Mentor Recruitment (MR)
        dues_bos = [bo for bo in bos if bo.get("entity_name") == "Dues"]
        for bo in dues_bos:
            assert bo["source_domains"][0] == "MR"

    @skip_fixture
    def test_discriminator_in_description(self, parsed):
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        client_contact = next(bo for bo in bos if bo["name"] == "Client Contact")
        assert "Discriminator:" in client_contact["description"]
        assert "contactType" in client_contact["description"]
        assert "Client" in client_contact["description"]

    @skip_fixture
    def test_multiple_domains(self, parsed):
        envelope, _report = parsed
        bos = envelope["payload"]["business_objects"]
        # Client Contact has domains "MN, CR"
        client_contact = next(bo for bo in bos if bo["name"] == "Client Contact")
        assert len(client_contact["source_domains"]) >= 2

    @skip_fixture
    def test_open_issues(self, parsed):
        envelope, _report = parsed
        issues = envelope["open_issues"]
        assert len(issues) >= 1
        for issue in issues:
            assert "identifier" in issue
            assert "description" in issue
            assert "resolution_path" in issue
            assert issue["status"] == "open"

    @skip_fixture
    def test_decisions_empty(self, parsed):
        envelope, _report = parsed
        assert envelope["decisions"] == []

    @skip_fixture
    def test_parsed_counts(self, parsed):
        _envelope, report = parsed
        assert "business_objects" in report.parsed_counts
        assert "detail_cards" in report.parsed_counts
        assert "open_issues" in report.parsed_counts


# =========================================================================
# Hard-failure synthetic tests
# =========================================================================


class TestHardFailures:
    """Tests for structural parse errors."""

    def test_wrong_work_item_type(self, tmp_path):
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        with pytest.raises(ValueError, match="business_object_discovery"):
            parse(p, {"id": 1, "item_type": "wrong_type"})

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            parse("/nonexistent/path.docx", _work_item())

    def test_no_entity_map_table(self, tmp_path):
        doc = _build_minimal_doc(include_entity_map=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityInventoryParseError, match="Entity Map"):
            parse(p, _work_item())


# =========================================================================
# Soft-warning synthetic tests
# =========================================================================


class TestSoftWarnings:
    """Tests for soft warnings."""

    def test_unclassified_bo_warning(self, tmp_path):
        doc = _build_minimal_doc(
            entity_map_rows=[
                ["Orphan Concept", "", "Custom", "Base", "—", "—", "MN"],
            ],
        )
        p = _save(doc, tmp_path)
        json_str, report = parse(p, _work_item())
        envelope = json.loads(json_str)

        bos = envelope["payload"]["business_objects"]
        assert bos[0]["status"] == "unclassified"
        assert any(w.category == "unclassified_bo" for w in report.warnings)

    def test_unmatched_detail_card_warning(self, tmp_path):
        doc = _build_minimal_doc(
            include_detail_cards=True,
            detail_cards=[{
                "entity_type_val": "Base",
                "singular": "NoMatch",
                "plural": "NoMatches",
                "owning_domain": "Test (MN)",
                "activity_stream": "Yes",
            }],
        )
        p = _save(doc, tmp_path)
        _json_str, report = parse(p, _work_item())
        assert any(w.category == "unmatched_detail_card" for w in report.warnings)

    def test_owning_domain_parse_failure_warning(self, tmp_path):
        doc = _build_minimal_doc(
            entity_map_rows=[
                ["Test Ent", "TestEntity", "Custom", "Base", "—", "—", "MN"],
            ],
            include_detail_cards=True,
            detail_cards=[{
                "entity_type_val": "Base",
                "singular": "TestEntity",
                "plural": "TestEntities",
                "owning_domain": "Bad Format No Parens",
                "activity_stream": "No",
            }],
        )
        p = _save(doc, tmp_path)
        _json_str, report = parse(p, _work_item())
        assert any(
            w.category == "owning_domain_parse_failure" for w in report.warnings
        )


# =========================================================================
# Structural tests
# =========================================================================


class TestStructure:
    """Tests for envelope structure."""

    def test_envelope_keys(self, tmp_path):
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        json_str, _report = parse(p, _work_item())
        envelope = json.loads(json_str)

        required_keys = {
            "output_version", "work_item_type", "work_item_id",
            "session_type", "payload", "decisions", "open_issues",
        }
        assert required_keys.issubset(envelope.keys())

    def test_payload_keys(self, tmp_path):
        doc = _build_minimal_doc()
        p = _save(doc, tmp_path)
        json_str, _report = parse(p, _work_item())
        envelope = json.loads(json_str)

        payload = envelope["payload"]
        assert "source_metadata" in payload
        assert "business_objects" in payload

    def test_json_roundtrip(self, tmp_path):
        doc = _build_minimal_doc(include_open_issues=True)
        p = _save(doc, tmp_path)
        json_str, _report = parse(p, _work_item())
        envelope = json.loads(json_str)
        # Re-serialize and parse again
        assert json.loads(json.dumps(envelope)) == envelope
