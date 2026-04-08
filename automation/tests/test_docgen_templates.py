"""Tests for automation.docgen.templates — document rendering.

Tests all 8 template modules by feeding hand-built data dictionaries
and verifying the output files are valid.
"""

from docx import Document

from automation.docgen.templates import (
    crm_evaluation_template,
    domain_overview_template,
    domain_prd_template,
    entity_inventory_template,
    entity_prd_template,
    master_prd_template,
    process_document_template,
    yaml_program_template,
)


def _master_prd_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "organization_overview": "We are a test org.\n\nSecond paragraph.",
        "personas": [
            {"id": 1, "name": "Admin", "code": "ADM", "description": "System admin"},
        ],
        "domains": [
            {
                "id": 1, "name": "Mentoring", "code": "MN",
                "description": "Mentoring domain", "sub_domains": [],
                "processes": [
                    {"id": 1, "name": "Client Intake", "code": "MN-INTAKE",
                     "description": "Intake process"},
                ],
            },
        ],
        "services": [],
    }


def _entity_inventory_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "entities": [
            {
                "id": 1, "name": "Contact", "code": "CONTACT",
                "entity_type": "Person", "is_native": True,
                "primary_domain": "Mentoring",
                "process_references": [{"name": "Intake", "code": "MN-INTAKE"}],
            },
            {
                "id": 2, "name": "Engagement", "code": "ENGAGE",
                "entity_type": "Base", "is_native": False,
                "primary_domain": "Mentoring",
                "process_references": [],
            },
        ],
        "business_objects": [
            {
                "id": 1, "name": "Mentor", "description": "A mentor",
                "status": "classified", "resolution": "entity",
                "resolution_detail": None,
                "resolved_entity": "Contact",
                "resolved_process": None,
                "resolved_persona": None,
            },
        ],
    }


def _entity_prd_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "entity": {
            "id": 1, "name": "Contact", "code": "CONTACT",
            "entity_type": "Person", "is_native": False,
            "singular_label": "Contact", "plural_label": "Contacts",
            "description": "A person in the system.",
        },
        "fields": [],
        "native_fields": [
            {
                "name": "firstName", "label": "First Name", "field_type": "varchar",
                "is_required": True, "is_native": True, "description": "First name",
                "process_references": [{"name": "Intake", "code": "MN-INTAKE", "usage": "collected"}],
                "options": [],
            },
        ],
        "custom_fields": [
            {
                "name": "contactType", "label": "Contact Type", "field_type": "multiEnum",
                "is_required": True, "is_native": False, "description": "Type discriminator",
                "category": "Shared",
                "default_value": None,
                "options": [
                    {"value": "Client", "label": "Client"},
                    {"value": "Mentor", "label": "Mentor"},
                ],
                "process_references": [],
            },
        ],
        "relationships": [
            {
                "id": 1, "name": "Contact-Account", "description": "Link to account",
                "entity_id": 1, "entity_foreign_id": 2,
                "link_type": "manyToMany", "link": "accounts", "link_foreign": "contacts",
                "label": "Accounts", "label_foreign": "Contacts",
                "entity_name": "Contact", "foreign_entity_name": "Account",
                "relation_name": None,
            },
        ],
        "layout_panels": [
            {
                "id": 1, "label": "Overview", "description": None,
                "tab_break": False, "tab_label": None, "style": None,
                "hidden": False, "sort_order": 1, "layout_mode": "rows",
                "dynamic_logic_attribute": None, "dynamic_logic_value": None,
                "rows": [
                    {"sort_order": 1, "cell1_name": "firstName", "cell1_label": "First Name",
                     "cell2_name": None, "cell2_label": None, "is_full_width": False},
                ],
                "tabs": [],
            },
        ],
        "list_columns": [],
        "decisions": [
            {"identifier": "DEC-001", "title": "Test", "description": "A decision", "status": "locked"},
        ],
        "open_issues": [
            {"identifier": "ISS-001", "title": "Test", "description": "An issue", "status": "open"},
        ],
        "contributing_domains": [{"name": "Mentoring", "code": "MN"}],
    }


def _domain_overview_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "domain": {"id": 1, "name": "Mentoring", "code": "MN", "description": "Mentoring domain"},
        "parent_domain": None,
        "domain_overview_text": "The Mentoring domain covers...",
        "processes": [
            {"id": 1, "name": "Client Intake", "code": "MN-INTAKE",
             "description": "Intake process", "sort_order": 1},
        ],
        "personas": [
            {"id": 1, "name": "Admin", "code": "ADM", "description": "Admin",
             "roles": [{"role": "performer", "process_name": "Client Intake"}]},
        ],
        "data_reference": [
            {
                "id": 1, "name": "Contact", "code": "CONTACT", "roles": ["primary"],
                "fields": [
                    {"id": 1, "name": "firstName", "label": "First Name",
                     "field_type": "varchar", "usage": "collected"},
                ],
            },
        ],
    }


def _process_document_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "process": {
            "id": 1, "name": "Client Intake", "code": "MN-INTAKE",
            "description": "Process for intake", "triggers": "Form submission",
            "completion_criteria": "Review done", "domain_id": 1,
        },
        "domain": {"id": 1, "name": "Mentoring", "code": "MN"},
        "personas": [
            {"id": 1, "name": "Admin", "code": "ADM", "description": "Admin", "role": "performer"},
        ],
        "steps": [
            {"id": 1, "name": "Submit", "description": "Client submits form",
             "step_type": "action", "sort_order": 1, "performer_name": "Admin"},
        ],
        "requirements": [
            {"identifier": "MN-INTAKE-REQ-001", "description": "Accept submissions",
             "priority": "must", "status": "approved"},
        ],
        "data_reference": [
            {
                "entity_name": "Contact", "entity_code": "CONTACT",
                "fields": [
                    {"name": "firstName", "label": "First Name",
                     "field_type": "varchar", "usage": "collected", "description": "First name"},
                ],
            },
        ],
        "decisions": [],
        "open_issues": [],
        "diagram_path": None,
    }


def _domain_prd_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "domain": {"id": 1, "name": "Mentoring", "code": "MN", "description": "Mentoring domain"},
        "reconciliation_text": "The domain reconciliation determined...",
        "processes": [
            {
                "id": 1, "name": "Client Intake", "code": "MN-INTAKE",
                "description": "Intake", "steps": [], "requirements": [
                    {"identifier": "REQ-001", "description": "Test req",
                     "priority": "must", "status": "approved"},
                ],
            },
        ],
        "personas": [],
        "data_reference": [],
        "decisions": [
            {"identifier": "DEC-001", "title": "T", "description": "A decision", "status": "locked"},
        ],
        "open_issues": [],
    }


def _yaml_program_data():
    return {
        "entities": [
            {
                "id": 1, "name": "Contact", "code": "CONTACT",
                "entity_type": "Person", "is_native": False,
                "singular_label": "Contact", "plural_label": "Contacts",
                "description": "A person",
                "fields": [
                    {
                        "name": "contactType", "label": "Contact Type",
                        "field_type": "multiEnum", "is_required": True,
                        "default_value": None, "read_only": False,
                        "audited": False, "max_length": None, "tooltip": None,
                        "description": "Type", "sort_order": 1,
                        "options": [
                            {"value": "Client", "label": "Client",
                             "style": None, "sort_order": 1, "is_default": False},
                        ],
                    },
                ],
                "relationships": [],
                "layout_panels": [],
                "list_columns": [],
            },
        ],
    }


def _crm_evaluation_data():
    return {
        "client_name": "Test Org",
        "client_short_name": "TO",
        "crm_platform": "EspoCRM",
        "decisions": [
            {"identifier": "CRM-DEC-001", "title": "T", "description": "Use EspoCRM", "status": "locked"},
        ],
        "open_issues": [],
        "requirements_summary": {"total": 10, "must": 5, "should": 3, "may": 2},
        "scale_summary": {"entity_count": 5, "field_count": 30, "relationship_count": 8},
    }


class TestMasterPrdTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "Master-PRD.docx"
        master_prd_template.generate(_master_prd_data(), out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test Org" in text
        assert "Mentoring" in text


class TestEntityInventoryTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "Entity-Inventory.docx"
        entity_inventory_template.generate(_entity_inventory_data(), out)
        assert out.exists()
        doc = Document(str(out))
        # Entity names appear in table cells
        all_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        combined = all_text + table_text
        assert "Contact" in combined


class TestEntityPrdTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "Contact-Entity-PRD.docx"
        entity_prd_template.generate(_entity_prd_data(), out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Contact" in text

    def test_draft_mode(self, tmp_path):
        out = tmp_path / "Contact-Entity-PRD-draft.docx"
        entity_prd_template.generate(_entity_prd_data(), out, is_draft=True)
        assert out.exists()


class TestDomainOverviewTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "Domain-Overview.docx"
        domain_overview_template.generate(_domain_overview_data(), out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Mentoring" in text


class TestProcessDocumentTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "MN-INTAKE.docx"
        process_document_template.generate(_process_document_data(), out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Client Intake" in text
        assert "MN-INTAKE" in text


class TestDomainPrdTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "Domain-PRD.docx"
        domain_prd_template.generate(_domain_prd_data(), out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Mentoring" in text


class TestYamlProgramTemplate:
    def test_generates_valid_yaml(self, tmp_path):
        out = tmp_path / "Contact.yaml"
        yaml_program_template.generate(_yaml_program_data(), out)
        assert out.exists()
        content = out.read_text()
        assert "contactType" in content
        assert "multiEnum" in content

    def test_generates_to_directory(self, tmp_path):
        out_dir = tmp_path / "programs"
        out_dir.mkdir()
        yaml_program_template.generate(_yaml_program_data(), out_dir)
        files = list(out_dir.glob("*.yaml"))
        assert len(files) == 1
        assert files[0].name == "Contact.yaml"


class TestCrmEvaluationTemplate:
    def test_generates_valid_docx(self, tmp_path):
        out = tmp_path / "CRM-Eval.docx"
        crm_evaluation_template.generate(_crm_evaluation_data(), out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "EspoCRM" in text  # Product name allowed in CRM Evaluation
        assert "Test Org" in text
