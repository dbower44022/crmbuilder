"""Tests for automation.importer.parsers.process_doc_docx."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from automation.importer.parsers import ParseReport, ProcessDocParseError
from automation.importer.parsers.process_doc_docx import parse

FIXTURE_MN = Path(__file__).parent / "fixtures" / "cbm-mn-intake-v2.3.docx"
FIXTURE_CR = Path(__file__).parent / "fixtures" / "cbm-cr-partner-manage-v1.0.docx"

# Skip real-document tests if fixture not present
skip_mn = pytest.mark.skipif(
    not FIXTURE_MN.exists(),
    reason="cbm-mn-intake-v2.3.docx fixture not committed",
)
skip_cr = pytest.mark.skipif(
    not FIXTURE_CR.exists(),
    reason="cbm-cr-partner-manage-v1.0.docx fixture not committed",
)


def _work_item(wi_id: int = 99, process_id: int = 1) -> dict:
    return {"id": wi_id, "item_type": "process_definition", "process_id": process_id}


# -------------------------------------------------------------------------
# Helpers for building synthetic documents
# -------------------------------------------------------------------------


def _save(doc: Document, tmp_path: Path, name: str = "test.docx") -> Path:
    p = tmp_path / name
    doc.save(str(p))
    return p


def _build_minimal_valid_doc(
    *,
    domain: str = "Testing (TST)",
    sub_domain: str | None = None,
    process_code: str = "TST-PROC",
    process_name: str = "Test Process",
    include_version: bool = True,
    include_status: bool = True,
    include_last_updated: bool = True,
) -> Document:
    """Build a minimal but structurally valid process document."""
    doc = Document()

    # Header table
    table = doc.add_table(rows=0, cols=2)
    _add_row(table, "Domain", domain)
    if sub_domain is not None:
        _add_row(table, "Sub-Domain", sub_domain)
    _add_row(table, "Process Code", process_code)
    _add_row(table, "Process Name", process_name)
    if include_version:
        _add_row(table, "Version", "1.0")
    if include_status:
        _add_row(table, "Status", "Draft")
    if include_last_updated:
        _add_row(table, "Last Updated", "04-10-26 12:00")

    # Section 1 — Process Purpose
    doc.add_heading("1. Process Purpose", level=1)
    doc.add_paragraph("This process handles test operations.")

    # Section 2 — Process Triggers
    doc.add_heading("2. Process Triggers", level=1)
    doc.add_paragraph("Triggered by test events.")

    # Section 3 — Personas Involved
    doc.add_heading("3. Personas Involved", level=1)
    doc.add_paragraph("Test Persona (MST-PER-001)")
    doc.add_paragraph("Performs the test activities.")

    # Section 4 — Process Workflow
    doc.add_heading("4. Process Workflow", level=1)
    p = doc.add_paragraph("Step one of the workflow.")
    p.style = doc.styles["List Paragraph"]
    p2 = doc.add_paragraph("Step two of the workflow.")
    p2.style = doc.styles["List Paragraph"]

    # Section 5 — Process Completion
    doc.add_heading("5. Process Completion", level=1)
    doc.add_paragraph("Process is complete when testing finishes.")

    # Section 6 — System Requirements
    doc.add_heading("6. System Requirements", level=1)
    req_table = doc.add_table(rows=2, cols=2)
    req_table.cell(0, 0).text = "ID"
    req_table.cell(0, 1).text = "Requirement"
    req_table.cell(1, 0).text = f"{process_code}-REQ-001"
    req_table.cell(1, 1).text = "The system must support testing."

    # Section 7 — Process Data
    doc.add_heading("7. Process Data", level=1)
    doc.add_paragraph("No pre-existing data referenced.")

    # Section 8 — Data Collected
    doc.add_heading("8. Data Collected", level=1)
    doc.add_paragraph("No data collected.")

    # Section 9 — Open Issues
    doc.add_heading("9. Open Issues", level=1)
    doc.add_paragraph("No open issues.")

    return doc


def _add_row(table, key: str, value: str) -> None:
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value


def _add_field_table(doc: Document, fields: list[tuple[str, str, str, str, str, str, str]]) -> None:
    """Add a six-column, two-row-per-field table.

    Each tuple: (name, type, required, values, default, identifier, description)
    """
    row_count = 1 + len(fields) * 2
    table = doc.add_table(rows=row_count, cols=6)
    headers = ["Field Name", "Type", "Required", "Values", "Default", "ID"]
    for ci, h in enumerate(headers):
        table.cell(0, ci).text = h

    for i, (name, ftype, req, vals, default, ident, desc) in enumerate(fields):
        meta_row = 1 + i * 2
        desc_row = 2 + i * 2
        table.cell(meta_row, 0).text = name
        table.cell(meta_row, 1).text = ftype
        table.cell(meta_row, 2).text = req
        table.cell(meta_row, 3).text = vals
        table.cell(meta_row, 4).text = default
        table.cell(meta_row, 5).text = ident
        for ci in range(6):
            table.cell(desc_row, ci).text = desc


# -------------------------------------------------------------------------
# 1 — Real document: MN-INTAKE
# -------------------------------------------------------------------------


@skip_mn
class TestMNIntake:

    def test_parse_returns_envelope_and_report(self):
        envelope_json, report = parse(FIXTURE_MN, _work_item())
        assert isinstance(envelope_json, str)
        assert isinstance(report, ParseReport)
        envelope = json.loads(envelope_json)
        assert "output_version" in envelope
        assert "payload" in envelope

    def test_source_metadata(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        meta = json.loads(envelope_json)["payload"]["source_metadata"]
        assert meta["domain_code"] == "MN"
        assert meta["process_code"] == "MN-INTAKE"
        assert "sub_domain_code" not in meta

    def test_personas(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        personas = json.loads(envelope_json)["payload"]["personas"]
        assert len(personas) == 2
        codes = {p["identifier"] for p in personas}
        assert codes == {"MST-PER-013", "MST-PER-003"}

    def test_workflow_format_a(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        steps = json.loads(envelope_json)["payload"]["workflow"]
        assert len(steps) == 9
        for s in steps:
            assert s["step_type"] == "action"
        sort_orders = [s["sort_order"] for s in steps]
        assert sort_orders == list(range(1, 10))

    def test_requirements(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        reqs = json.loads(envelope_json)["payload"]["system_requirements"]
        assert len(reqs) == 9
        for r in reqs:
            assert r["identifier"].startswith("MN-INTAKE-REQ-")
            assert r["priority"] == "must"

    def test_data_collected_entities(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        dc = json.loads(envelope_json)["payload"]["data_collected"]
        names = [e["entity_name"] for e in dc]
        assert names == ["Client Organization", "Client Contact", "Engagement"]
        assert len(dc[0]["new_fields"]) == 8   # Client Organization
        assert len(dc[1]["new_fields"]) == 11  # Client Contact
        assert len(dc[2]["new_fields"]) == 7   # Engagement

    def test_process_data_empty(self):
        envelope_json, report = parse(FIXTURE_MN, _work_item())
        pd = json.loads(envelope_json)["payload"]["process_data"]
        assert pd == []
        assert any(
            w.category == "empty_section" and "Section 7" in w.location
            for w in report.warnings
        )

    def test_open_issues(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        issues = json.loads(envelope_json)["open_issues"]
        assert len(issues) == 2
        ids = {i["identifier"] for i in issues}
        assert "MN-INTAKE-ISS-001" in ids
        assert "MN-INTAKE-ISS-002" in ids

    def test_process_code_prefix_validates(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        meta = json.loads(envelope_json)["payload"]["source_metadata"]
        assert meta["process_code"].startswith(f"{meta['domain_code']}-")


# -------------------------------------------------------------------------
# 2 — Real document: CR-PARTNER-MANAGE (skipped until fixture exists)
# -------------------------------------------------------------------------


@skip_cr
class TestCRPartnerManage:

    def test_parse_returns_envelope(self):
        envelope_json, report = parse(FIXTURE_CR, _work_item())
        assert isinstance(envelope_json, str)
        envelope = json.loads(envelope_json)
        assert "payload" in envelope

    def test_source_metadata_with_subdomain(self):
        envelope_json, _ = parse(FIXTURE_CR, _work_item())
        meta = json.loads(envelope_json)["payload"]["source_metadata"]
        assert meta["domain_code"] == "CR"
        assert meta["sub_domain_code"] == "CR-PARTNER"
        assert meta["process_code"].startswith("CR-PARTNER-")

    def test_workflow_format_b(self):
        envelope_json, _ = parse(FIXTURE_CR, _work_item())
        steps = json.loads(envelope_json)["payload"]["workflow"]
        assert len(steps) >= 5
        for s in steps:
            assert s["step_type"] == "action"
        sort_orders = [s["sort_order"] for s in steps]
        assert sort_orders == list(range(1, len(steps) + 1))


# -------------------------------------------------------------------------
# 3–15 — Hard-failure synthetic tests
# -------------------------------------------------------------------------


class TestHardFailures:

    def test_missing_header_table(self, tmp_path):
        doc = Document()
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose text.")
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="No header table"):
            parse(path, _work_item())

    def test_missing_process_code_row(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Process Code"):
            parse(path, _work_item())

    def test_process_code_prefix_mismatch_domain(self, tmp_path):
        doc = _build_minimal_valid_doc(
            domain="Client Recruiting (CR)",
            process_code="MN-FOO",
        )
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="does not start with.*domain code prefix"):
            parse(path, _work_item())

    def test_process_code_prefix_mismatch_subdomain(self, tmp_path):
        doc = _build_minimal_valid_doc(
            domain="Client Recruiting (CR)",
            sub_domain="Partner (CR-PARTNER)",
            process_code="CR-MARKETING-FOO",
        )
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="does not start with.*sub-domain code prefix"):
            parse(path, _work_item())

    def test_domain_bad_pattern(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "BadDomain"
        table.cell(1, 0).text = "Process Code"
        table.cell(1, 1).text = "BD-PROC"
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Name \\(CODE\\)"):
            parse(path, _work_item())

    def test_subdomain_bad_pattern(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        table.cell(1, 0).text = "Sub-Domain"
        table.cell(1, 1).text = "BadSubDomain"
        table.cell(2, 0).text = "Process Code"
        table.cell(2, 1).text = "TST-PROC"
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Sub-Domain.*Name \\(CODE\\)"):
            parse(path, _work_item())

    def test_missing_section1_purpose(self, tmp_path):
        # Build without Section 1
        doc2 = Document()
        table = doc2.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        table.cell(1, 0).text = "Process Code"
        table.cell(1, 1).text = "TST-PROC"
        doc2.add_heading("2. Process Triggers", level=1)
        doc2.add_paragraph("Triggers.")
        path = _save(doc2, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Section 1.*Process Purpose"):
            parse(path, _work_item())

    def test_missing_section3_personas(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        table.cell(1, 0).text = "Process Code"
        table.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        # Skip Section 3
        doc.add_heading("4. Process Workflow", level=1)
        doc.add_paragraph("Workflow.")
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Section 3.*Personas"):
            parse(path, _work_item())

    def test_section3_zero_personas(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        table.cell(1, 0).text = "Process Code"
        table.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("No personas match the expected pattern.")
        doc.add_heading("4. Process Workflow", level=1)
        doc.add_paragraph("Workflow.")
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="zero personas"):
            parse(path, _work_item())

    def test_missing_section4_workflow(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        table.cell(1, 0).text = "Process Code"
        table.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("Test Persona (MST-PER-001)")
        doc.add_paragraph("Description.")
        # Skip Section 4
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Section 4.*Workflow"):
            parse(path, _work_item())

    def test_section4_zero_steps(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Domain"
        table.cell(0, 1).text = "Testing (TST)"
        table.cell(1, 0).text = "Process Code"
        table.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("Test Persona (MST-PER-001)")
        doc.add_paragraph("Description.")
        doc.add_heading("4. Process Workflow", level=1)
        # No list paragraphs or H2 headings — only normal text
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        path = _save(doc, tmp_path)
        with pytest.raises(ProcessDocParseError, match="zero steps"):
            parse(path, _work_item())

    def test_missing_section6_requirements(self, tmp_path):
        # Build without Section 6
        doc2 = Document()
        tbl = doc2.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc2.add_heading("1. Process Purpose", level=1)
        doc2.add_paragraph("Purpose.")
        doc2.add_heading("2. Process Triggers", level=1)
        doc2.add_paragraph("Triggers.")
        doc2.add_heading("3. Personas Involved", level=1)
        doc2.add_paragraph("Test Persona (MST-PER-001)")
        doc2.add_paragraph("Description.")
        doc2.add_heading("4. Process Workflow", level=1)
        p = doc2.add_paragraph("Step one.")
        p.style = doc2.styles["List Paragraph"]
        doc2.add_heading("5. Process Completion", level=1)
        doc2.add_paragraph("Done.")
        # Skip Section 6 — go to Section 7
        doc2.add_heading("7. Process Data", level=1)
        doc2.add_paragraph("No data.")
        doc2.add_heading("8. Data Collected", level=1)
        doc2.add_paragraph("None.")
        doc2.add_heading("9. Open Issues", level=1)
        doc2.add_paragraph("None.")
        path = _save(doc2, tmp_path)
        with pytest.raises(ProcessDocParseError, match="Section 6.*System Requirements"):
            parse(path, _work_item())

    def test_section6_no_matching_table(self, tmp_path):
        # Build with Section 6 but no matching table
        doc2 = Document()
        tbl = doc2.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc2.add_heading("1. Process Purpose", level=1)
        doc2.add_paragraph("Purpose.")
        doc2.add_heading("2. Process Triggers", level=1)
        doc2.add_paragraph("Triggers.")
        doc2.add_heading("3. Personas Involved", level=1)
        doc2.add_paragraph("Test Persona (MST-PER-001)")
        doc2.add_paragraph("Description.")
        doc2.add_heading("4. Process Workflow", level=1)
        p = doc2.add_paragraph("Step one.")
        p.style = doc2.styles["List Paragraph"]
        doc2.add_heading("5. Process Completion", level=1)
        doc2.add_paragraph("Done.")
        doc2.add_heading("6. System Requirements", level=1)
        doc2.add_paragraph("No table here.")
        doc2.add_heading("7. Process Data", level=1)
        doc2.add_paragraph("No data.")
        doc2.add_heading("8. Data Collected", level=1)
        doc2.add_paragraph("None.")
        doc2.add_heading("9. Open Issues", level=1)
        doc2.add_paragraph("None.")
        path = _save(doc2, tmp_path)
        with pytest.raises(ProcessDocParseError, match="no matching"):
            parse(path, _work_item())

    def test_wrong_work_item_type(self, tmp_path):
        path = tmp_path / "dummy.docx"
        path.write_bytes(b"")
        with pytest.raises(ValueError, match="process_definition"):
            parse(path, {"id": 1, "item_type": "master_prd"})

    def test_file_not_found(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            parse(tmp_path / "nonexistent.docx", _work_item())


# -------------------------------------------------------------------------
# 16–22 — Soft-warning synthetic tests
# -------------------------------------------------------------------------


class TestSoftWarnings:

    def test_missing_version_status_last_updated(self, tmp_path):
        doc = _build_minimal_valid_doc(
            include_version=False,
            include_status=False,
            include_last_updated=False,
        )
        path = _save(doc, tmp_path)
        _, report = parse(path, _work_item())
        field_warnings = [
            w for w in report.warnings if w.category == "missing_field"
        ]
        texts = " ".join(w.message for w in field_warnings)
        assert "Version" in texts
        assert "Status" in texts
        assert "Last Updated" in texts

    def test_empty_section7_no_entity_subsections(self, tmp_path):
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        _, report = parse(path, _work_item())
        pd = json.loads(parse(path, _work_item())[0])["payload"]["process_data"]
        assert pd == []
        assert any(
            w.category == "empty_section" and "Section 7" in w.location
            for w in report.warnings
        )

    def test_empty_section9_no_issues_table(self, tmp_path):
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        envelope_json, _ = parse(path, _work_item())
        issues = json.loads(envelope_json)["open_issues"]
        assert issues == []

    def test_requirement_identifier_mismatch(self, tmp_path):
        # Build with mismatched requirement identifiers
        doc2 = Document()
        tbl = doc2.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc2.add_heading("1. Process Purpose", level=1)
        doc2.add_paragraph("Purpose.")
        doc2.add_heading("2. Process Triggers", level=1)
        doc2.add_paragraph("Triggers.")
        doc2.add_heading("3. Personas Involved", level=1)
        doc2.add_paragraph("Test Persona (MST-PER-001)")
        doc2.add_paragraph("Desc.")
        doc2.add_heading("4. Process Workflow", level=1)
        p = doc2.add_paragraph("Step one.")
        p.style = doc2.styles["List Paragraph"]
        doc2.add_heading("5. Process Completion", level=1)
        doc2.add_paragraph("Done.")
        doc2.add_heading("6. System Requirements", level=1)
        req_table = doc2.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "OTHER-PROC-REQ-001"
        req_table.cell(1, 1).text = "Some requirement."
        doc2.add_heading("7. Process Data", level=1)
        doc2.add_paragraph("No data.")
        doc2.add_heading("8. Data Collected", level=1)
        doc2.add_paragraph("None.")
        doc2.add_heading("9. Open Issues", level=1)
        doc2.add_paragraph("None.")
        path = _save(doc2, tmp_path)
        _, report = parse(path, _work_item())
        id_warnings = [
            w for w in report.warnings if w.category == "identifier_mismatch"
        ]
        assert len(id_warnings) == 1
        assert "OTHER-PROC-REQ-001" in id_warnings[0].location

    def test_field_table_odd_row_count(self, tmp_path):
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("Test Persona (MST-PER-001)")
        doc.add_paragraph("Desc.")
        doc.add_heading("4. Process Workflow", level=1)
        p = doc.add_paragraph("Step one.")
        p.style = doc.styles["List Paragraph"]
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        doc.add_heading("6. System Requirements", level=1)
        req_table = doc.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "TST-PROC-REQ-001"
        req_table.cell(1, 1).text = "A requirement."
        doc.add_heading("7. Process Data", level=1)
        doc.add_paragraph("No data.")
        doc.add_heading("8. Data Collected", level=1)
        doc.add_heading("Entity: TestEntity", level=2)
        doc.add_paragraph("Description.")
        # Field table with odd data rows (header + 1 meta row, no desc row)
        ft = doc.add_table(rows=2, cols=6)
        for ci, h in enumerate(["Field Name", "Type", "Required", "Values", "Default", "ID"]):
            ft.cell(0, ci).text = h
        ft.cell(1, 0).text = "testField"
        ft.cell(1, 1).text = "varchar"
        ft.cell(1, 2).text = "Yes"
        ft.cell(1, 3).text = "—"
        ft.cell(1, 4).text = "—"
        ft.cell(1, 5).text = "TST-PROC-DAT-001"
        doc.add_heading("9. Open Issues", level=1)
        doc.add_paragraph("None.")
        path = _save(doc, tmp_path)
        _, report = parse(path, _work_item())
        odd_warnings = [w for w in report.warnings if w.category == "odd_row_count"]
        assert len(odd_warnings) >= 1

    def test_section9_multiple_tables(self, tmp_path):
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("Test Persona (MST-PER-001)")
        doc.add_paragraph("Desc.")
        doc.add_heading("4. Process Workflow", level=1)
        p = doc.add_paragraph("Step one.")
        p.style = doc.styles["List Paragraph"]
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        doc.add_heading("6. System Requirements", level=1)
        req_table = doc.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "TST-PROC-REQ-001"
        req_table.cell(1, 1).text = "A requirement."
        doc.add_heading("7. Process Data", level=1)
        doc.add_paragraph("No data.")
        doc.add_heading("8. Data Collected", level=1)
        doc.add_paragraph("None.")
        doc.add_heading("9. Open Issues", level=1)
        # First table (process-owned issues)
        iss1 = doc.add_table(rows=2, cols=2)
        iss1.cell(0, 0).text = "ID"
        iss1.cell(0, 1).text = "Issue"
        iss1.cell(1, 0).text = "TST-PROC-ISS-001"
        iss1.cell(1, 1).text = "First issue."
        # Second table (inherited — should be warned about)
        iss2 = doc.add_table(rows=2, cols=2)
        iss2.cell(0, 0).text = "ID"
        iss2.cell(0, 1).text = "Issue"
        iss2.cell(1, 0).text = "OTHER-ISS-001"
        iss2.cell(1, 1).text = "Inherited issue."
        path = _save(doc, tmp_path)
        envelope_json, report = parse(path, _work_item())
        issues = json.loads(envelope_json)["open_issues"]
        # Only first table parsed
        assert len(issues) == 1
        assert issues[0]["identifier"] == "TST-PROC-ISS-001"
        multi_warnings = [
            w for w in report.warnings if w.category == "multiple_tables"
        ]
        assert len(multi_warnings) == 1

    def test_unclear_required_value(self, tmp_path):
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("Test Persona (MST-PER-001)")
        doc.add_paragraph("Desc.")
        doc.add_heading("4. Process Workflow", level=1)
        p = doc.add_paragraph("Step one.")
        p.style = doc.styles["List Paragraph"]
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        doc.add_heading("6. System Requirements", level=1)
        req_table = doc.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "TST-PROC-REQ-001"
        req_table.cell(1, 1).text = "A requirement."
        doc.add_heading("7. Process Data", level=1)
        doc.add_paragraph("No data.")
        doc.add_heading("8. Data Collected", level=1)
        doc.add_heading("Entity: TestEntity", level=2)
        doc.add_paragraph("Entity desc.")
        # Field table with unclear required value
        ft = doc.add_table(rows=3, cols=6)
        for ci, h in enumerate(["Field Name", "Type", "Required", "Values", "Default", "ID"]):
            ft.cell(0, ci).text = h
        ft.cell(1, 0).text = "testField"
        ft.cell(1, 1).text = "varchar"
        ft.cell(1, 2).text = "Yes (system-generated)"
        ft.cell(1, 3).text = "—"
        ft.cell(1, 4).text = "—"
        ft.cell(1, 5).text = "TST-PROC-DAT-001"
        for ci in range(6):
            ft.cell(2, ci).text = "Description of testField."
        doc.add_heading("9. Open Issues", level=1)
        doc.add_paragraph("None.")
        path = _save(doc, tmp_path)
        envelope_json, report = parse(path, _work_item())
        dc = json.loads(envelope_json)["payload"]["data_collected"]
        assert len(dc) == 1
        field = dc[0]["new_fields"][0]
        # "Yes (system-generated)" starts with "yes" → treated as True, no warning
        # Actually this starts with "Yes" so it IS "yes"-prefixed, no warning
        assert field["is_required"] is True
        # Now test with truly unclear value
        ft.cell(1, 2).text = "Conditional"
        doc.save(str(path))
        _, report2 = parse(path, _work_item())
        unclear_warnings = [
            w for w in report2.warnings if w.category == "unclear_required"
        ]
        assert len(unclear_warnings) >= 1


# -------------------------------------------------------------------------
# 23–26 — Round-trip and structural tests
# -------------------------------------------------------------------------


class TestStructural:

    def test_envelope_json_round_trips(self, tmp_path):
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        envelope_json, _ = parse(path, _work_item())
        parsed = json.loads(envelope_json)
        assert isinstance(parsed, dict)

    def test_envelope_has_required_keys(self, tmp_path):
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        envelope_json, _ = parse(path, _work_item(42))
        envelope = json.loads(envelope_json)
        required = {"output_version", "work_item_type", "work_item_id",
                     "session_type", "payload", "decisions", "open_issues"}
        assert required.issubset(set(envelope.keys()))
        assert envelope["output_version"] == "1.0"
        assert envelope["work_item_type"] == "process_definition"
        assert envelope["work_item_id"] == 42
        assert envelope["session_type"] == "initial"

    def test_payload_has_required_keys(self, tmp_path):
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        envelope_json, _ = parse(path, _work_item())
        payload = json.loads(envelope_json)["payload"]
        required = {"process_purpose", "triggers", "personas", "workflow",
                     "completion", "system_requirements", "process_data",
                     "data_collected"}
        assert required.issubset(set(payload.keys()))
        # triggers and completion are dicts (to pass Layer 3 validation)
        assert isinstance(payload["triggers"], dict)
        assert isinstance(payload["completion"], dict)

    def test_adapter_does_not_need_database(self, tmp_path):
        """Verify parse() works with no database connection."""
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        # work_item has no 'conn' — parse should not attempt DB lookups
        envelope_json, report = parse(path, {"id": 1, "item_type": "process_definition"})
        assert isinstance(envelope_json, str)
        assert isinstance(report, ParseReport)


# -------------------------------------------------------------------------
# Extra: Persona format tests
# -------------------------------------------------------------------------


class TestPersonaFormats:

    def test_format_a_paragraph_scan(self, tmp_path):
        doc = _build_minimal_valid_doc()
        path = _save(doc, tmp_path)
        envelope_json, _ = parse(path, _work_item())
        personas = json.loads(envelope_json)["payload"]["personas"]
        assert len(personas) == 1
        assert personas[0]["identifier"] == "MST-PER-001"
        assert personas[0]["name"] == "Test Persona"
        assert personas[0]["role"] == "performer"

    def test_format_b_persona_table(self, tmp_path):
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        # Persona table
        ptable = doc.add_table(rows=3, cols=3)
        ptable.cell(0, 0).text = "ID"
        ptable.cell(0, 1).text = "Persona"
        ptable.cell(0, 2).text = "Role"
        ptable.cell(1, 0).text = "MST-PER-001"
        ptable.cell(1, 1).text = "Admin User"
        ptable.cell(1, 2).text = "Initiates and approves the process."
        ptable.cell(2, 0).text = "MST-PER-002"
        ptable.cell(2, 1).text = "Reviewer"
        ptable.cell(2, 2).text = "Receives notification of completion."
        doc.add_heading("4. Process Workflow", level=1)
        p = doc.add_paragraph("Step one.")
        p.style = doc.styles["List Paragraph"]
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        doc.add_heading("6. System Requirements", level=1)
        req_table = doc.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "TST-PROC-REQ-001"
        req_table.cell(1, 1).text = "A requirement."
        doc.add_heading("7. Process Data", level=1)
        doc.add_paragraph("No data.")
        doc.add_heading("8. Data Collected", level=1)
        doc.add_paragraph("None.")
        doc.add_heading("9. Open Issues", level=1)
        doc.add_paragraph("None.")
        path = _save(doc, tmp_path)
        envelope_json, _ = parse(path, _work_item())
        personas = json.loads(envelope_json)["payload"]["personas"]
        assert len(personas) == 2
        assert personas[0]["identifier"] == "MST-PER-001"
        assert personas[0]["name"] == "Admin User"
        assert personas[0]["role"] == "initiator"  # "initiates and approves"
        assert personas[1]["identifier"] == "MST-PER-002"
        assert personas[1]["role"] == "recipient"  # "receives notification"

    def test_role_derivation(self, tmp_path):
        # Build with persona that has "approve" in description
        doc2 = Document()
        tbl = doc2.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc2.add_heading("1. Process Purpose", level=1)
        doc2.add_paragraph("Purpose.")
        doc2.add_heading("2. Process Triggers", level=1)
        doc2.add_paragraph("Triggers.")
        doc2.add_heading("3. Personas Involved", level=1)
        doc2.add_paragraph("Approver Person (MST-PER-001)")
        doc2.add_paragraph("Reviews and approves all submissions.")
        doc2.add_heading("4. Process Workflow", level=1)
        p = doc2.add_paragraph("Step one.")
        p.style = doc2.styles["List Paragraph"]
        doc2.add_heading("5. Process Completion", level=1)
        doc2.add_paragraph("Done.")
        doc2.add_heading("6. System Requirements", level=1)
        req_table = doc2.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "TST-PROC-REQ-001"
        req_table.cell(1, 1).text = "A requirement."
        doc2.add_heading("7. Process Data", level=1)
        doc2.add_paragraph("No data.")
        doc2.add_heading("8. Data Collected", level=1)
        doc2.add_paragraph("None.")
        doc2.add_heading("9. Open Issues", level=1)
        doc2.add_paragraph("None.")
        path = _save(doc2, tmp_path)
        envelope_json, _ = parse(path, _work_item())
        personas = json.loads(envelope_json)["payload"]["personas"]
        assert personas[0]["role"] == "approver"


# -------------------------------------------------------------------------
# Workflow format B tests
# -------------------------------------------------------------------------


class TestWorkflowFormatB:

    def test_activity_area_steps(self, tmp_path):
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Domain"
        tbl.cell(0, 1).text = "Testing (TST)"
        tbl.cell(1, 0).text = "Process Code"
        tbl.cell(1, 1).text = "TST-PROC"
        doc.add_heading("1. Process Purpose", level=1)
        doc.add_paragraph("Purpose.")
        doc.add_heading("2. Process Triggers", level=1)
        doc.add_paragraph("Triggers.")
        doc.add_heading("3. Personas Involved", level=1)
        doc.add_paragraph("Test Persona (MST-PER-001)")
        doc.add_paragraph("Desc.")
        doc.add_heading("4. Process Workflow", level=1)
        doc.add_paragraph("Some framing text before activities.")
        doc.add_heading("4.1 First Activity", level=2)
        doc.add_paragraph("Description of first activity.")
        doc.add_heading("4.2 Second Activity", level=2)
        doc.add_paragraph("Description of second activity.")
        p = doc.add_paragraph("A bullet under activity 2.")
        p.style = doc.styles["List Paragraph"]
        doc.add_heading("5. Process Completion", level=1)
        doc.add_paragraph("Done.")
        doc.add_heading("6. System Requirements", level=1)
        req_table = doc.add_table(rows=2, cols=2)
        req_table.cell(0, 0).text = "ID"
        req_table.cell(0, 1).text = "Requirement"
        req_table.cell(1, 0).text = "TST-PROC-REQ-001"
        req_table.cell(1, 1).text = "A requirement."
        doc.add_heading("7. Process Data", level=1)
        doc.add_paragraph("No data.")
        doc.add_heading("8. Data Collected", level=1)
        doc.add_paragraph("None.")
        doc.add_heading("9. Open Issues", level=1)
        doc.add_paragraph("None.")
        path = _save(doc, tmp_path)
        envelope_json, report = parse(path, _work_item())
        steps = json.loads(envelope_json)["payload"]["workflow"]
        assert len(steps) == 2
        assert steps[0]["name"] == "First Activity"
        assert steps[0]["sort_order"] == 1
        assert steps[1]["name"] == "Second Activity"
        assert steps[1]["sort_order"] == 2
        # Bullet included in step 2's description
        assert "bullet" in steps[1]["description"].lower()
        # Framing prose warned
        framing_warnings = [
            w for w in report.warnings if w.category == "framing_prose"
        ]
        assert len(framing_warnings) == 1


# -------------------------------------------------------------------------
# Path B validation round-trip
# -------------------------------------------------------------------------


@skip_mn
class TestPathBRoundTrip:

    def test_envelope_passes_layer3_validation(self):
        envelope_json, _ = parse(FIXTURE_MN, _work_item())
        from automation.importer.parser import parse_and_validate
        parse_and_validate(
            envelope_json,
            expected_item_type="process_definition",
            expected_work_item_id=99,
            expected_session_type="initial",
        )

    def test_parsed_counts_keys(self):
        _, report = parse(FIXTURE_MN, _work_item())
        assert "persona" in report.parsed_counts
        assert "workflow_step" in report.parsed_counts
        assert "requirement" in report.parsed_counts
