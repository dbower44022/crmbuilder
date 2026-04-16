"""Tests for automation.importer.parsers.entity_prd_docx."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from automation.importer.parsers import EntityPrdParseError, ParseReport
from automation.importer.parsers.entity_prd_docx import parse

FIXTURE = Path(__file__).parent / "fixtures" / "cbm-contact-entity-prd-v1.3.docx"

skip_fixture = pytest.mark.skipif(
    not FIXTURE.exists(),
    reason="cbm-contact-entity-prd-v1.3.docx fixture not committed",
)


def _work_item(wi_id: int = 99, entity_id: int = 1) -> dict:
    return {"id": wi_id, "item_type": "entity_prd", "entity_id": entity_id}


# -------------------------------------------------------------------------
# Helpers for building synthetic documents
# -------------------------------------------------------------------------


def _save(doc: Document, tmp_path: Path, name: str = "test.docx") -> Path:
    p = tmp_path / name
    doc.save(str(p))
    return p


def _add_row(table, key: str, value: str) -> None:
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value


def _add_field_row(table, name, ftype, required, values, default, fid):
    """Add a metadata row to a 6-column field table."""
    row = table.add_row()
    row.cells[0].text = name
    row.cells[1].text = ftype
    row.cells[2].text = required
    row.cells[3].text = values
    row.cells[4].text = default
    row.cells[5].text = fid


def _add_desc_row(table, desc: str) -> None:
    """Add a description row spanning all columns."""
    row = table.add_row()
    row.cells[0].text = desc


def _build_minimal_valid_doc(
    *,
    entity_value: str = "TestEntity (Custom — Base Type)",
    include_header: bool = True,
    include_overview: bool = True,
    overview_entity_name: str = "TestEntity",
    contributing_domains: str = "Testing (TST)",
    primary_domain: str | None = None,
    include_section1: bool = True,
    include_section2: bool = True,
    native_field_count: int = 2,
    include_section3: bool = True,
    custom_field_count: int = 1,
    include_section4: bool = True,
    rel_count: int = 1,
    include_section5: bool = False,
    include_section6: bool = False,
    include_section7: bool = False,
    include_section8: bool = True,
    include_section9: bool = True,
    include_version: bool = True,
) -> Document:
    """Build a minimal but structurally valid entity PRD document."""
    doc = Document()

    # Header table
    if include_header:
        table = doc.add_table(rows=0, cols=2)
        _add_row(table, "Document Type", "Entity PRD")
        _add_row(table, "Entity", entity_value)
        _add_row(table, "Implementation", "Test")
        if include_version:
            _add_row(table, "Version", "1.0")
        _add_row(table, "Status", "Draft")
        _add_row(table, "Last Updated", "04-10-26 12:00")
        _add_row(table, "Source Documents", "Test Source Doc v1.0")

    # Entity Overview table
    if include_overview:
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", overview_entity_name)
        _add_row(ov, "Native / Custom", "Custom")
        _add_row(ov, "Entity Type", "Base")
        _add_row(ov, "Display Label (Singular)", overview_entity_name)
        _add_row(ov, "Display Label (Plural)", f"{overview_entity_name}s")
        _add_row(ov, "Activity Stream", "No")
        if primary_domain is not None:
            _add_row(ov, "Primary Domain", primary_domain)
        _add_row(ov, "Contributing Domains", contributing_domains)

    # Section 1 — Entity Overview prose
    if include_section1:
        doc.add_heading("1. Entity Overview", level=1)
        doc.add_paragraph("Test entity overview description.")

    # Section 2 — Native Fields
    if include_section2:
        doc.add_heading("2. Native Fields", level=1)
        nf_table = doc.add_table(rows=1, cols=4)
        nf_table.cell(0, 0).text = "Native Field"
        nf_table.cell(0, 1).text = "Type"
        nf_table.cell(0, 2).text = "PRD Name(s) / Mapping"
        nf_table.cell(0, 3).text = "Referenced By"
        for i in range(native_field_count):
            row = nf_table.add_row()
            row.cells[0].text = f"field{i}"
            row.cells[1].text = "varchar"
            row.cells[2].text = f"Test field {i}"
            row.cells[3].text = ""

    # Section 3 — Custom Fields
    if include_section3:
        doc.add_heading("3. Custom Fields", level=1)
        doc.add_heading("3.1 Test Fields", level=2)
        cf_table = doc.add_table(rows=1, cols=6)
        cf_table.cell(0, 0).text = "Field Name"
        cf_table.cell(0, 1).text = "Type"
        cf_table.cell(0, 2).text = "Required"
        cf_table.cell(0, 3).text = "Values"
        cf_table.cell(0, 4).text = "Default"
        cf_table.cell(0, 5).text = "ID"
        for i in range(custom_field_count):
            _add_field_row(
                cf_table, f"customField{i}", "varchar", "No", "", "—",
                f"TST-DAT-{i:03d}",
            )
            _add_desc_row(cf_table, f"Description for custom field {i}.")

    # Section 4 — Relationships
    if include_section4:
        doc.add_heading("4. Relationships", level=1)
        rel_table = doc.add_table(rows=1, cols=5)
        rel_table.cell(0, 0).text = "Relationship"
        rel_table.cell(0, 1).text = "Related Entity"
        rel_table.cell(0, 2).text = "Link Type"
        rel_table.cell(0, 3).text = "PRD Reference"
        rel_table.cell(0, 4).text = "Domain(s)"
        for i in range(rel_count):
            row = rel_table.add_row()
            row.cells[0].text = f"TestEntity → Related{i}"
            row.cells[1].text = f"Related{i}"
            row.cells[2].text = "oneToMany"
            row.cells[3].text = "Test reference"
            row.cells[4].text = "TST"

    # Section 5 — Dynamic Logic Rules
    if include_section5:
        doc.add_heading("5. Dynamic Logic Rules", level=1)
        doc.add_heading("5.1 Test Rule", level=2)
        doc.add_paragraph("Test dynamic logic rule description.")

    # Section 6 — Layout Guidance
    if include_section6:
        doc.add_heading("6. Layout Guidance", level=1)
        doc.add_paragraph("Test layout guidance.")

    # Section 7 — Implementation Notes
    if include_section7:
        doc.add_heading("7. Implementation Notes", level=1)
        doc.add_paragraph("Test implementation notes.")

    # Section 8 — Open Issues
    if include_section8:
        doc.add_heading("8. Open Issues", level=1)
        oi_table = doc.add_table(rows=2, cols=2)
        oi_table.cell(0, 0).text = "ID"
        oi_table.cell(0, 1).text = "Issue"
        oi_table.cell(1, 0).text = "TST-ISS-001"
        oi_table.cell(1, 1).text = "Test open issue."

    # Section 9 — Decisions Made
    if include_section9:
        doc.add_heading("9. Decisions Made", level=1)
        dec_table = doc.add_table(rows=2, cols=2)
        dec_table.cell(0, 0).text = "ID"
        dec_table.cell(0, 1).text = "Decision"
        dec_table.cell(1, 0).text = "TST-DEC-001"
        dec_table.cell(1, 1).text = "Test decision."

    return doc


# =========================================================================
# Fixture tests against Contact v1.3
# =========================================================================


@skip_fixture
class TestContactFixture:
    """Tests against the real Contact Entity PRD v1.3 fixture."""

    @pytest.fixture()
    def parsed(self):
        envelope_json, report = parse(FIXTURE, _work_item())
        env = json.loads(envelope_json)
        return env, report

    # -- Test 1 --
    def test_parse_succeeds_and_roundtrips(self, parsed):
        env, _report = parsed
        # Round-trip through JSON
        assert json.loads(json.dumps(env)) == env

    # -- Test 2 --
    def test_source_metadata(self, parsed):
        env, _ = parsed
        sm = env["payload"]["source_metadata"]
        assert sm["entity_name"] == "Contact"
        assert sm["version"] == "1.3"

    # -- Test 3 --
    def test_entity_metadata(self, parsed):
        env, _ = parsed
        em = env["payload"]["entity_metadata"]
        assert em["is_native"] is True
        assert em["entity_type"] == "Person"
        assert em["singular_label"] == "Contact"
        assert em["plural_label"] == "Contacts"
        assert em["activity_stream"] is True

    # -- Test 4 --
    def test_primary_domain_code_fallback(self, parsed):
        env, _ = parsed
        em = env["payload"]["entity_metadata"]
        assert em["primary_domain_code"] == "MN"

    # -- Test 5 --
    def test_primary_domain_fallback_warning(self, parsed):
        _, report = parsed
        cats = [w.category for w in report.warnings]
        assert "primary_domain_fallback" in cats

    # -- Test 6 --
    def test_contributing_domain_codes(self, parsed):
        env, _ = parsed
        em = env["payload"]["entity_metadata"]
        assert em["contributing_domain_codes"] == ["MN", "MR", "CR", "FU"]

    # -- Test 7 --
    def test_native_fields_count(self, parsed):
        env, _ = parsed
        assert len(env["payload"]["native_fields"]) == 16

    # -- Test 8 --
    def test_custom_fields_count(self, parsed):
        env, _ = parsed
        # Contact v1.3 has 50 custom fields across all subsections
        assert len(env["payload"]["custom_fields"]) == 50

    # -- Test 9 --
    def test_custom_fields_have_subsection(self, parsed):
        env, _ = parsed
        for f in env["payload"]["custom_fields"]:
            assert f.get("subsection"), f"Field {f['name']} has empty subsection"

    # -- Test 10 --
    def test_incomplete_domain_subsection_warning(self, parsed):
        _, report = parsed
        # Section 3.4 has no field tables, so subsection_no_field_tables fires
        cats = [w.category for w in report.warnings]
        assert "subsection_no_field_tables" in cats
        # Verify the warning mentions "Incomplete Domain"
        matching = [
            w for w in report.warnings
            if w.category == "subsection_no_field_tables"
            and "Incomplete Domain" in w.location
        ]
        assert len(matching) >= 1

    # -- Test 11 --
    def test_relationships_count(self, parsed):
        env, _ = parsed
        assert len(env["payload"]["relationships"]) >= 10

    # -- Test 12 --
    def test_dynamic_logic_count(self, parsed):
        env, _ = parsed
        assert len(env["payload"]["dynamic_logic"]) >= 7

    # -- Test 13 --
    def test_layout_guidance_present(self, parsed):
        env, _ = parsed
        assert env["payload"]["layout_guidance"]

    # -- Test 14 --
    def test_implementation_notes_nonempty(self, parsed):
        env, _ = parsed
        assert isinstance(env["payload"]["implementation_notes"], str)
        assert env["payload"]["implementation_notes"]

    # -- Test 15 --
    def test_open_issues_count(self, parsed):
        env, _ = parsed
        assert len(env["open_issues"]) == 8

    # -- Test 16 --
    def test_decisions_count(self, parsed):
        env, _ = parsed
        assert len(env["decisions"]) == 14


# =========================================================================
# Hard-failure synthetic tests
# =========================================================================


class TestHardFailures:
    """Tests that verify hard-fail conditions raise appropriate errors."""

    # -- Test 17 --
    def test_wrong_item_type_raises_value_error(self, tmp_path):
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        wi = {"id": 1, "item_type": "process_definition", "entity_id": 1}
        with pytest.raises(ValueError, match="must be 'entity_prd'"):
            parse(p, wi)

    # -- Test 18 --
    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            parse("/nonexistent/path.docx", _work_item())

    # -- Test 19 --
    def test_no_header_table(self, tmp_path):
        doc = Document()
        doc.add_paragraph("No tables at all.")
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="No header table"):
            parse(p, _work_item())

    # -- Test 20 --
    def test_missing_entity_row_in_header(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        _add_row(table, "Document Type", "Entity PRD")
        _add_row(table, "Version", "1.0")
        # No "Entity" row
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="missing an 'Entity' row"):
            parse(p, _work_item())

    # -- Test 21 --
    def test_no_entity_overview_table(self, tmp_path):
        doc = Document()
        # Header table
        table = doc.add_table(rows=0, cols=2)
        _add_row(table, "Entity", "Test (Custom — Base)")
        # No overview table — skip straight to sections
        doc.add_heading("1. Entity Overview", level=1)
        doc.add_paragraph("Description.")
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="No Entity Overview table"):
            parse(p, _work_item())

    # -- Test 22 --
    def test_overview_missing_entity_name(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        _add_row(table, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        # First cell should be "CRM Entity Name" but it's something else
        _add_row(ov, "CRM Entity Name", "")
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="missing 'CRM Entity Name'"):
            parse(p, _work_item())

    # -- Test 23 --
    def test_both_domains_absent(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        _add_row(table, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        # No Primary Domain AND no Contributing Domains
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Contributing Domains"):
            parse(p, _work_item())

    # -- Test 24 --
    def test_contributing_domains_empty(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        _add_row(table, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        _add_row(ov, "Contributing Domains", "   ")
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Contributing Domains"):
            parse(p, _work_item())

    # -- Test 25 --
    def test_missing_section1(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section1=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Section 1"):
            parse(p, _work_item())

    # -- Test 26 --
    def test_missing_section2(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section2=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Section 2"):
            parse(p, _work_item())

    # -- Test 27 --
    def test_section2_no_matching_table(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section2=False)
        # Add Section 2 heading but wrong table format
        doc.add_heading("2. Native Fields", level=1)
        wrong = doc.add_table(rows=2, cols=2)
        wrong.cell(0, 0).text = "X"
        wrong.cell(0, 1).text = "Y"
        # Re-add sections 3-9 since build skipped them
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="no matching native fields"):
            parse(p, _work_item())

    # -- Test 28 --
    def test_missing_section3(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section3=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Section 3"):
            parse(p, _work_item())

    # -- Test 29 --
    def test_section3_zero_fields(self, tmp_path):
        doc2 = Document()
        t = doc2.add_table(rows=0, cols=2)
        _add_row(t, "Entity", "Test (Custom — Base)")
        ov = doc2.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        _add_row(ov, "Contributing Domains", "Testing (TST)")
        doc2.add_heading("1. Entity Overview", level=1)
        doc2.add_paragraph("Description.")
        doc2.add_heading("2. Native Fields", level=1)
        nf = doc2.add_table(rows=2, cols=4)
        nf.cell(0, 0).text = "Native Field"
        nf.cell(0, 1).text = "Type"
        nf.cell(0, 2).text = "Mapping"
        nf.cell(0, 3).text = "Ref"
        nf.cell(1, 0).text = "name"
        nf.cell(1, 1).text = "varchar"
        doc2.add_heading("3. Custom Fields", level=1)
        doc2.add_heading("3.1 Empty Subsection", level=2)
        doc2.add_paragraph("No field tables here.")
        doc2.add_heading("4. Relationships", level=1)
        rel = doc2.add_table(rows=2, cols=5)
        rel.cell(0, 0).text = "Relationship"
        rel.cell(0, 1).text = "Related Entity"
        rel.cell(0, 2).text = "Link Type"
        rel.cell(0, 3).text = "PRD Reference"
        rel.cell(0, 4).text = "Domain(s)"
        rel.cell(1, 0).text = "Test → Other"
        rel.cell(1, 1).text = "Other"
        rel.cell(1, 2).text = "oneToMany"
        rel.cell(1, 3).text = "Test"
        rel.cell(1, 4).text = "TST"
        doc2.add_heading("8. Open Issues", level=1)
        oi = doc2.add_table(rows=2, cols=2)
        oi.cell(0, 0).text = "ID"
        oi.cell(0, 1).text = "Issue"
        oi.cell(1, 0).text = "TST-ISS-001"
        oi.cell(1, 1).text = "Test"
        doc2.add_heading("9. Decisions Made", level=1)
        dec = doc2.add_table(rows=2, cols=2)
        dec.cell(0, 0).text = "ID"
        dec.cell(0, 1).text = "Decision"
        dec.cell(1, 0).text = "TST-DEC-001"
        dec.cell(1, 1).text = "Test"
        p = _save(doc2, tmp_path)
        with pytest.raises(EntityPrdParseError, match="zero fields"):
            parse(p, _work_item())

    # -- Test 30 --
    def test_missing_section4(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section4=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Section 4"):
            parse(p, _work_item())

    # -- Test 31 --
    def test_missing_section8(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section8=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Section 8"):
            parse(p, _work_item())

    # -- Test 32 --
    def test_missing_section9(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section9=False)
        p = _save(doc, tmp_path)
        with pytest.raises(EntityPrdParseError, match="Section 9"):
            parse(p, _work_item())


# =========================================================================
# Soft-warning synthetic tests
# =========================================================================


class TestSoftWarnings:
    """Tests that verify soft-warning conditions produce warnings."""

    # -- Test 33 --
    def test_header_entity_no_parens(self, tmp_path):
        doc = _build_minimal_valid_doc(
            entity_value="BareEntity",
            overview_entity_name="BareEntity",
        )
        p = _save(doc, tmp_path)
        envelope_json, report = parse(p, _work_item())
        env = json.loads(envelope_json)
        # Name passes through
        assert env["payload"]["entity_metadata"]["name"] == "BareEntity"
        # is_native is null (header couldn't determine), entity_type comes
        # from overview table ("Base") since header had no parens
        assert env["payload"]["entity_metadata"]["is_native"] is None
        # Warning emitted
        cats = [w.category for w in report.warnings]
        assert "entity_value_no_parens" in cats

    # -- Test 34 --
    def test_native_mismatch_cross_check(self, tmp_path):
        # Header says Native, overview says Custom
        doc = _build_minimal_valid_doc(
            entity_value="TestEntity (Native — Base Type)",
        )
        # The overview says "Custom" by default
        p = _save(doc, tmp_path)
        _, report = parse(p, _work_item())
        cats = [w.category for w in report.warnings]
        assert "native_mismatch" in cats

    # -- Test 35 --
    def test_primary_domain_not_in_contributing(self, tmp_path):
        doc = _build_minimal_valid_doc(
            primary_domain="Other (OTH)",
            contributing_domains="Testing (TST)",
        )
        p = _save(doc, tmp_path)
        _, report = parse(p, _work_item())
        cats = [w.category for w in report.warnings]
        assert "primary_domain_not_in_contributing" in cats

    # -- Test 36 --
    def test_missing_section5_info_note(self, tmp_path):
        doc = _build_minimal_valid_doc(include_section5=False)
        p = _save(doc, tmp_path)
        envelope_json, report = parse(p, _work_item())
        env = json.loads(envelope_json)
        # dynamic_logic is empty list
        assert env["payload"]["dynamic_logic"] == []
        # Info-level note emitted
        info_entries = [
            w for w in report.warnings
            if w.severity == "info"
            and w.category == "optional_section_absent"
        ]
        assert len(info_entries) >= 1

    # -- Test 37 --
    def test_subsection_no_field_tables_warning(self, tmp_path):
        doc = Document()
        t = doc.add_table(rows=0, cols=2)
        _add_row(t, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        _add_row(ov, "Contributing Domains", "Testing (TST)")
        doc.add_heading("1. Entity Overview", level=1)
        doc.add_paragraph("Description.")
        doc.add_heading("2. Native Fields", level=1)
        nf = doc.add_table(rows=2, cols=4)
        nf.cell(0, 0).text = "Native Field"
        nf.cell(0, 1).text = "Type"
        nf.cell(0, 2).text = "Mapping"
        nf.cell(0, 3).text = "Ref"
        nf.cell(1, 0).text = "name"
        nf.cell(1, 1).text = "varchar"
        doc.add_heading("3. Custom Fields", level=1)
        # H2 with no field tables
        doc.add_heading("3.1 Empty Sub", level=2)
        doc.add_paragraph("No tables.")
        # H2 with a field table
        doc.add_heading("3.2 Has Fields", level=2)
        cf = doc.add_table(rows=1, cols=6)
        cf.cell(0, 0).text = "Field Name"
        cf.cell(0, 1).text = "Type"
        cf.cell(0, 2).text = "Required"
        cf.cell(0, 3).text = "Values"
        cf.cell(0, 4).text = "Default"
        cf.cell(0, 5).text = "ID"
        _add_field_row(cf, "testField", "varchar", "No", "", "—", "TST-001")
        _add_desc_row(cf, "Test desc.")
        doc.add_heading("4. Relationships", level=1)
        rel = doc.add_table(rows=2, cols=5)
        rel.cell(0, 0).text = "Relationship"
        rel.cell(0, 1).text = "Related Entity"
        rel.cell(0, 2).text = "Link Type"
        rel.cell(0, 3).text = "PRD Reference"
        rel.cell(0, 4).text = "Domain(s)"
        rel.cell(1, 0).text = "T → O"
        rel.cell(1, 1).text = "O"
        rel.cell(1, 2).text = "oneToMany"
        rel.cell(1, 3).text = "T"
        rel.cell(1, 4).text = "TST"
        doc.add_heading("8. Open Issues", level=1)
        oi = doc.add_table(rows=2, cols=2)
        oi.cell(0, 0).text = "ID"
        oi.cell(0, 1).text = "Issue"
        oi.cell(1, 0).text = "TST-ISS-001"
        oi.cell(1, 1).text = "Test"
        doc.add_heading("9. Decisions Made", level=1)
        dec = doc.add_table(rows=2, cols=2)
        dec.cell(0, 0).text = "ID"
        dec.cell(0, 1).text = "Decision"
        dec.cell(1, 0).text = "TST-DEC-001"
        dec.cell(1, 1).text = "Test"
        p = _save(doc, tmp_path)
        _, report = parse(p, _work_item())
        matching = [
            w for w in report.warnings
            if w.category == "subsection_no_field_tables"
        ]
        assert len(matching) >= 1

    # -- Test 38 --
    def test_invalid_link_type_warning(self, tmp_path):
        doc = Document()
        t = doc.add_table(rows=0, cols=2)
        _add_row(t, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        _add_row(ov, "Contributing Domains", "Testing (TST)")
        doc.add_heading("1. Entity Overview", level=1)
        doc.add_paragraph("Description.")
        doc.add_heading("2. Native Fields", level=1)
        nf = doc.add_table(rows=2, cols=4)
        nf.cell(0, 0).text = "Native Field"
        nf.cell(0, 1).text = "Type"
        nf.cell(0, 2).text = "Mapping"
        nf.cell(0, 3).text = "Ref"
        nf.cell(1, 0).text = "name"
        nf.cell(1, 1).text = "varchar"
        doc.add_heading("3. Custom Fields", level=1)
        doc.add_heading("3.1 Fields", level=2)
        cf = doc.add_table(rows=1, cols=6)
        cf.cell(0, 0).text = "Field Name"
        cf.cell(0, 1).text = "Type"
        cf.cell(0, 2).text = "Required"
        cf.cell(0, 3).text = "Values"
        cf.cell(0, 4).text = "Default"
        cf.cell(0, 5).text = "ID"
        _add_field_row(cf, "f", "varchar", "No", "", "—", "T-001")
        _add_desc_row(cf, "d")
        doc.add_heading("4. Relationships", level=1)
        rel = doc.add_table(rows=2, cols=5)
        rel.cell(0, 0).text = "Relationship"
        rel.cell(0, 1).text = "Related Entity"
        rel.cell(0, 2).text = "Link Type"
        rel.cell(0, 3).text = "PRD Reference"
        rel.cell(0, 4).text = "Domain(s)"
        rel.cell(1, 0).text = "Test → Other"
        rel.cell(1, 1).text = "Other"
        rel.cell(1, 2).text = "invalidLinkType"
        rel.cell(1, 3).text = "Test"
        rel.cell(1, 4).text = "TST"
        doc.add_heading("8. Open Issues", level=1)
        oi = doc.add_table(rows=2, cols=2)
        oi.cell(0, 0).text = "ID"
        oi.cell(0, 1).text = "Issue"
        oi.cell(1, 0).text = "TST-ISS-001"
        oi.cell(1, 1).text = "Test"
        doc.add_heading("9. Decisions Made", level=1)
        dec = doc.add_table(rows=2, cols=2)
        dec.cell(0, 0).text = "ID"
        dec.cell(0, 1).text = "Decision"
        dec.cell(1, 0).text = "TST-DEC-001"
        dec.cell(1, 1).text = "Test"
        p = _save(doc, tmp_path)
        _, report = parse(p, _work_item())
        cats = [w.category for w in report.warnings]
        assert "invalid_link_type" in cats

    # -- Test 39 --
    def test_decision_id_format_warning(self, tmp_path):
        doc = Document()
        t = doc.add_table(rows=0, cols=2)
        _add_row(t, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        _add_row(ov, "Contributing Domains", "Testing (TST)")
        doc.add_heading("1. Entity Overview", level=1)
        doc.add_paragraph("Description.")
        doc.add_heading("2. Native Fields", level=1)
        nf = doc.add_table(rows=2, cols=4)
        nf.cell(0, 0).text = "Native Field"
        nf.cell(0, 1).text = "Type"
        nf.cell(0, 2).text = "Mapping"
        nf.cell(0, 3).text = "Ref"
        nf.cell(1, 0).text = "name"
        nf.cell(1, 1).text = "varchar"
        doc.add_heading("3. Custom Fields", level=1)
        doc.add_heading("3.1 Fields", level=2)
        cf = doc.add_table(rows=1, cols=6)
        cf.cell(0, 0).text = "Field Name"
        cf.cell(0, 1).text = "Type"
        cf.cell(0, 2).text = "Required"
        cf.cell(0, 3).text = "Values"
        cf.cell(0, 4).text = "Default"
        cf.cell(0, 5).text = "ID"
        _add_field_row(cf, "f", "varchar", "No", "", "—", "T-001")
        _add_desc_row(cf, "d")
        doc.add_heading("4. Relationships", level=1)
        rel = doc.add_table(rows=2, cols=5)
        rel.cell(0, 0).text = "Relationship"
        rel.cell(0, 1).text = "Related Entity"
        rel.cell(0, 2).text = "Link Type"
        rel.cell(0, 3).text = "PRD Reference"
        rel.cell(0, 4).text = "Domain(s)"
        rel.cell(1, 0).text = "T → O"
        rel.cell(1, 1).text = "O"
        rel.cell(1, 2).text = "oneToMany"
        rel.cell(1, 3).text = "T"
        rel.cell(1, 4).text = "TST"
        doc.add_heading("8. Open Issues", level=1)
        oi = doc.add_table(rows=2, cols=2)
        oi.cell(0, 0).text = "ID"
        oi.cell(0, 1).text = "Issue"
        oi.cell(1, 0).text = "TST-ISS-001"
        oi.cell(1, 1).text = "Test"
        doc.add_heading("9. Decisions Made", level=1)
        dec = doc.add_table(rows=2, cols=2)
        dec.cell(0, 0).text = "ID"
        dec.cell(0, 1).text = "Decision"
        # Non-matching identifier
        dec.cell(1, 0).text = "BAD-FORMAT"
        dec.cell(1, 1).text = "Bad decision ID."
        p = _save(doc, tmp_path)
        _, report = parse(p, _work_item())
        cats = [w.category for w in report.warnings]
        assert "decision_id_format" in cats


# =========================================================================
# Structural / round-trip tests
# =========================================================================


class TestStructural:
    """Tests for envelope structure and round-trip validity."""

    # -- Test 40 --
    def test_envelope_json_valid(self, tmp_path):
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        envelope_json, _ = parse(p, _work_item())
        env = json.loads(envelope_json)
        assert json.loads(json.dumps(env)) == env

    # -- Test 41 --
    def test_required_top_level_keys(self, tmp_path):
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        envelope_json, _ = parse(p, _work_item())
        env = json.loads(envelope_json)
        assert "output_version" in env
        assert "work_item_type" in env
        assert "work_item_id" in env
        assert "session_type" in env
        assert "payload" in env
        assert "decisions" in env
        assert "open_issues" in env

    # -- Test 42 --
    def test_entity_metadata_required_keys(self, tmp_path):
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        envelope_json, _ = parse(p, _work_item())
        env = json.loads(envelope_json)
        em = env["payload"]["entity_metadata"]
        required = [
            "name", "entity_type", "is_native", "singular_label",
            "plural_label", "activity_stream", "primary_domain_code",
            "contributing_domain_codes", "discriminator_field",
            "discriminator_values", "description",
        ]
        for key in required:
            assert key in em, f"Missing entity_metadata key: {key}"

    # -- Test 43 --
    def test_adapter_does_not_perform_db_lookups(self, tmp_path):
        """Adapter only produces JSON — no conn or DB interaction."""
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        # work_item has no conn key, adapter should not fail
        wi = {"id": 99, "item_type": "entity_prd", "entity_id": 1}
        envelope_json, report = parse(p, wi)
        env = json.loads(envelope_json)
        assert env["work_item_id"] == 99
        assert isinstance(report, ParseReport)

    # -- Test 44 --
    def test_work_item_id_in_envelope(self, tmp_path):
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        wi = {"id": 42, "item_type": "entity_prd", "entity_id": 7}
        envelope_json, _ = parse(p, wi)
        env = json.loads(envelope_json)
        assert env["work_item_id"] == 42
        assert env["work_item_type"] == "entity_prd"
        assert env["session_type"] == "initial"

    # -- Test 45 --
    def test_custom_session_type(self, tmp_path):
        doc = _build_minimal_valid_doc()
        p = _save(doc, tmp_path)
        envelope_json, _ = parse(p, _work_item(), session_type="revision")
        env = json.loads(envelope_json)
        assert env["session_type"] == "revision"

    # -- Test 46 --
    def test_primary_domain_from_explicit_row(self, tmp_path):
        doc = _build_minimal_valid_doc(
            primary_domain="Alpha (ALP)",
            contributing_domains="Alpha (ALP), Beta (BET)",
        )
        p = _save(doc, tmp_path)
        envelope_json, report = parse(p, _work_item())
        env = json.loads(envelope_json)
        em = env["payload"]["entity_metadata"]
        assert em["primary_domain_code"] == "ALP"
        # No fallback warning
        cats = [w.category for w in report.warnings]
        assert "primary_domain_fallback" not in cats

    # -- Test 47 --
    def test_enum_value_options_parsed(self, tmp_path):
        doc = Document()
        t = doc.add_table(rows=0, cols=2)
        _add_row(t, "Entity", "Test (Custom — Base)")
        ov = doc.add_table(rows=0, cols=2)
        _add_row(ov, "CRM Entity Name", "Test")
        _add_row(ov, "Contributing Domains", "Testing (TST)")
        doc.add_heading("1. Entity Overview", level=1)
        doc.add_paragraph("Description.")
        doc.add_heading("2. Native Fields", level=1)
        nf = doc.add_table(rows=2, cols=4)
        nf.cell(0, 0).text = "Native Field"
        nf.cell(0, 1).text = "Type"
        nf.cell(0, 2).text = "Mapping"
        nf.cell(0, 3).text = "Ref"
        nf.cell(1, 0).text = "name"
        nf.cell(1, 1).text = "varchar"
        doc.add_heading("3. Custom Fields", level=1)
        doc.add_heading("3.1 Fields", level=2)
        cf = doc.add_table(rows=1, cols=6)
        cf.cell(0, 0).text = "Field Name"
        cf.cell(0, 1).text = "Type"
        cf.cell(0, 2).text = "Required"
        cf.cell(0, 3).text = "Values"
        cf.cell(0, 4).text = "Default"
        cf.cell(0, 5).text = "ID"
        _add_field_row(
            cf, "status", "enum", "Yes",
            "Active, Inactive, Pending", "Active", "TST-DAT-001",
        )
        _add_desc_row(cf, "Status of the record.")
        doc.add_heading("4. Relationships", level=1)
        rel = doc.add_table(rows=2, cols=5)
        rel.cell(0, 0).text = "Relationship"
        rel.cell(0, 1).text = "Related Entity"
        rel.cell(0, 2).text = "Link Type"
        rel.cell(0, 3).text = "PRD Reference"
        rel.cell(0, 4).text = "Domain(s)"
        rel.cell(1, 0).text = "T → O"
        rel.cell(1, 1).text = "O"
        rel.cell(1, 2).text = "oneToMany"
        rel.cell(1, 3).text = "T"
        rel.cell(1, 4).text = "TST"
        doc.add_heading("8. Open Issues", level=1)
        oi = doc.add_table(rows=2, cols=2)
        oi.cell(0, 0).text = "ID"
        oi.cell(0, 1).text = "Issue"
        oi.cell(1, 0).text = "TST-ISS-001"
        oi.cell(1, 1).text = "Test"
        doc.add_heading("9. Decisions Made", level=1)
        dec = doc.add_table(rows=2, cols=2)
        dec.cell(0, 0).text = "ID"
        dec.cell(0, 1).text = "Decision"
        dec.cell(1, 0).text = "TST-DEC-001"
        dec.cell(1, 1).text = "Test"
        p = _save(doc, tmp_path)
        envelope_json, _ = parse(p, _work_item())
        env = json.loads(envelope_json)
        field = env["payload"]["custom_fields"][0]
        assert field["name"] == "status"
        assert field["value_options"] == ["Active", "Inactive", "Pending"]
        assert field["default_value"] == "Active"
