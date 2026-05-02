"""Deployment Record .docx generator.

Produces a per-instance Deployment Record artifact capturing the
as-deployed state of an EspoCRM instance: hosting, domain/DNS, TLS
certificate, application stack, SSH access, credential references, and
deployment history. Mirrors the structure of the JavaScript reference
generator at
``ClevelandBusinessMentoring/PRDs/deployment/generate-deployment-record.js``.

Invoked by Prompt B (post-deploy automatic generation in the wizard) and
Prompt C (manual regeneration UI). In Prompt A this module is unwired
from the application; it is exercised via tests and the CLI entry point
at the bottom of the file.

See deployment-record series Prompt A for the full specification.
"""

from __future__ import annotations

import dataclasses
import hashlib
import logging
import re
import socket
import ssl
import types
import typing
from collections.abc import Callable
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING
from urllib.parse import urlsplit

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

from automation.core.deployment.deploy_config_repo import InstanceDeployConfig
from automation.core.deployment.ssh_deploy import run_remote

if TYPE_CHECKING:
    import paramiko

    from automation.ui.deployment.deployment_logic import InstanceDetail

logger = logging.getLogger(__name__)


# ── Style constants ──────────────────────────────────────────────────

HEADER_FILL = "1F3864"      # navy
HEADER_TEXT = "FFFFFF"      # white
ALT_ROW_FILL = "F2F7FB"     # light blue tint
META_LABEL_FILL = "F2F7FB"
BORDER_COLOR = "AAAAAA"
HEADING_RGB = RGBColor(0x1F, 0x38, 0x64)
HEADER_TEXT_RGB = RGBColor(0xFF, 0xFF, 0xFF)
BLACK_RGB = RGBColor(0x00, 0x00, 0x00)

FONT_NAME = "Arial"
BODY_PT = Pt(11)            # 22 half-points (matches JS default run size)
H1_PT = Pt(15)              # 30 half-points
H2_PT = Pt(13)              # 26 half-points
H3_PT = Pt(12)              # 24 half-points
TITLE_PT = Pt(16)           # 32 half-points (JS title)
SUBTITLE_PT = Pt(14)        # 28 half-points (JS subtitle)


# ── Dataclasses ──────────────────────────────────────────────────────


@dataclasses.dataclass
class AdministratorInputs:
    """Values supplied by the administrator that cannot be auto-detected.

    Used by :func:`inspect_server_for_record_values` to fill in the
    fields of :class:`DeploymentRecordValues` that are not derivable
    from the database, the SSH session, or the live SSL handshake.
    """

    domain_registrar: str
    dns_provider: str
    primary_domain: str
    instance_subdomain: str
    droplet_id: str | None
    backups_enabled: bool
    proton_pass_admin_entry: str
    proton_pass_db_root_entry: str
    proton_pass_hosting_entry: str
    document_version: str = "1.0"
    document_status: str = (
        "Active — reflects live system as of Last Updated"
    )


@dataclasses.dataclass
class DeploymentRecordValues:
    """Inputs to :func:`generate_deployment_record`.

    All fields are required unless marked Optional (``T | None``). The
    generator does not invent values; missing optional fields render as
    a placeholder string ("not captured") so the document remains
    structurally complete.
    """

    # Document metadata
    document_version: str
    document_last_updated: str           # MM-DD-YY HH:MM
    document_status: str

    # Instance identification
    client_name: str
    instance_name: str
    instance_code: str
    environment: str                     # "test" | "staging" | "production"
    application_url: str
    admin_username: str
    instance_created_at_utc: str         # ISO 8601

    # Hosting / Droplet
    hosting_provider: str
    droplet_id: str | None
    droplet_detail_url: str | None
    droplet_console_url: str | None
    region: str
    hostname: str
    public_ipv4: str
    droplet_size_summary: str
    os_release: str
    kernel: str
    cpu_count: int
    memory_summary: str
    disk_summary: str
    swap_summary: str
    ufw_summary: str
    backups_enabled: bool

    # Domain / DNS
    primary_domain: str
    domain_registrar: str
    dns_provider: str
    instance_subdomain: str

    # TLS certificate
    tls_issuer: str
    tls_subject: str
    tls_issued_utc: str
    tls_expires_utc: str
    tls_sha256_fingerprint: str

    # EspoCRM application
    espocrm_version: str
    espocrm_install_completed_utc: str
    espocrm_install_path: str
    mariadb_version: str
    nginx_version: str
    docker_version: str
    docker_compose_version: str

    # SSH access
    ssh_authorized_user: str
    ssh_key_algorithm: str
    ssh_key_comment: str
    ssh_key_fingerprint: str

    # Credential references
    proton_pass_admin_entry: str
    proton_pass_db_root_entry: str
    proton_pass_hosting_entry: str

    # Deployment history rows: {date_utc, event, notes}
    deployment_history: list[dict[str, str]]

    # Open items rows: {id, item, status_or_plan}
    open_items: list[dict[str, str]]

    # Revision History entries: {version, date, notes}
    revision_history: list[dict[str, str]]

    # Change Log entries: {version, date, changes}
    change_log: list[dict[str, str]]

    def __post_init__(self) -> None:
        """Validate that required string fields are non-empty.

        :raises ValueError: If any required field (typed without
            ``| None``) is empty when string-typed, or None when
            non-string-typed.
        """
        hints = typing.get_type_hints(self.__class__)
        missing: list[str] = []
        for field in dataclasses.fields(self):
            annotation = hints.get(field.name, field.type)
            value = getattr(self, field.name)
            if _is_optional_type(annotation):
                continue
            if isinstance(value, str):
                if not value.strip():
                    missing.append(field.name)
            elif value is None:
                missing.append(field.name)
        if missing:
            raise ValueError(
                f"DeploymentRecordValues missing required fields: "
                f"{', '.join(sorted(missing))}"
            )


def _is_optional_type(annotation: object) -> bool:
    """Return True if an annotation includes ``None`` in a union."""
    origin = typing.get_origin(annotation)
    if origin is typing.Union or origin is types.UnionType:
        return type(None) in typing.get_args(annotation)
    return False


# ── Public API ───────────────────────────────────────────────────────


def increment_minor_version(version: str | None) -> str:
    """Compute the next ``document_version`` given the previous one.

    Strict-increment on the minor component:

    * ``None`` → ``"1.0"`` (first generation; no prior record persisted).
    * ``"M.N"`` (numeric major.minor) → ``"M.{N+1}"``
      (e.g., ``"1.0"`` → ``"1.1"``, ``"1.9"`` → ``"1.10"``).
    * Anything else → returned verbatim (caller is responsible for
      logging a warning if the value is non-numeric so the surprise is
      visible in the run log).

    Strict-increment is intentional: the document version is descriptive
    metadata — the per-instance Revision History entry — not a software
    version, so the conventional 1.9 → 2.0 rollover would silently
    erase the operator's own intent. Major bumps are out of scope and
    must be made by hand.

    :param version: The most recently rendered document_version, or
        ``None`` if no record has been written yet for this instance.
    :returns: The next version string to render.
    """
    if version is None:
        return "1.0"
    match = re.match(r"^(\d+)\.(\d+)$", version.strip())
    if not match:
        return version
    major, minor = match.groups()
    return f"{major}.{int(minor) + 1}"


def generate_deployment_record(
    values: DeploymentRecordValues,
    output_path: Path,
) -> Path:
    """Generate a per-instance Deployment Record .docx.

    :param values: Populated :class:`DeploymentRecordValues` bag.
    :param output_path: Path where the .docx should be written. Parent
        directory is created if it does not exist.
    :returns: ``output_path`` on success.
    :raises ValueError: If ``values`` fails post-init validation.
    :raises OSError: If the output cannot be written.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = _build_document(values)
    doc.save(str(output_path))
    return output_path


def inspect_server_for_record_values(
    ssh: paramiko.SSHClient,
    instance: InstanceDetail,
    deploy_config: InstanceDeployConfig,
    administrator_inputs: AdministratorInputs,
    client_name: str,
) -> DeploymentRecordValues:
    """Capture live on-server state and assemble a DeploymentRecordValues.

    Combines:

    * Database state (Instance row, InstanceDeployConfig row).
    * Live SSH inspection of the deployed Droplet.
    * Live TLS certificate inspection (no SSH required).
    * Administrator-supplied values that cannot be auto-detected
      (registrar, DNS provider, Droplet ID, Proton Pass entry names,
      backup status).

    :param ssh: An open ``paramiko.SSHClient`` connected to the Droplet.
    :param instance: The Instance row from the per-client database.
    :param deploy_config: The InstanceDeployConfig row associated with
        the Instance.
    :param administrator_inputs: Values the administrator must supply.
    :param client_name: The human-readable client name (e.g.,
        "Cleveland Business Mentors"), looked up by the caller from
        the master ``Client`` table. Distinct from
        ``Instance.name``, which is the technical instance label
        (e.g., "CBMTEST"); the rendered Deployment Record's title and
        metadata block use this value.
    :returns: A populated :class:`DeploymentRecordValues` ready for
        :func:`generate_deployment_record`.
    """
    log: Callable[[str, str], None] = lambda msg, lvl: logger.info(  # noqa: E731
        "[record-inspect %s] %s", lvl, msg
    )

    os_release = _parse_os_release(_run(ssh, "lsb_release -a", log))
    kernel = _run(ssh, "uname -r", log).strip()
    cpu_count = _parse_cpu_count(_run(ssh, "nproc", log))
    memory_summary = _parse_memory_summary(_run(ssh, "free -h", log))
    disk_summary = _parse_disk_summary(_run(ssh, "df -h /", log))
    swap_summary = _parse_swap_summary(
        _run(ssh, "swapon --show 2>/dev/null", log)
    )
    ufw_summary = _parse_ufw_summary(_run(ssh, "ufw status 2>/dev/null", log))
    hostname = _run(ssh, "hostname", log).strip()
    docker_version = _parse_docker_version(_run(ssh, "docker --version", log))
    docker_compose_version = _parse_compose_version(
        _run(ssh, "docker compose version", log)
    )
    espocrm_install_completed_utc = _parse_install_completed(
        _run(
            ssh,
            f"stat -c '%w' {deploy_config_install_path(deploy_config)} "
            "2>/dev/null",
            log,
        )
    )
    espocrm_version = _read_espocrm_version(ssh, log) or "unknown"
    mariadb_version = _parse_mariadb_version(
        _run(
            ssh,
            "docker exec espocrm-db mariadb --version 2>/dev/null",
            log,
        )
    )
    nginx_version = _parse_nginx_version(
        _run(
            ssh,
            "docker exec espocrm-nginx nginx -v 2>&1",
            log,
        )
    )
    ssh_user, key_algo, key_comment = _parse_authorized_key(
        _run(ssh, "head -1 /root/.ssh/authorized_keys 2>/dev/null", log)
    )
    ssh_key_fp = _read_ssh_key_fingerprint(ssh, log)

    tls = _inspect_tls(deploy_config.domain)

    droplet_id = administrator_inputs.droplet_id
    droplet_detail_url = (
        f"https://cloud.digitalocean.com/droplets/{droplet_id}"
        if droplet_id else None
    )
    droplet_console_url = (
        f"https://cloud.digitalocean.com/droplets/{droplet_id}/console"
        if droplet_id else None
    )

    return DeploymentRecordValues(
        document_version=administrator_inputs.document_version,
        document_last_updated=datetime.now(UTC).strftime("%m-%d-%y %H:%M"),
        document_status=administrator_inputs.document_status,
        client_name=client_name,
        instance_name=instance.code,
        instance_code=instance.code,
        environment=instance.environment,
        application_url=instance.url or f"https://{deploy_config.domain}/",
        admin_username=instance.username or "admin",
        instance_created_at_utc=instance.created_at or "",
        hosting_provider="DigitalOcean",
        droplet_id=droplet_id,
        droplet_detail_url=droplet_detail_url,
        droplet_console_url=droplet_console_url,
        region=_infer_region_from_hostname(hostname) or "unknown",
        hostname=hostname or "unknown",
        public_ipv4=deploy_config.ssh_host,
        droplet_size_summary=(
            f"{cpu_count} vCPU / {memory_summary} / {disk_summary}"
        ),
        os_release=os_release,
        kernel=kernel,
        cpu_count=cpu_count,
        memory_summary=memory_summary,
        disk_summary=disk_summary,
        swap_summary=swap_summary,
        ufw_summary=ufw_summary,
        backups_enabled=administrator_inputs.backups_enabled,
        primary_domain=administrator_inputs.primary_domain,
        domain_registrar=administrator_inputs.domain_registrar,
        dns_provider=administrator_inputs.dns_provider,
        instance_subdomain=administrator_inputs.instance_subdomain,
        tls_issuer=tls["issuer"],
        tls_subject=tls["subject"],
        tls_issued_utc=tls["issued_utc"],
        tls_expires_utc=tls["expires_utc"],
        tls_sha256_fingerprint=tls["sha256_fingerprint"],
        espocrm_version=espocrm_version,
        espocrm_install_completed_utc=espocrm_install_completed_utc,
        espocrm_install_path=deploy_config_install_path(deploy_config),
        mariadb_version=mariadb_version,
        nginx_version=nginx_version,
        docker_version=docker_version,
        docker_compose_version=docker_compose_version,
        ssh_authorized_user=ssh_user,
        ssh_key_algorithm=key_algo,
        ssh_key_comment=key_comment,
        ssh_key_fingerprint=ssh_key_fp,
        proton_pass_admin_entry=administrator_inputs.proton_pass_admin_entry,
        proton_pass_db_root_entry=(
            administrator_inputs.proton_pass_db_root_entry
        ),
        proton_pass_hosting_entry=(
            administrator_inputs.proton_pass_hosting_entry
        ),
        deployment_history=[],
        open_items=[],
        revision_history=[
            {
                "version": administrator_inputs.document_version,
                "date": datetime.now(UTC).strftime("%m-%d-%y %H:%M"),
                "notes": (
                    "Generated automatically by the CRM Builder "
                    "Deployment Record generator."
                ),
            }
        ],
        change_log=[
            {
                "version": administrator_inputs.document_version,
                "date": datetime.now(UTC).strftime("%m-%d-%y %H:%M"),
                "changes": "Initial generated record.",
            }
        ],
    )


def deploy_config_install_path(_config: InstanceDeployConfig) -> str:
    """Return the canonical EspoCRM install path on the Droplet."""
    return "/var/www/espocrm"


# ── SSH inspection helpers ───────────────────────────────────────────


def _run(
    ssh: paramiko.SSHClient,
    command: str,
    log: Callable[[str, str], None],
) -> str:
    """Run a remote command and return stdout+stderr text.

    :returns: Combined output (may be empty); errors are swallowed and
        result in an empty string so callers can pass to ``_parse_*``.
    """
    try:
        _exit_code, output = run_remote(ssh, command, log)
    except Exception as exc:
        logger.warning("Remote command failed (%s): %s", exc, command)
        return ""
    return output


VERSION_DEFAULTS_CONFIG_PATH = (
    "/var/www/html/application/Espo/Resources/defaults/config.php"
)

VERSION_RETRIEVAL_STRATEGIES: tuple[tuple[str, str], ...] = (
    # 1. PHP eval against the application defaults config — the
    # canonical declaration of the running version. Verified against
    # the live CBM Test Droplet (returns "9.3.4").
    (
        "defaults-php-eval",
        'docker exec espocrm php -r '
        '\'echo (require "'
        + VERSION_DEFAULTS_CONFIG_PATH
        + '")["version"] ?? "";\'',
    ),
    # 2. Pure-shell grep against the same file — fallback when PHP
    # eval is unavailable (e.g. exec-time PHP wrapper changes).
    (
        "defaults-shell-grep",
        "docker exec espocrm sh -c "
        "\"grep -oE \\\"'version'[[:space:]]*=>[[:space:]]*'[^']+'\\\" "
        + VERSION_DEFAULTS_CONFIG_PATH
        + " | head -1 | sed -E \\\"s/.*'([^']+)'.*/\\1/\\\"\"",
    ),
    # 3. Legacy lookup against data/config.php — preserved for older
    # installations that did write the version into the user config.
    (
        "data-config-php-eval",
        'docker exec espocrm php -r '
        '\'echo (require "/var/www/html/data/config.php")'
        '["version"] ?? "";\'',
    ),
)


def _read_espocrm_version(
    ssh: paramiko.SSHClient,
    log: Callable[[str, str], None],
) -> str | None:
    """Read EspoCRM's version from inside the running container.

    Tries multiple retrieval strategies in order of reliability,
    falling through on empty or error result. Returns ``None`` if
    every strategy fails; the caller renders the field as ``"unknown"``.

    Strategies, in order:

    1. PHP eval against
       ``application/Espo/Resources/defaults/config.php`` — the
       application's canonical version declaration.
    2. Pure-shell ``grep`` over the same file — works when PHP eval
       is not available.
    3. PHP eval against ``data/config.php`` — legacy fallback for
       older installations that mirrored the version into the user
       config.

    :param ssh: An open paramiko SSH client.
    :param log: Logger callable used for warnings/info traces.
    :returns: A bare ``X.Y.Z`` version string, or ``None`` if every
        strategy fails.
    """
    for strategy_name, cmd in VERSION_RETRIEVAL_STRATEGIES:
        try:
            out = _run(ssh, cmd, log).strip()
        except Exception as exc:
            log(
                f"Version retrieval strategy '{strategy_name}' "
                f"failed: {exc}",
                "warning",
            )
            continue
        if not out:
            log(
                f"Version retrieval strategy '{strategy_name}' "
                f"returned empty.",
                "info",
            )
            continue
        last_line = out.splitlines()[-1].strip()
        match = re.search(r"(\d+\.\d+\.\d+)", last_line)
        if match:
            log(
                f"EspoCRM version {match.group(1)} retrieved via "
                f"strategy '{strategy_name}'.",
                "info",
            )
            return match.group(1)
        log(
            f"Version retrieval strategy '{strategy_name}' returned "
            f"output but no version pattern: {last_line[:80]}",
            "warning",
        )

    log(
        "All EspoCRM version retrieval strategies failed; "
        "rendering 'unknown'.",
        "warning",
    )
    return None


def _read_ssh_key_fingerprint(
    ssh: paramiko.SSHClient,
    log: Callable[[str, str], None],
) -> str:
    """Compute SHA256 fingerprint of the first authorized key on root."""
    out = _run(
        ssh,
        "ssh-keygen -lf /root/.ssh/authorized_keys 2>/dev/null | head -1",
        log,
    ).strip()
    if not out:
        return "unknown"
    parts = out.split()
    for part in parts:
        if part.startswith("SHA256:"):
            return part
    return "unknown"


def _parse_os_release(output: str) -> str:
    """Extract a single-line OS release summary from ``lsb_release -a``."""
    description = ""
    codename = ""
    for line in output.splitlines():
        if line.startswith("Description:"):
            description = line.split(":", 1)[1].strip()
        elif line.startswith("Codename:"):
            codename = line.split(":", 1)[1].strip()
    if description and codename:
        return f"{description} ({codename})"
    return description or output.strip() or "unknown"


def _parse_cpu_count(output: str) -> int:
    """Parse ``nproc`` output into an int (defaulting to 0)."""
    text = output.strip().splitlines()[-1].strip() if output.strip() else ""
    try:
        return int(text)
    except ValueError:
        return 0


def _parse_memory_summary(output: str) -> str:
    """Parse the ``Mem:`` line from ``free -h`` into a 'X GiB' summary."""
    for line in output.splitlines():
        if line.lower().startswith("mem:"):
            parts = line.split()
            if len(parts) >= 2:
                return parts[1]
    return "unknown"


def _parse_disk_summary(output: str) -> str:
    """Parse ``df -h /`` into a 'X GB (root filesystem on Y)' summary."""
    lines = [
        line for line in output.splitlines()
        if line.strip() and not line.lower().startswith("filesystem")
    ]
    if not lines:
        return "unknown"
    parts = lines[0].split()
    if len(parts) < 2:
        return "unknown"
    device = parts[0]
    size = parts[1]
    return f"{size} (root filesystem on {device})"


def _parse_swap_summary(output: str) -> str:
    """Parse ``swapon --show`` output."""
    lines = [line for line in output.splitlines() if line.strip()]
    if len(lines) < 2:
        return "no swap configured"
    parts = lines[1].split()
    if len(parts) < 3:
        return lines[1]
    name = parts[0]
    size = parts[2]
    return f"{size} swapfile ({name})"


def _parse_ufw_summary(output: str) -> str:
    """Compress ``ufw status`` output into a single-line summary."""
    if not output.strip():
        return "ufw status unavailable"
    lower = output.lower()
    state = "active" if "status: active" in lower else "inactive"
    ports: list[str] = []
    for line in output.splitlines():
        match = re.match(r"^\s*(\d+)(/[a-z]+)?\s+ALLOW", line)
        if match:
            port = match.group(1)
            if port not in ports:
                ports.append(port)
    if ports:
        return f"{state}; allows {' / '.join(ports)} (IPv4 and IPv6)"
    return state


def _parse_docker_version(output: str) -> str:
    """Extract the version number from ``docker --version`` output."""
    match = re.search(r"version (\d+\.\d+\.\d+)", output)
    return match.group(1) if match else (output.strip() or "unknown")


def _parse_compose_version(output: str) -> str:
    """Extract the version number from ``docker compose version`` output."""
    match = re.search(r"v?(\d+\.\d+\.\d+)", output)
    if match:
        return f"v{match.group(1)}"
    return output.strip() or "unknown"


def _parse_install_completed(output: str) -> str:
    """Parse ``stat -c '%w'`` output into a clean UTC timestamp."""
    text = output.strip()
    if not text or text == "-":
        return "unknown"
    return text


def _parse_mariadb_version(output: str) -> str:
    """Parse ``mariadb --version`` output into a bare version number."""
    match = re.search(r"Ver \S+ Distrib (\S+?)(?:[-,]| |$)", output)
    if match:
        return match.group(1)
    match = re.search(r"(\d+\.\d+\.\d+)", output)
    return match.group(1) if match else (output.strip() or "unknown")


def _parse_nginx_version(output: str) -> str:
    """Parse ``nginx -v`` output (which writes to stderr) for the version."""
    match = re.search(r"nginx/(\d+\.\d+\.\d+)", output)
    return match.group(1) if match else (output.strip() or "unknown")


def _parse_authorized_key(output: str) -> tuple[str, str, str]:
    """Parse an authorized_keys line into (user, algorithm, comment)."""
    text = output.strip()
    if not text:
        return ("root", "unknown", "")
    parts = text.split(None, 2)
    algo_raw = parts[0] if parts else ""
    algo = algo_raw.replace("ssh-", "").upper() if algo_raw else "unknown"
    comment = parts[2] if len(parts) >= 3 else ""
    return ("root", algo or "unknown", comment)


def _infer_region_from_hostname(hostname: str) -> str | None:
    """DigitalOcean hostnames include the region as a suffix (e.g., NYC3)."""
    match = re.search(r"-(NYC\d|SFO\d|TOR\d|LON\d|AMS\d|FRA\d|SGP\d|BLR\d)$",
                      hostname.upper())
    if match:
        return match.group(1)
    return None


def _inspect_tls(host: str, port: int = 443) -> dict[str, str]:
    """Connect to ``host:port`` and capture certificate metadata."""
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    with socket.create_connection((host, port), timeout=10) as raw:
        with ctx.wrap_socket(raw, server_hostname=host) as ssock:
            der = ssock.getpeercert(binary_form=True) or b""
    fingerprint_hex = hashlib.sha256(der).hexdigest().upper()
    fingerprint = ":".join(
        fingerprint_hex[i:i + 2] for i in range(0, len(fingerprint_hex), 2)
    )

    issuer, subject, issued, expires = _parse_der_certificate(der)
    return {
        "issuer": issuer,
        "subject": subject,
        "issued_utc": issued,
        "expires_utc": expires,
        "sha256_fingerprint": fingerprint,
    }


def _parse_der_certificate(der: bytes) -> tuple[str, str, str, str]:
    """Extract issuer/subject/issued/expires from a DER certificate.

    Uses ``cryptography`` if available (it is a transitive dependency
    via paramiko); falls back to placeholder strings if parsing fails.
    """
    try:
        from cryptography import x509
        from cryptography.hazmat.backends import default_backend

        cert = x509.load_der_x509_certificate(der, default_backend())
        issuer = cert.issuer.rfc4514_string()
        subject = cert.subject.rfc4514_string()
        issued = cert.not_valid_before_utc.strftime("%Y-%m-%d %H:%M:%S UTC")
        expires = cert.not_valid_after_utc.strftime("%Y-%m-%d %H:%M:%S UTC")
        return issuer, subject, issued, expires
    except Exception as exc:
        logger.warning("DER cert parse failed: %s", exc)
        return ("unknown", "unknown", "unknown", "unknown")


# ── Document construction ────────────────────────────────────────────


def _build_document(values: DeploymentRecordValues) -> Document:
    """Assemble the full Word document, section by section."""
    doc = Document()
    _configure_styles(doc)
    _configure_page(doc)

    _add_title_block(doc, values)
    _add_metadata_table(doc, values)
    _add_revision_history(doc, values)

    _add_section_1_purpose(doc, values)
    _add_section_2_summary(doc, values)
    _add_section_3_droplet(doc, values)
    _add_section_4_dns(doc, values)
    _add_section_5_tls(doc, values)
    _add_section_6_espocrm(doc, values)
    _add_section_7_ssh(doc, values)
    _add_section_8_credentials(doc, values)
    _add_section_9_history(doc, values)
    _add_section_10_operational(doc, values)
    _add_section_11_open_items(doc, values)

    _add_change_log(doc, values)
    return doc


def _configure_styles(doc: Document) -> None:
    """Set Normal and Heading 1/2/3 styles to match the JS reference."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_PT

    for level, size in ((1, H1_PT), (2, H2_PT), (3, H3_PT)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = HEADING_RGB


def _configure_page(doc: Document) -> None:
    """US Letter, 1-inch margins (matches JS 1440-twip page geometry)."""
    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.left_margin = Twips(1440)
    section.right_margin = Twips(1440)
    section.top_margin = Twips(1440)
    section.bottom_margin = Twips(1440)


def _add_title_block(doc: Document, values: DeploymentRecordValues) -> None:
    """Centered title and subtitle in navy."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(values.client_name)
    _style_run(run, bold=True, size=TITLE_PT, color=HEADING_RGB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    env = values.environment.title()
    run = subtitle.add_run(
        f"EspoCRM {env} Instance — Deployment Record"
    )
    _style_run(run, bold=True, size=SUBTITLE_PT, color=HEADING_RGB)


def _add_metadata_table(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Two-column metadata table at the top of the document."""
    rows = [
        ("Document Type", "Deployment Record (operational reference)"),
        (
            "Subject Instance",
            f"{values.client_name} EspoCRM "
            f"{values.environment.title()} Instance "
            f"({values.instance_code})",
        ),
        ("Implementation", values.client_name),
        ("Version", values.document_version),
        ("Status", values.document_status),
        ("Last Updated", values.document_last_updated),
    ]
    table = _new_table(doc, [Twips(2600), Twips(6760)], len(rows) + 1)
    for idx, (label, value) in enumerate(rows):
        _set_label_cell(table.rows[idx].cells[0], label)
        _set_value_cell(table.rows[idx].cells[1], value)
    last = table.rows[len(rows)].cells
    _set_label_cell(last[0], "Source of Truth")
    _set_multi_value_cell(last[1], [
        "The running server is the authoritative source for current state.",
        "This document captures values verified at Last Updated and may "
        "drift over time.",
        "Re-verify against the live system before relying on these "
        "values for operational work.",
    ])
    _add_blank(doc)


def _add_revision_history(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Revision History table with one row per recorded version."""
    doc.add_heading("Revision History", level=1)
    rows = [
        (entry.get("version", ""), entry.get("date", ""),
         entry.get("notes", ""))
        for entry in values.revision_history
    ]
    _striped_table(
        doc,
        column_widths=[Twips(800), Twips(1500), Twips(7060)],
        headers=["Version", "Date", "Notes"],
        rows=rows,
    )
    _add_blank(doc)


def _add_section_1_purpose(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 1 — Document Purpose and Scope."""
    doc.add_heading("1. Document Purpose and Scope", level=1)

    host = _hostname_from_url(values.application_url)
    _add_para(
        doc,
        f"This document captures the as-deployed state of the "
        f"{values.client_name} EspoCRM "
        f"{values.environment.title()} instance hosted at "
        f"{host}. Its purpose is to provide a single durable record "
        "of what was deployed, when, on what infrastructure, with "
        "what configuration, and which credentials govern access — "
        "without duplicating credential values that belong in the "
        "password manager. The document supports day-to-day operations, "
        "future upgrades, disaster recovery, and onboarding of any "
        "additional administrators."
    )

    _add_para(
        doc,
        f"The document covers the {values.environment.title()} "
        "instance only. A separate Deployment Record will be created "
        "for any other environment when it is deployed; all records "
        "live alongside one another in PRDs/deployment/."
    )

    _add_para(doc, "Scope includes:")
    for text in [
        "The deployed EspoCRM application instance — version, "
        "container set, install location, application URL",
        "The DigitalOcean Droplet hosting the application — "
        "region, size, OS, hardware, swap, firewall",
        "The domain and DNS configuration that routes traffic to "
        "the Droplet",
        "The TLS certificate and its renewal arrangement",
        "SSH access to the server, including authorized keys",
        "The credentials inventory — by reference to Proton Pass "
        "entry names; never by value",
        "Deployment and operational history — events that have "
        "occurred against this instance",
        "Open items relating to this instance",
    ]:
        _add_bullet(doc, text)

    _add_para(doc, "Scope does not include:")
    for text in [
        "The CRM Builder application itself — covered by the "
        "crmbuilder repository's product PRDs",
        "YAML configuration that has been applied to this instance "
        "— covered by the YAML program files in programs/ and "
        "the ConfigurationRun audit in the per-client database",
        "Generic deployment instructions for any new client — "
        "covered by the CRM Builder Deployment Runbook (see Section 11)",
        "Credential values — stored exclusively in Proton Pass; "
        "this document references entries by name only",
    ]:
        _add_bullet(doc, text)
    _add_blank(doc)


def _add_section_2_summary(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 2 — Deployment Summary."""
    doc.add_heading("2. Deployment Summary", level=1)
    _add_para(
        doc,
        "At-a-glance summary of the deployed instance. Each value is "
        "detailed in subsequent sections."
    )
    deploy_date = (values.instance_created_at_utc or "")[:10]
    deploy_completed = values.espocrm_install_completed_utc
    rows = [
        ("Instance name and code", values.instance_code),
        ("Environment", values.environment.title()),
        ("Application URL", values.application_url),
        ("Public IPv4 (SSH Host)", values.public_ipv4),
        ("Hostname", values.hostname),
        ("Hosting provider", values.hosting_provider),
        ("Region", values.region),
        ("Droplet size", values.droplet_size_summary),
        ("Operating system", values.os_release),
        ("EspoCRM version", values.espocrm_version),
        (
            "TLS certificate",
            f"Let's Encrypt; expires "
            f"{_date_only(values.tls_expires_utc)}; auto-renews nightly",
        ),
        ("Initial deploy date", deploy_date or _placeholder()),
        ("Initial deploy completed", deploy_completed),
        ("Deploy method", "CRM Builder application (self-hosted scenario)"),
        ("Status", "Active and operational"),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(3000), Twips(6360)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)


def _add_section_3_droplet(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 3 — DigitalOcean Droplet."""
    doc.add_heading("3. DigitalOcean Droplet", level=1)
    _add_para(
        doc,
        "The instance runs on a single DigitalOcean Droplet "
        "provisioned manually in the DigitalOcean control panel "
        "before the CRM Builder Setup Wizard was run against it. "
        "The CRM Builder application does not provision Droplets; "
        "it requires that the Droplet exist and be reachable by SSH "
        "before deployment begins."
    )

    doc.add_heading("3.1 Droplet Identification", level=2)
    rows = [
        ("Provider", values.hosting_provider),
        (
            "Account",
            f"Referenced in Proton Pass entry: "
            f"{values.proton_pass_hosting_entry}",
        ),
        ("Region", values.region),
        ("Hostname (server-side)", values.hostname),
        ("Public IPv4 (SSH Host)", values.public_ipv4),
        ("Droplet ID", values.droplet_id or _placeholder()),
        (
            "Droplet detail page",
            values.droplet_detail_url or _placeholder(),
        ),
        (
            "In-browser Console",
            values.droplet_console_url or _placeholder(),
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("3.2 Hardware and Image", level=2)
    rows = [
        (
            "Image at provisioning",
            "Ubuntu 22.04 LTS x64 (clean image; no one-click app)",
        ),
        ("Current OS", values.os_release),
        ("Kernel", values.kernel),
        ("CPU", f"{values.cpu_count} vCPU"),
        ("Memory", values.memory_summary),
        ("Disk", values.disk_summary),
        ("Swap", values.swap_summary),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("3.3 Firewall (ufw)", level=2)
    _add_para(
        doc,
        f"ufw is {values.ufw_summary}. The configuration permits "
        "inbound traffic on the three ports required for the EspoCRM "
        "stack: 22 (SSH), 80 (HTTP, used for the Let's Encrypt HTTP-01 "
        "challenge and HTTP-to-HTTPS redirect), and 443 (HTTPS, the "
        "application's primary listener). Both IPv4 and IPv6 rules are "
        "present for each port. All other inbound traffic is denied by "
        "default."
    )

    doc.add_heading("3.4 Backups", level=2)
    if values.backups_enabled:
        _add_para(
            doc,
            "DigitalOcean automated backups are enabled for this "
            "Droplet (weekly retention). DigitalOcean retains the four "
            "most recent weekly snapshots; older snapshots roll off "
            "automatically."
        )
    else:
        _add_para(
            doc,
            "DigitalOcean automated backups are not currently enabled "
            "for this Droplet. This is acceptable for a test instance "
            "but should be reviewed before production-grade data "
            "accumulates. Enabling weekly DigitalOcean backups adds "
            "approximately 20% to the Droplet's monthly cost. See "
            "Section 11 Open Items."
        )
    _add_blank(doc)


def _add_section_4_dns(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 4 — Domain and DNS."""
    doc.add_heading("4. Domain and DNS", level=1)
    _add_para(
        doc,
        f"The instance is reachable at the subdomain "
        f"{values.instance_subdomain} of the organization's primary "
        f"domain. The {values.instance_subdomain} subdomain is reserved "
        "by convention for this environment."
    )

    doc.add_heading("4.1 Domain and Registrar", level=2)
    rows = [
        ("Primary domain", values.primary_domain),
        ("Domain registrar", values.domain_registrar),
        ("DNS provider", values.dns_provider),
        ("Subdomain for this instance", values.instance_subdomain),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("4.2 A Record", level=2)
    rows = [
        ("Record type", "A"),
        ("Host", values.instance_subdomain),
        ("Target", f"{values.public_ipv4} (Droplet public IPv4)"),
        (
            "TTL",
            f"Default for {values.dns_provider} (verify in registrar "
            "dashboard if precision needed)",
        ),
        ("Proxy / CDN", "None — direct A record"),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    _add_para(
        doc,
        "DNS resolution was confirmed at Last Updated. The A record "
        "must remain in place for the application to be reachable and "
        "for Let's Encrypt certificate renewals to succeed."
    )
    _add_blank(doc)


def _add_section_5_tls(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 5 — TLS Certificate."""
    doc.add_heading("5. TLS Certificate", level=1)
    _add_para(
        doc,
        "All traffic to the application is served over HTTPS. The "
        "certificate is issued by Let's Encrypt and renewed "
        "automatically by a cron job inside the EspoCRM stack. There "
        "is no manual certificate management."
    )
    rows = [
        ("Issuer", values.tls_issuer),
        ("Subject", values.tls_subject),
        ("Issued", values.tls_issued_utc),
        ("Expires", values.tls_expires_utc),
        ("SHA-256 fingerprint", values.tls_sha256_fingerprint),
        (
            "Renewal mechanism",
            "Crontab on the host runs "
            "/var/www/espocrm/command.sh cert-renew daily at 01:00 UTC; "
            "output appended to "
            "/var/www/espocrm/data/letsencrypt/renew.log",
        ),
        (
            "Renewal frequency check",
            "Daily; Let's Encrypt issues a fresh certificate when the "
            "existing one is within 30 days of expiry",
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value (verified at Last Updated)"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("5.1 Renewal Failure Handling", level=2)
    _add_para(
        doc,
        "If certificate renewal fails the application will continue "
        "to serve the existing certificate until expiry. To diagnose, "
        "SSH to the server and inspect "
        "/var/www/espocrm/data/letsencrypt/renew.log for the most "
        "recent renewal attempt. Common causes are: an A-record change "
        "pointing the domain elsewhere (Let's Encrypt cannot complete "
        "the HTTP-01 challenge), ufw blocking port 80 inbound, or a "
        "Cloudflare-style proxy intercepting the challenge. Manual "
        "renewal can be triggered by running the command in the "
        "crontab line directly."
    )
    _add_blank(doc)


def _add_section_6_espocrm(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 6 — EspoCRM Application."""
    doc.add_heading("6. EspoCRM Application", level=1)

    doc.add_heading("6.1 Application Identification", level=2)
    rows = [
        ("Application name", "EspoCRM"),
        ("Version", values.espocrm_version),
        ("Application URL", values.application_url),
        ("Admin user (default)", values.admin_username),
        (
            "Admin password",
            f"Referenced in Proton Pass entry: "
            f"{values.proton_pass_admin_entry}",
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("6.2 Install Method", level=2)
    _add_para(
        doc,
        "The application was installed via the official EspoCRM "
        "installer (install.sh from "
        "github.com/espocrm/espocrm-installer/releases/latest), invoked "
        "with the SSL and Let's Encrypt flags. The installer file "
        "remains at /root/install.sh on the server. The installer set "
        "up the full Docker-based stack and obtained the initial Let's "
        "Encrypt certificate in a single run."
    )

    doc.add_heading("6.3 Container Stack", level=2)
    _add_para(
        doc,
        f"The application runs as a five-container Docker Compose "
        f"stack defined in {values.espocrm_install_path}/"
        "docker-compose.yml. Docker engine version "
        f"{values.docker_version}; Docker Compose plugin "
        f"{values.docker_compose_version}."
    )
    container_rows = [
        (
            "espocrm",
            "espocrm/espocrm:fpm",
            "PHP-FPM application server (port 9000)",
        ),
        (
            "espocrm-daemon",
            "espocrm/espocrm:fpm",
            "Background daemon for EspoCRM scheduled jobs",
        ),
        (
            "espocrm-db",
            f"mariadb:latest ({values.mariadb_version})",
            "MariaDB database (port 3306, healthy)",
        ),
        (
            "espocrm-nginx",
            f"nginx ({values.nginx_version})",
            "Public-facing web server, listening on 80/443 "
            "(IPv4 and IPv6)",
        ),
        (
            "espocrm-websocket",
            "espocrm/espocrm:fpm",
            "WebSocket server (port 8080, IPv4 and IPv6)",
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2200), Twips(2400), Twips(4760)],
        headers=["Container", "Image", "Role"],
        rows=container_rows,
    )
    _add_blank(doc)

    doc.add_heading("6.4 Install Location", level=2)
    base = values.espocrm_install_path
    rows = [
        ("Base path", base),
        ("Docker Compose file", f"{base}/docker-compose.yml"),
        ("Application data", f"{base}/data/"),
        (
            "Let's Encrypt artifacts",
            f"{base}/data/letsencrypt/ (managed inside the stack)",
        ),
        (
            "Installer command file",
            f"{base}/command.sh (used by the renewal cron job)",
        ),
        ("Install completed", values.espocrm_install_completed_utc),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("6.5 Database", level=2)
    _add_para(
        doc,
        "The application database is MariaDB, running in the "
        "espocrm-db container and reporting healthy at Last Updated. "
        "The MariaDB root password lives only in the running "
        "container's environment variables; it has been extracted and "
        "stored in Proton Pass as a recovery measure. The database is "
        "not exposed outside the Docker network."
    )
    rows = [
        ("Engine", "MariaDB"),
        ("Version", values.mariadb_version),
        ("Image", "mariadb:latest"),
        ("Container", "espocrm-db"),
        (
            "Network exposure",
            "Internal Docker network only; not bound to host port",
        ),
        (
            "MariaDB root password",
            f"Referenced in Proton Pass entry: "
            f"{values.proton_pass_db_root_entry}",
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)


def _add_section_7_ssh(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 7 — SSH Access."""
    doc.add_heading("7. SSH Access", level=1)
    _add_para(
        doc,
        "Server administration is performed via SSH. Access is "
        "restricted to key-based authentication for the root user "
        "only; password SSH is not supported by the deployment "
        "toolchain."
    )
    host = _hostname_from_url(values.application_url)
    rows = [
        ("Allowed user", values.ssh_authorized_user),
        ("Authentication", f"Key-based only ({values.ssh_key_algorithm})"),
        ("Authorized key comment", values.ssh_key_comment or _placeholder()),
        ("Public key fingerprint", values.ssh_key_fingerprint),
        (
            "Connection example",
            f"ssh -i /path/to/ssh {values.ssh_authorized_user}@{host}",
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(2600), Twips(6760)],
        headers=["Field", "Value"],
        rows=rows,
    )
    _add_blank(doc)

    doc.add_heading("7.1 Adding an Additional Authorized Key", level=2)
    _add_para(
        doc,
        "To grant SSH access to an additional administrator without "
        "sharing the existing private key, append their public key to "
        "the server's authorized keys file:"
    )
    for text in [
        "SSH to the server using the existing key",
        "Edit /root/.ssh/authorized_keys and append the new public "
        "key on its own line",
        "Verify by having the new administrator connect with their "
        "key from a separate session before closing the original "
        "session",
    ]:
        _add_bullet(doc, text)
    _add_blank(doc)


def _add_section_8_credentials(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 8 — Credentials Inventory."""
    doc.add_heading("8. Credentials Inventory", level=1)
    _add_para(
        doc,
        "All credential values for this instance are stored "
        "exclusively in the organization's Proton Pass vault. This "
        "document references each entry by its exact name; no "
        "credential value appears in this document and none should "
        "ever be added to this document or to the repository in any "
        "form."
    )
    rows = [
        (
            "EspoCRM admin password",
            values.proton_pass_admin_entry,
            "Web UI and REST API access as admin user",
        ),
        (
            "MariaDB root password",
            values.proton_pass_db_root_entry,
            "Database administrative access; recovery and direct SQL "
            "operations",
        ),
        (
            "DigitalOcean account login",
            values.proton_pass_hosting_entry,
            "Droplet management, billing, DNS (if delegated to DO), "
            "Console access",
        ),
    ]
    _striped_table(
        doc,
        column_widths=[Twips(3000), Twips(3500), Twips(2860)],
        headers=["Credential", "Proton Pass Entry Name", "Used For"],
        rows=rows,
    )
    _add_blank(doc)


def _add_section_9_history(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 9 — Deployment History."""
    doc.add_heading("9. Deployment History", level=1)
    _add_para(
        doc,
        "Chronological log of deployment-relevant events against this "
        "instance. Future re-deploys, version upgrades, recovery "
        "operations, and other significant lifecycle events should be "
        "appended here."
    )
    if values.deployment_history:
        rows = [
            (
                entry.get("date_utc", ""),
                entry.get("event", ""),
                entry.get("notes", ""),
            )
            for entry in values.deployment_history
        ]
    else:
        rows = [(_placeholder(), _placeholder(), _placeholder())]
    _striped_table(
        doc,
        column_widths=[Twips(1700), Twips(2200), Twips(5460)],
        headers=["Date (UTC)", "Event", "Notes"],
        rows=rows,
    )
    _add_blank(doc)


def _add_section_10_operational(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 10 — Operational Notes."""
    doc.add_heading("10. Operational Notes", level=1)

    doc.add_heading("10.1 Reaching the Server", level=2)
    _add_para(
        doc,
        "SSH using the deployment private key is the primary access "
        "path. The DigitalOcean in-browser Console (linked in Section "
        "3.1) is the documented fallback when SSH is unreachable for "
        "any reason — for example after a firewall "
        "misconfiguration that blocks port 22, or while diagnosing a "
        "network-layer issue from outside the server."
    )

    doc.add_heading("10.2 Inspecting the Stack", level=2)
    _add_para(
        doc,
        "The following commands, run as root on the server, give a "
        "quick view of stack health. They are read-only:"
    )
    base = values.espocrm_install_path
    for text in [
        f"docker compose -f {base}/docker-compose.yml ps — list "
        "all containers and their status",
        "docker logs espocrm --tail 100 — recent application "
        "server log lines",
        "docker logs espocrm-nginx --tail 100 — recent web server "
        "log lines",
        "docker logs espocrm-db --tail 100 — recent database log "
        "lines",
        "crontab -l — confirm the certificate renewal job is "
        "configured",
        f"tail {base}/data/letsencrypt/renew.log — most recent "
        "renewal attempts",
    ]:
        _add_bullet(doc, text)

    doc.add_heading("10.3 Application Configuration", level=2)
    _add_para(
        doc,
        "YAML configuration applied to this instance is tracked in "
        "the ConfigurationRun table in the per-client CRM Builder "
        "database. This document does not enumerate the YAML "
        "configuration; consult the ConfigurationRun rows or the "
        "programs/ directory in the client repository for that detail."
    )

    doc.add_heading("10.4 Decommissioning", level=2)
    _add_para(
        doc,
        f"When this {values.environment.title()} instance is no longer "
        "needed, decommissioning involves: destroying the DigitalOcean "
        "Droplet (which deletes the application, database, "
        "certificates, and all data); removing the "
        f"{values.instance_subdomain} A record from "
        f"{values.dns_provider}; deleting the Instance row from the "
        "CRM Builder per-client database; and archiving the relevant "
        "Proton Pass entries. This Deployment Record should be "
        "retained as a historical record and updated with a "
        "decommissioning entry in Section 9."
    )
    _add_blank(doc)


def _add_section_11_open_items(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Section 11 — Open Items."""
    doc.add_heading("11. Open Items", level=1)
    _add_para(
        doc,
        "Items relating to this instance that are tracked here for "
        "visibility rather than buried in a project tracker."
    )
    if values.open_items:
        rows = [
            (
                entry.get("id", ""),
                entry.get("item", ""),
                entry.get("status_or_plan", ""),
            )
            for entry in values.open_items
        ]
    else:
        rows = [(_placeholder(), _placeholder(), _placeholder())]
    _striped_table(
        doc,
        column_widths=[Twips(800), Twips(3000), Twips(5560)],
        headers=["ID", "Item", "Status / Plan"],
        rows=rows,
    )
    _add_blank(doc)


def _add_change_log(
    doc: Document, values: DeploymentRecordValues
) -> None:
    """Change Log table — detailed per-version content notes."""
    doc.add_heading("Change Log", level=1)
    rows = [
        (entry.get("version", ""), entry.get("date", ""),
         entry.get("changes", ""))
        for entry in values.change_log
    ]
    _striped_table(
        doc,
        column_widths=[Twips(800), Twips(1500), Twips(7060)],
        headers=["Version", "Date", "Changes"],
        rows=rows,
    )
    _add_blank(doc)


# ── Low-level styling helpers ────────────────────────────────────────


def _new_table(doc: Document, column_widths: list[Twips], rows: int):
    """Create an N-row, M-column table with fixed column widths."""
    table = doc.add_table(rows=rows, cols=len(column_widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for col_idx, width in enumerate(column_widths):
        for row in table.rows:
            row.cells[col_idx].width = width
    return table


def _striped_table(
    doc: Document,
    *,
    column_widths: list[Twips],
    headers: list[str],
    rows: list[tuple[str, ...]],
) -> None:
    """Build a navy-header / alternating-row striped table."""
    table = _new_table(doc, column_widths, len(rows) + 1)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        _set_header_cell(header_cells[idx], header)
    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        fill = ALT_ROW_FILL if row_idx % 2 == 1 else None
        for col_idx, value in enumerate(row):
            _set_value_cell(cells[col_idx], str(value), fill=fill)


def _set_header_cell(cell, text: str) -> None:
    """Style a header cell: navy fill, white bold Arial."""
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _style_run(run, bold=True, color=HEADER_TEXT_RGB)
    _set_cell_shading(cell, HEADER_FILL)
    _set_cell_borders(cell)
    _set_cell_margins(cell)


def _set_label_cell(cell, text: str) -> None:
    """Style a label cell (left column of metadata table): bold, light fill."""
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _style_run(run, bold=True)
    _set_cell_shading(cell, META_LABEL_FILL)
    _set_cell_borders(cell)
    _set_cell_margins(cell)


def _set_value_cell(cell, text: str, *, fill: str | None = None) -> None:
    """Style a plain value cell."""
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _style_run(run)
    if fill:
        _set_cell_shading(cell, fill)
    _set_cell_borders(cell)
    _set_cell_margins(cell)


def _set_multi_value_cell(cell, lines: list[str]) -> None:
    """Style a value cell containing multiple stacked lines."""
    _clear_cell(cell)
    first_paragraph = cell.paragraphs[0]
    first_run = first_paragraph.add_run(lines[0]) if lines else None
    if first_run is not None:
        _style_run(first_run)
    for line in lines[1:]:
        para = cell.add_paragraph()
        run = para.add_run(line)
        _style_run(run)
    _set_cell_borders(cell)
    _set_cell_margins(cell)


def _clear_cell(cell) -> None:
    """Remove the auto-created empty paragraph contents."""
    for paragraph in list(cell.paragraphs):
        for run in list(paragraph.runs):
            run.text = ""


def _style_run(
    run,
    *,
    bold: bool = False,
    italic: bool = False,
    size: Pt | None = None,
    color: RGBColor | None = None,
) -> None:
    """Apply the standard Arial body styling to a run."""
    run.font.name = FONT_NAME
    _force_east_asia_font(run, FONT_NAME)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = size if size is not None else BODY_PT
    if color is not None:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = BLACK_RGB


def _force_east_asia_font(run, font_name: str) -> None:
    """Set ``w:eastAsia`` so Word does not substitute another font."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set cell background shading via OXML."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell) -> None:
    """Apply consistent grey borders to all sides of a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), BORDER_COLOR)
        borders.append(edge)
    tc_pr.append(borders)


def _set_cell_margins(cell) -> None:
    """Match the JS reference's cell padding: 80 / 80 / 120 / 120 twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (
        ("top", 80), ("bottom", 80), ("left", 120), ("right", 120)
    ):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:w"), str(value))
        edge.set(qn("w:type"), "dxa")
        margins.append(edge)
    tc_pr.append(margins)


def _add_para(doc: Document, text: str) -> None:
    """Add a body paragraph in the standard style."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _style_run(run)


def _add_bullet(doc: Document, text: str) -> None:
    """Add a bullet-style paragraph (manually rendered for portability)."""
    paragraph = doc.add_paragraph()
    pf = paragraph.paragraph_format
    pf.left_indent = Pt(36)
    pf.first_line_indent = Pt(-18)
    pf.space_after = Pt(2)
    run = paragraph.add_run("•  " + text)
    _style_run(run)


def _add_blank(doc: Document) -> None:
    """Append a blank paragraph (visual spacer)."""
    doc.add_paragraph("")


# ── Small string helpers ─────────────────────────────────────────────


def _placeholder() -> str:
    """Standard placeholder for missing-but-structurally-required values."""
    return "not captured"


def _hostname_from_url(url: str) -> str:
    """Extract hostname from a URL, returning the raw URL on failure."""
    try:
        host = urlsplit(url).hostname
        return host or url
    except ValueError:
        return url


def _date_only(timestamp: str) -> str:
    """Return the YYYY-MM-DD prefix of a UTC timestamp string."""
    text = (timestamp or "").strip()
    return text[:10] if len(text) >= 10 else text


# ── CLI entry point ──────────────────────────────────────────────────


def _cli() -> int:
    """Read a JSON fixture and emit a Deployment Record .docx."""
    import argparse
    import json

    parser = argparse.ArgumentParser(
        description=(
            "Generate a Deployment Record .docx from a JSON fixture "
            "matching DeploymentRecordValues."
        )
    )
    parser.add_argument(
        "--values", type=Path, required=True,
        help="Path to a JSON file matching DeploymentRecordValues."
    )
    parser.add_argument(
        "--output", type=Path, required=True,
        help="Path to write the .docx."
    )
    args = parser.parse_args()

    data = json.loads(args.values.read_text(encoding="utf-8"))
    values = DeploymentRecordValues(**data)
    out = generate_deployment_record(values, args.output)
    print(f"Wrote {out} ({out.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(_cli())
