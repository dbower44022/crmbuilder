"""Shared python-docx helper functions for Document Generator templates.

Provides primitives that match the Node.js template helpers: paragraph
creation, heading creation, table cell creation, field tables, requirement
tables, and metadata tables. All formatting constants come from formatting.py.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from automation.docgen.templates.formatting import (
    ALT_ROW_BG,
    BODY_SIZE,
    BORDER_COLOR,
    CELL_MARGIN_BOTTOM,
    CELL_MARGIN_LEFT,
    CELL_MARGIN_RIGHT,
    CELL_MARGIN_TOP,
    DESC_MARGIN_BOTTOM,
    DESC_MARGIN_TOP,
    DRAFT_HEADER_TEXT,
    FIELD_COL_HEADERS,
    FIELD_COL_WIDTHS,
    FONT_NAME,
    GRAY_TEXT_RGB,
    H1_SIZE,
    H2_SIZE,
    H3_SIZE,
    HEADER_BG,
    HEADER_FONT_SIZE,
    HEADER_TEXT_RGB,
    MARGIN_BOTTOM,
    MARGIN_LEFT,
    MARGIN_RIGHT,
    MARGIN_TOP,
    META_KEY_BG,
    PAGE_HEIGHT,
    PAGE_WIDTH,
    SMALL_SIZE,
    TABLE_WIDTH_DXA,
    TITLE_COLOR_RGB,
    XS_SIZE,
)

# ---------------------------------------------------------------------------
# Document creation
# ---------------------------------------------------------------------------

def create_document(is_draft: bool = False) -> Document:
    """Create a new Document with standard page setup and default styles."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE

    # Heading styles
    for level, size, color in [
        ("Heading 1", H1_SIZE, TITLE_COLOR_RGB),
        ("Heading 2", H2_SIZE, TITLE_COLOR_RGB),
        ("Heading 3", H3_SIZE, TITLE_COLOR_RGB),
    ]:
        if level in doc.styles:
            hs = doc.styles[level]
            hs.font.name = FONT_NAME
            hs.font.size = size
            hs.font.bold = True
            hs.font.color.rgb = color

    return doc


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def add_paragraph(
    doc_or_cell,
    text: str = "",
    bold: bool = False,
    size=None,
    color: RGBColor | None = None,
    alignment=None,
    space_after: int | None = None,
    space_before: int | None = None,
    italic: bool = False,
):
    """Add a paragraph with formatted text."""
    p = doc_or_cell.add_paragraph()
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after / 20)  # Convert twips
    if space_before is not None:
        pf.space_before = Pt(space_before / 20)

    if text:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size or BODY_SIZE
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        run.font.italic = italic
    return p


def add_heading(doc, text: str, level: int = 1):
    """Add a heading paragraph."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = TITLE_COLOR_RGB
    return h


def add_labeled_paragraph(doc, label: str, text: str):
    """Add a paragraph with a bold label prefix followed by normal text."""
    p = doc.add_paragraph()
    run_bold = p.add_run(label)
    run_bold.font.name = FONT_NAME
    run_bold.font.size = BODY_SIZE
    run_bold.font.bold = True
    run_text = p.add_run(text)
    run_text.font.name = FONT_NAME
    run_text.font.size = BODY_SIZE
    return p


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right with val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), attrs.get("val", "single"))
        element.set(qn("w:sz"), str(attrs.get("sz", "4")))
        element.set(qn("w:color"), attrs.get("color", BORDER_COLOR))
        element.set(qn("w:space"), "0")
        borders.append(element)
    tcPr.append(borders)


def _set_cell_shading(cell, color: str):
    """Set cell background shading."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """Set cell margins in DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            margins.append(el)
    tcPr.append(margins)


def _set_cell_width(cell, width_dxa: int):
    """Set cell width in DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(width_dxa))
    w.set(qn("w:type"), "dxa")
    tcPr.append(w)


def _apply_standard_borders(cell):
    """Apply standard borders to a cell."""
    border_attrs = {"val": "single", "sz": "4", "color": BORDER_COLOR}
    _set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                     left=border_attrs, right=border_attrs)


def _apply_standard_margins(cell):
    """Apply standard cell margins."""
    _set_cell_margins(cell, CELL_MARGIN_TOP, CELL_MARGIN_BOTTOM,
                      CELL_MARGIN_LEFT, CELL_MARGIN_RIGHT)


def _apply_desc_margins(cell):
    """Apply description row margins."""
    _set_cell_margins(cell, DESC_MARGIN_TOP, DESC_MARGIN_BOTTOM,
                      CELL_MARGIN_LEFT, CELL_MARGIN_RIGHT)


def _format_cell(cell, text: str, bold: bool = False, size=None,
                 color: RGBColor | None = None, shaded: bool = False):
    """Write text into a cell with formatting."""
    # Clear default paragraph
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size or SMALL_SIZE
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

    _apply_standard_borders(cell)
    _apply_standard_margins(cell)

    if shaded:
        _set_cell_shading(cell, ALT_ROW_BG)


def add_header_row(table, headers: list[str], widths: list[int]):
    """Add a header row to a table with standard header formatting."""
    row = table.rows[0] if table.rows else table.add_row()
    for i, (header, width) in enumerate(zip(headers, widths, strict=False)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_NAME
        run.font.size = SMALL_SIZE
        run.font.bold = True
        run.font.color.rgb = HEADER_TEXT_RGB

        _apply_standard_borders(cell)
        _apply_standard_margins(cell)
        _set_cell_shading(cell, HEADER_BG)
        _set_cell_width(cell, width)


def add_data_row(table, values: list[str], widths: list[int],
                 shaded: bool = False, bold_indices: set[int] | None = None,
                 size_overrides: dict[int, object] | None = None,
                 color_overrides: dict[int, RGBColor] | None = None):
    """Add a data row to a table."""
    row = table.add_row()
    bold_indices = bold_indices or set()
    size_overrides = size_overrides or {}
    color_overrides = color_overrides or {}

    for i, (value, width) in enumerate(zip(values, widths, strict=False)):
        cell = row.cells[i]
        _format_cell(
            cell, value,
            bold=i in bold_indices,
            size=size_overrides.get(i),
            color=color_overrides.get(i),
            shaded=shaded,
        )
        _set_cell_width(cell, width)
    return row


def create_table(doc, num_cols: int) -> object:
    """Create a table with the right number of columns."""
    table = doc.add_table(rows=0, cols=num_cols)
    table.autofit = False
    # Add a header row
    table.add_row()
    return table


# ---------------------------------------------------------------------------
# Field table (two-row format matching JS template)
# ---------------------------------------------------------------------------

def add_field_table(doc, fields: list[dict]) -> object:
    """Add a field table in the two-row format.

    Each field produces two rows:
      Row 1: Field Name (bold) | Type | Required | Values | Default | ID (gray)
      Row 2: Description (full-width span, gray font)

    Alternating shading per field pair (odd fields are shaded).
    """
    num_cols = len(FIELD_COL_WIDTHS)
    table = doc.add_table(rows=1, cols=num_cols)
    table.autofit = False

    # Header row
    add_header_row(table, FIELD_COL_HEADERS, FIELD_COL_WIDTHS)

    for idx, field in enumerate(fields):
        shaded = idx % 2 == 1

        # Row 1: data
        values = field.get("values", "\u2014")
        if isinstance(values, list):
            values = ", ".join(str(v) for v in values)
        default = field.get("default_value") or "\u2014"
        field_id = field.get("identifier", field.get("name", ""))

        add_data_row(
            table,
            [
                field.get("label", field.get("name", "")),
                field.get("field_type", ""),
                "Yes" if field.get("is_required") else "No",
                str(values),
                str(default),
                str(field_id),
            ],
            FIELD_COL_WIDTHS,
            shaded=shaded,
            bold_indices={0},
            size_overrides={5: XS_SIZE},
            color_overrides={5: GRAY_TEXT_RGB},
        )

        # Row 2: description spanning all columns
        desc_row = table.add_row()
        # Merge all cells in the description row
        desc_cell = desc_row.cells[0]
        for ci in range(1, num_cols):
            desc_cell.merge(desc_row.cells[ci])

        desc_text = field.get("description", "") or ""
        desc_cell.text = ""
        p = desc_cell.paragraphs[0]
        run = p.add_run(desc_text)
        run.font.name = FONT_NAME
        run.font.size = SMALL_SIZE
        run.font.color.rgb = GRAY_TEXT_RGB

        _apply_standard_borders(desc_cell)
        _apply_desc_margins(desc_cell)
        if shaded:
            _set_cell_shading(desc_cell, ALT_ROW_BG)

    return table


# ---------------------------------------------------------------------------
# Requirement table (two columns: ID + Requirement)
# ---------------------------------------------------------------------------

def add_requirement_table(doc, requirements: list[dict]) -> object:
    """Add a requirement table with ID and Requirement columns."""
    from automation.docgen.templates.formatting import REQ_COL_WIDTHS

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    add_header_row(table, ["ID", "Requirement"], REQ_COL_WIDTHS)

    for idx, req in enumerate(requirements):
        shaded = idx % 2 == 1
        add_data_row(
            table,
            [req.get("identifier", ""), req.get("description", "")],
            REQ_COL_WIDTHS,
            shaded=shaded,
            bold_indices={0},
        )

    return table


# ---------------------------------------------------------------------------
# Two-column table (generic: ID + text)
# ---------------------------------------------------------------------------

def add_two_col_table(doc, h1: str, h2: str, rows: list[tuple[str, str]],
                      widths: list[int] | None = None) -> object:
    """Add a two-column table with header."""
    from automation.docgen.templates.formatting import TWO_COL_WIDTHS
    w = widths or TWO_COL_WIDTHS

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    add_header_row(table, [h1, h2], w)

    for idx, (col1, col2) in enumerate(rows):
        shaded = idx % 2 == 1
        add_data_row(table, [col1, col2], w, shaded=shaded, bold_indices={0})

    return table


# ---------------------------------------------------------------------------
# Metadata table (key-value pairs)
# ---------------------------------------------------------------------------

def add_meta_table(doc, pairs: list[tuple[str, str | list[str]]],
                   widths: list[int] | None = None) -> object:
    """Add a metadata key-value table (no header row, shaded keys).

    Values may be a plain string or a list of strings.  When a list is
    provided, each item is rendered as a bulleted paragraph inside the
    value cell.
    """
    from automation.docgen.templates.formatting import META_COL_WIDTHS_PROCESS
    w = widths or META_COL_WIDTHS_PROCESS

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False

    # Set the table grid column widths so Word respects the requested sizes
    tbl_grid = table._tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        for i, grid_col in enumerate(tbl_grid.findall(qn("w:gridCol"))):
            if i < len(w):
                grid_col.set(qn("w:w"), str(w[i]))

    for key, value in pairs:
        row = table.add_row()
        # Key cell (shaded)
        key_cell = row.cells[0]
        _format_cell(key_cell, key, bold=True, size=SMALL_SIZE)
        _set_cell_shading(key_cell, META_KEY_BG)
        _set_cell_width(key_cell, w[0])

        # Value cell
        val_cell = row.cells[1]
        if isinstance(value, list):
            _format_cell_bulleted(val_cell, value)
        else:
            _format_cell(val_cell, value, size=SMALL_SIZE)
        _apply_standard_borders(val_cell)
        _apply_standard_margins(val_cell)
        _set_cell_width(val_cell, w[1])

    return table


def _format_cell_bulleted(cell, items: list[str]) -> None:
    """Write a list of items as bulleted paragraphs inside a cell."""
    cell.text = ""
    # Use the first (default) paragraph for the first bullet
    for i, item in enumerate(items):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(f"\u2022 {item}")
        run.font.name = FONT_NAME
        run.font.size = SMALL_SIZE
        # Tight spacing between bullets
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)


# ---------------------------------------------------------------------------
# Header/Footer
# ---------------------------------------------------------------------------

def set_header(section, left_text: str, right_text: str):
    """Set page header with left and right aligned text."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()

    run_left = p.add_run(left_text)
    run_left.font.name = FONT_NAME
    run_left.font.size = HEADER_FONT_SIZE
    run_left.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run_tab = p.add_run("\t")
    run_tab.font.name = FONT_NAME

    run_right = p.add_run(right_text)
    run_right.font.name = FONT_NAME
    run_right.font.size = HEADER_FONT_SIZE
    run_right.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Right-align tab stop
    from docx.oxml.ns import qn as _qn
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(_qn("w:val"), "right")
    tab.set(_qn("w:pos"), str(TABLE_WIDTH_DXA))
    tab.set(_qn("w:leader"), "none")
    tabs.append(tab)
    pPr.append(tabs)


def set_footer(section, text: str):
    """Set page footer with centered text."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = XS_SIZE
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def set_draft_header(section):
    """Add a DRAFT indicator to the header."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    run = p.add_run(f"  [{DRAFT_HEADER_TEXT}]")
    run.font.name = FONT_NAME
    run.font.size = HEADER_FONT_SIZE
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
