"""Low-level docx file reading utilities.

Wraps python-docx access into simple data structures (lists of strings,
lists of lists) that parser_logic.py can process.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def load_document(path: str | Path) -> Document:
    """Load a .docx file.

    :param path: Path to the .docx file.
    :returns: python-docx Document object.
    :raises FileNotFoundError: If the file does not exist.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    return Document(str(path))


def extract_paragraphs(doc: Document) -> list[str]:
    """Extract all paragraph text from a document.

    :param doc: python-docx Document.
    :returns: List of paragraph text strings.
    """
    return [p.text for p in doc.paragraphs]


def extract_tables(doc: Document) -> list[list[list[str]]]:
    """Extract all tables from a document as lists of rows of cells.

    :param doc: python-docx Document.
    :returns: List of tables, each table is a list of rows,
        each row is a list of cell text strings.
    """
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


def extract_heading_text(doc: Document) -> list[tuple[str, int | None]]:
    """Extract paragraphs with their heading level.

    :param doc: python-docx Document.
    :returns: List of (text, heading_level) tuples.
        heading_level is None for non-heading paragraphs.
    """
    result = []
    for p in doc.paragraphs:
        level = None
        if p.style and p.style.name:
            style_name = p.style.name
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    pass
        result.append((p.text, level))
    return result


def get_first_table(doc: Document) -> list[list[str]] | None:
    """Get the first table in the document (typically the header table).

    :param doc: python-docx Document.
    :returns: Table rows, or None if no tables exist.
    """
    if not doc.tables:
        return None
    rows = []
    for row in doc.tables[0].rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return rows
