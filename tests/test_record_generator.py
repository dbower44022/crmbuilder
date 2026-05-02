"""Tests for automation.core.deployment.record_generator."""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from automation.core.deployment.record_generator import (
    DeploymentRecordValues,
    _read_espocrm_version,
    generate_deployment_record,
    increment_minor_version,
)

FIXTURE_PATH = (
    Path(__file__).parent
    / "fixtures"
    / "deployment_record_values_cbmtest.json"
)


def _make_fixture_values(**overrides) -> DeploymentRecordValues:
    """Build a DeploymentRecordValues from the CBM Test fixture.

    :param overrides: Optional keyword overrides for any field.
    :returns: A populated DeploymentRecordValues, with overrides applied.
    """
    data = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
    data.update(overrides)
    return DeploymentRecordValues(**data)


# ── Validation ───────────────────────────────────────────────────────


def test_values_validates_required_fields_non_empty():
    """Empty required string fields trigger ValueError listing the field."""
    with pytest.raises(ValueError) as excinfo:
        _make_fixture_values(instance_code="")
    assert "instance_code" in str(excinfo.value)


def test_values_accepts_optional_none():
    """Optional fields (typed `T | None`) may be None."""
    values = _make_fixture_values(
        droplet_id=None,
        droplet_detail_url=None,
        droplet_console_url=None,
    )
    assert values.droplet_id is None
    assert values.droplet_detail_url is None
    assert values.droplet_console_url is None


# ── Generation ───────────────────────────────────────────────────────


def test_generate_writes_file(tmp_path):
    """Generator writes a .docx file of substantive size."""
    values = _make_fixture_values()
    out = tmp_path / "out.docx"
    result = generate_deployment_record(values, out)
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 10_000


def test_generated_docx_opens_in_python_docx(tmp_path):
    """The generated .docx round-trips through python-docx cleanly."""
    values = _make_fixture_values()
    out = tmp_path / "out.docx"
    generate_deployment_record(values, out)
    document = Document(str(out))
    assert len(document.paragraphs) >= 100


def _all_text(document: Document) -> str:
    """Concatenate every paragraph and table-cell string in the document."""
    chunks: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def test_generated_docx_contains_expected_strings(tmp_path):
    """Key fixture values appear in the rendered document."""
    values = _make_fixture_values()
    out = tmp_path / "out.docx"
    generate_deployment_record(values, out)
    text = _all_text(Document(str(out)))

    expected_substrings = [
        values.instance_code,
        values.application_url,
        values.proton_pass_admin_entry,
        values.proton_pass_db_root_entry,
        values.proton_pass_hosting_entry,
        values.espocrm_version,
        values.mariadb_version,
        values.nginx_version,
        values.tls_sha256_fingerprint,
        values.docker_version,
        values.docker_compose_version,
        # Deploy date appears as a YYYY-MM-DD prefix in the summary
        values.instance_created_at_utc[:10],
    ]
    missing = [s for s in expected_substrings if s not in text]
    assert not missing, f"Missing substrings: {missing}"


def test_generated_docx_section_structure(tmp_path):
    """All eleven numbered sections plus Revision History / Change Log."""
    values = _make_fixture_values()
    out = tmp_path / "out.docx"
    generate_deployment_record(values, out)
    text = _all_text(Document(str(out)))

    headings = [
        "Revision History",
        "1. Document Purpose and Scope",
        "2. Deployment Summary",
        "3. DigitalOcean Droplet",
        "4. Domain and DNS",
        "5. TLS Certificate",
        "6. EspoCRM Application",
        "7. SSH Access",
        "8. Credentials Inventory",
        "9. Deployment History",
        "10. Operational Notes",
        "11. Open Items",
        "Change Log",
    ]
    missing = [h for h in headings if h not in text]
    assert not missing, f"Missing section headings: {missing}"


def test_public_ipv4_label_includes_ssh_host(tmp_path):
    """The IPv4 row label is "Public IPv4 (SSH Host)" in §2 and §3.1."""
    values = _make_fixture_values()
    out = tmp_path / "out.docx"
    generate_deployment_record(values, out)
    text = _all_text(Document(str(out)))

    occurrences = text.count("Public IPv4 (SSH Host)")
    assert occurrences >= 2, (
        f"Expected at least 2 occurrences of 'Public IPv4 (SSH Host)' "
        f"(Section 2 and Section 3.1); got {occurrences}."
    )
    bare = text.count("Public IPv4") - text.count("Public IPv4 (SSH Host)")
    assert bare == 0, (
        "Found bare 'Public IPv4' label without the (SSH Host) "
        "parenthetical."
    )


def test_inspect_server_uses_provided_client_name():
    """The client_name parameter overrides Instance.name in the rendered doc."""
    from automation.core.deployment.record_generator import (
        AdministratorInputs,
        inspect_server_for_record_values,
    )

    instance = MagicMock()
    instance.name = "CBMTEST"
    instance.code = "CBMTEST"
    instance.environment = "test"
    instance.url = "https://crm-test.example.com/"
    instance.username = "admin"
    instance.created_at = "2026-04-01T00:00:00+00:00"

    deploy_config = MagicMock()
    deploy_config.ssh_host = "1.2.3.4"
    deploy_config.domain = "crm-test.example.com"

    administrator_inputs = AdministratorInputs(
        domain_registrar="Porkbun",
        dns_provider="Porkbun",
        primary_domain="example.com",
        instance_subdomain="crm-test",
        droplet_id="999",
        backups_enabled=False,
        proton_pass_admin_entry="pp-admin",
        proton_pass_db_root_entry="pp-db",
        proton_pass_hosting_entry="pp-host",
    )

    ssh = MagicMock()

    # Return a non-empty stub so the parsers produce values that pass
    # DeploymentRecordValues post-init validation; the exact contents
    # don't matter for this test, only that client_name flows through.
    def fake_run(_ssh, cmd, _log):
        if "authorized_keys" in cmd and "head" in cmd:
            return "ssh-ed25519 AAAA crm-deploy"
        if "ssh-keygen" in cmd:
            return "256 SHA256:abc root@host (ED25519)"
        if "uname -r" in cmd:
            return "5.15.0-generic"
        return "stub"

    with patch(
        "automation.core.deployment.record_generator._run",
        side_effect=fake_run,
    ), patch(
        "automation.core.deployment.record_generator._inspect_tls",
        return_value={
            "issuer": "Let's Encrypt",
            "subject": "x",
            "issued_utc": "2026-04-01 00:00:00 UTC",
            "expires_utc": "2026-07-01 00:00:00 UTC",
            "sha256_fingerprint": "AA:BB",
        },
    ):
        values = inspect_server_for_record_values(
            ssh,
            instance,
            deploy_config,
            administrator_inputs,
            client_name="Cleveland Business Mentors",
        )

    assert values.client_name == "Cleveland Business Mentors"
    assert instance.name == "CBMTEST"


def test_espocrm_version_strategy_first_returns_version():
    """First strategy succeeding returns the parsed version."""
    log_calls: list[tuple[str, str]] = []

    def log(msg: str, lvl: str) -> None:
        log_calls.append((msg, lvl))

    with patch(
        "automation.core.deployment.record_generator._run",
        return_value="9.3.4\n",
    ) as mock_run:
        result = _read_espocrm_version(MagicMock(), log)

    assert result == "9.3.4"
    assert mock_run.call_count == 1


def test_espocrm_version_strategy_falls_through_to_second():
    """Empty first strategy falls through; second-strategy hit returns version."""
    log_calls: list[tuple[str, str]] = []

    def log(msg: str, lvl: str) -> None:
        log_calls.append((msg, lvl))

    outputs = iter(["", "9.3.4\n"])

    def fake_run(_ssh, _cmd, _log):
        return next(outputs)

    with patch(
        "automation.core.deployment.record_generator._run",
        side_effect=fake_run,
    ):
        result = _read_espocrm_version(MagicMock(), log)

    assert result == "9.3.4"


def test_espocrm_version_all_strategies_empty_returns_none():
    """Every strategy returning empty yields None."""
    log_calls: list[tuple[str, str]] = []

    def log(msg: str, lvl: str) -> None:
        log_calls.append((msg, lvl))

    with patch(
        "automation.core.deployment.record_generator._run",
        return_value="",
    ):
        result = _read_espocrm_version(MagicMock(), log)

    assert result is None
    # Final summary log line must appear at warning level.
    assert any(
        lvl == "warning" and "All EspoCRM version retrieval" in msg
        for msg, lvl in log_calls
    )


def test_generated_docx_validates_with_office_validator_when_available(
    tmp_path,
):
    """Run the Anthropic office validator if available; otherwise skip."""
    validator = Path("/mnt/skills/public/docx/scripts/office/validate.py")
    if not validator.exists():
        pytest.skip("Office validator not available in this environment")

    values = _make_fixture_values()
    out = tmp_path / "out.docx"
    generate_deployment_record(values, out)
    result = subprocess.run(
        [sys.executable, str(validator), str(out)],
        capture_output=True,
        text=True,
        check=False,
    )
    assert result.returncode == 0, (
        f"Office validator failed: {result.stdout}\n{result.stderr}"
    )


# ── CLI ──────────────────────────────────────────────────────────────


def test_cli_round_trip(tmp_path):
    """Running the module via -m produces a .docx from the fixture."""
    out = tmp_path / "cli.docx"
    result = subprocess.run(
        [
            sys.executable,
            "-m",
            "automation.core.deployment.record_generator",
            "--values",
            str(FIXTURE_PATH),
            "--output",
            str(out),
        ],
        capture_output=True,
        text=True,
        check=False,
        cwd=str(Path(__file__).resolve().parents[1]),
    )
    assert result.returncode == 0, (
        f"CLI exited {result.returncode}: {result.stdout}\n{result.stderr}"
    )
    assert out.exists()
    assert out.stat().st_size > 10_000


# ── Prompt I FU-2: increment_minor_version helper ───────────────────


def test_increment_minor_version_initial():
    """A NULL/None previous version starts the document at 1.0."""
    assert increment_minor_version(None) == "1.0"


def test_increment_minor_version_basic():
    """Standard minor increment: 1.0 → 1.1, 2.3 → 2.4."""
    assert increment_minor_version("1.0") == "1.1"
    assert increment_minor_version("2.3") == "2.4"
    assert increment_minor_version("1.5") == "1.6"


def test_increment_minor_version_handles_double_digit():
    """1.9 → 1.10 (strict-increment, not semver rollover)."""
    assert increment_minor_version("1.9") == "1.10"
    assert increment_minor_version("1.10") == "1.11"


def test_increment_minor_version_passes_through_non_numeric():
    """Non-matching strings are returned unchanged so the caller can warn."""
    assert increment_minor_version("draft") == "draft"
    assert increment_minor_version("1.0-rc1") == "1.0-rc1"
    assert increment_minor_version("v2.0") == "v2.0"
    # Empty string is non-matching too
    assert increment_minor_version("") == ""


def test_increment_minor_version_strips_whitespace():
    """Surrounding whitespace doesn't break the parse."""
    assert increment_minor_version("  1.2  ") == "1.3"
